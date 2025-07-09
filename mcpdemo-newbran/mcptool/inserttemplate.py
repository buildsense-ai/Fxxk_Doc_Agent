#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FastAPI 应用：文档处理工具箱
提供独立的、可作为 MCP 工具使用的文档处理接口。
"""

import os
import re
import logging
import subprocess
import shutil
from typing import List, Dict, Any
from pathlib import Path
import uuid
from datetime import datetime
import json
import requests
from urllib.parse import urlparse
import tempfile

from fastapi import FastAPI, UploadFile, File, HTTPException, Query, Body
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field
from docx import Document
from mcp.server.fastmcp import FastMCP
from minio import Minio

# --- 日志配置 ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# --- FastAPI 应用设置 ---
# 为工具API使用专用的标题和描述
app = FastAPI(
    title="文档处理工具箱 API",
    description="提供一系列独立的文档处理工具，例如文件格式转换、模板分析等。可以被 MCP (Model-as-a-Service Consumer Platform) 或其他应用集成。",
    version="1.0.0"
)

# --- 临时文件目录 ---
TEMP_DIR = Path("temp_tool_files")
TEMP_DIR.mkdir(exist_ok=True)

mcp = FastMCP("template tools")

# --- Pydantic 模型定义 (用于请求和响应) ---

class DocConversionResponse(BaseModel):
    """DOC转换工具的响应模型（用于成功时的JSON消息）"""
    message: str = Field("文件转换成功，请查收下载的文件。", description="操作结果消息。")
    original_filename: str = Field(..., description="上传的原始 .doc 文件名。")
    converted_filename: str = Field(..., description="转换后的 .docx 文件名。")


class GenerateDocRequest(BaseModel):
    """从JSON生成文档的请求模型"""
    filename: str = Field(..., description="希望生成的文档名，无需文件后缀。", example="monthly_report")
    data: Dict[str, Any] = Field(..., description="用于填充文档内容的JSON对象。")


class GenerateDocResponse(BaseModel):
    """从JSON生成文档的响应模型"""
    message: str = Field("文档生成并上传成功。", description="操作结果消息。")
    filename: str = Field(..., description="在对象存储中保存的文件名。")
    url: str = Field(..., description="文件的访问URL。",
                     example="https://minio.example.com/documents/monthly_report_20250625162300.docx")


# --- 核心工具逻辑 (从 main.py 中抽离并优化) ---

@mcp.tool()
def convert_doc_to_docx(doc_path: str) -> dict:
    """
    使用 LibreOffice 将 .doc 文件转换为 .docx 文件。
    Args:
        doc_path: 输入的 .doc 文件路径。
    Returns:
        包含转换结果的字典。
    """
    doc_path = Path(doc_path)
    if not doc_path.exists():
        return {"error": f"文件不存在: {doc_path}"}

    try:
        # 定义输出路径
        output_dir = doc_path.parent
        docx_path = output_dir / f"{doc_path.stem}.docx"

        # 检查 LibreOffice
        libreoffice_cmd = None
        for cmd in ['libreoffice', 'soffice']:
            try:
                result = subprocess.run([cmd, '--version'], capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    libreoffice_cmd = cmd
                    break
            except:
                continue

        if not libreoffice_cmd:
            return {"error": "LibreOffice未安装或无法访问"}

        # 执行转换
        cmd = [
            libreoffice_cmd,
            '--headless',
            '--convert-to', 'docx',
            '--outdir', str(output_dir),
            str(doc_path)
        ]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            return {"error": f"转换失败: {result.stderr}"}

        if not docx_path.exists():
            return {"error": "转换命令执行成功但未找到输出文件"}

        return {
            "success": True,
            "original_filename": doc_path.name,
            "converted_filename": docx_path.name,
            "converted_path": str(docx_path)
        }

    except Exception as e:
        return {"error": f"转换过程中发生错误: {str(e)}"}


@mcp.tool()
def generate_doc_from_json(filename: str, data: dict) -> dict:
    """
    根据JSON数据生成Word文档。
    Args:
        filename: 希望生成的文档名（无需后缀）
        data: 用于填充文档的数据字典
    Returns:
        包含生成结果的字典
    """
    try:
        # 创建临时目录
        temp_dir = Path("temp_docx_files")
        temp_dir.mkdir(exist_ok=True)
        session_id = str(uuid.uuid4())
        session_dir = temp_dir / session_id
        session_dir.mkdir()

        # 生成文档
        doc = Document()
        doc.add_heading(f"数据报告: {filename}", level=0)
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        for key, value in data.items():
            doc.add_heading(str(key), level=2)
            if isinstance(value, (dict, list)):
                doc.add_paragraph(json.dumps(value, ensure_ascii=False, indent=2))
            else:
                doc.add_paragraph(str(value))

        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        final_filename = f"{filename}_{timestamp}.docx"
        temp_doc_path = session_dir / final_filename
        doc.save(temp_doc_path)

        # 上传到MinIO
        minio_url = upload_to_local_minio(temp_doc_path, final_filename)

        result = {
            "success": True,
            "filename": final_filename,
            "url": minio_url,
            "local_path": str(temp_doc_path)
        }

        return result

    except Exception as e:
        return {"error": f"生成文档时发生错误: {str(e)}"}
    finally:
        # 清理临时目录
        if 'session_dir' in locals() and session_dir.exists():
            shutil.rmtree(session_dir)


def download_template_from_url(template_url: str) -> Path:
    """从URL下载模板文件"""
    try:
        response = requests.get(template_url, timeout=30)
        response.raise_for_status()
        
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(response.content)
        temp_file.close()
        
        return Path(temp_file.name)
    except Exception as e:
        raise Exception(f"下载模板失败: {str(e)}")

def fill_template_with_data(template_path: Path, data: dict) -> Path:
    """使用JSON数据填充模板，参考feishu/main.py的复杂填充逻辑"""
    try:
        doc = Document(template_path)
        filled_count = 0
        
        # 1. 准备附件信息
        attachments_map = data.pop('attachments_map', {})
        attachment_ref_map = {}
        ordered_attachments = []
        if attachments_map and isinstance(attachments_map, dict):
            logger.info(f"🖼️  找到 {len(attachments_map)} 个图片附件待处理。")
            ordered_attachments = list(attachments_map.items())
            for i, (key, _) in enumerate(ordered_attachments):
                attachment_ref_map[key.strip()] = i + 1
        else:
            attachments_map = {}

        # 2. 分离文本填充数据
        placeholder_data = {k: v for k, v in data.items() if k.startswith(('label_', 'inline_', 'blank_'))}
        structure_data = {k: v for k, v in data.items() if k.startswith(('table_', 'paragraph_'))}
        
        # 3. 替换所有文本占位符（包括图片引用）
        image_placeholder_pattern = re.compile(r'\{\{image:([^}]+)\}\}')
        text_placeholder_pattern = re.compile(r'\{(label_[^}]+|inline_[^}]+|blank_[^}]+)\}')

        def process_element_text(element):
            nonlocal filled_count
            if '{' not in element.text:
                return

            original_text = element.text
            new_text = ""
            last_end = 0

            # 创建一个组合的正则表达式来查找所有类型的占位符
            combined_pattern = re.compile(f"({image_placeholder_pattern.pattern}|{text_placeholder_pattern.pattern})")
            
            for match in combined_pattern.finditer(original_text):
                new_text += original_text[last_end:match.start()]
                last_end = match.end()
                
                image_key_match = image_placeholder_pattern.match(match.group(0))
                text_key_match = text_placeholder_pattern.match(match.group(0))

                if image_key_match:
                    key = image_key_match.group(1).strip()
                    if key in attachment_ref_map:
                        number = attachment_ref_map[key]
                        replacement = f"（详见附件{number}）"
                        new_text += replacement
                        logger.info(f"   🖼️  图片引用替换: '{match.group(0)}' -> '{replacement}'")
                    else:
                        logger.warning(f"   ⚠️  找到图片占位符 {match.group(0)} 但无匹配图片，已移除。")
                
                elif text_key_match:
                    placeholder_key = text_key_match.group(1)
                    placeholder = f"{{{placeholder_key}}}"

                    if placeholder_key in placeholder_data:
                        value = str(placeholder_data[placeholder_key])
                        new_text += value
                        logger.info(f"   ✏️  占位符填充: {placeholder} -> {value[:50]}...")
                        filled_count += 1
                    else: # 未匹配的文本占位符
                        if placeholder_key.startswith('label_'):
                            logger.info(f"   🔘  移除未匹配标签占位符: {placeholder}")
                            # The replacement is empty string, so we add nothing
                        elif placeholder_key.startswith(('inline_', 'blank_')):
                            original_underscore = "____"  # 默认下划线
                            new_text += original_underscore
                            logger.info(f"   🔘  恢复未匹配占位符: {placeholder} -> '{original_underscore}'")

            new_text += original_text[last_end:]
            
            if new_text != original_text:
                element.text = new_text

        # 遍历段落和表格进行统一替换
        for para in doc.paragraphs:
            process_element_text(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_element_text(cell)

        # 4. 填充原始结构
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    cell_key = f"table_{i}_row_{j}_col_{k}"
                    if cell_key in structure_data:
                        cell.text = str(structure_data[cell_key])
                        logger.info(f"   ✏️  结构填充(表格): {cell_key} -> {str(structure_data[cell_key])[:50]}...")
                        filled_count += 1

        for i, para in enumerate(doc.paragraphs):
            para_key = f"paragraph_{i}"
            if para_key in structure_data:
                # 只有在段落中没有占位符的情况下才进行结构填充
                if not re.search(r'\{.*\}', para.text):
                    para.text = str(structure_data[para_key])
                    logger.info(f"   ✏️  结构填充(段落): {para_key} -> {str(structure_data[para_key])[:50]}...")
                    filled_count += 1

        # 5. 将图片作为附件附加到文档末尾
        if ordered_attachments:
            logger.info("📎 开始在文档末尾附加图片...")
            try:
                from docx.shared import Inches
                doc.add_page_break()
                doc.add_heading('附件列表', level=1)
                
                for i, (key, image_path) in enumerate(ordered_attachments):
                    attachment_counter = i + 1
                    if not image_path or not isinstance(image_path, str) or not os.path.exists(image_path):
                        logger.warning(f"⚠️ 图片路径不存在或无效，跳过附件 '{key}': {image_path}")
                        continue
                    
                    try:
                        heading_text = f"附件 {attachment_counter}: {key}"
                        doc.add_heading(heading_text, level=2)
                        doc.add_picture(image_path, width=Inches(6.0))
                        doc.add_paragraph()
                        logger.info(f"   ✅ 成功附加图片: {heading_text} ({image_path})")
                    except Exception as pic_e:
                        logger.error(f"❌ 附加图片 '{key}' ({image_path}) 时出错: {pic_e}")
            except Exception as e:
                logger.error(f"❌ 处理附件时发生意外错误: {e}")
        
        # 保存填充后的文档
        output_path = template_path.parent / f"filled_{template_path.name}"
        doc.save(output_path)
        logger.info(f"✅ 混合模式填充完成，共填充 {filled_count} 个字段: {output_path}")
        
        return output_path
        
    except Exception as e:
        raise Exception(f"填充模板失败: {str(e)}")

def upload_to_local_minio(file_path: Path, filename: str) -> str:
    """上传文件到MinIO服务器"""
    try:
        # MinIO配置
        minio_config = {
            "endpoint": "43.139.19.144:9000",  # MinIO API端口（不是控制台端口9001）
            "access_key": "minioadmin",         # 访问密钥
            "secret_key": "minioadmin",         # 秘密密钥
            "bucket": "templates",              # 存储桶名
            "secure": False                     # 是否使用HTTPS
        }
        
        # 使用MinIO Python客户端上传
        client = Minio(
            minio_config["endpoint"],
            access_key=minio_config["access_key"],
            secret_key=minio_config["secret_key"],
            secure=minio_config["secure"]
        )
        
        # 确保存储桶存在
        if not client.bucket_exists(minio_config["bucket"]):
            client.make_bucket(minio_config["bucket"])
        
        # 上传文件
        client.fput_object(
            minio_config["bucket"], 
            filename, 
            str(file_path)
        )
        
        # 返回访问URL
        return f"http://{minio_config['endpoint']}/{minio_config['bucket']}/{filename}"
        
    except ImportError:
        raise Exception("MinIO客户端库未安装，请运行: pip install minio")
    except Exception as e:
        raise Exception(f"上传到MinIO失败: {str(e)}")

@mcp.tool()
def fill_template_from_url(template_url: str, data: dict, output_filename: str = None) -> dict:
    """
    从URL获取模板，填充JSON数据，上传到MinIO并返回地址。
    参考feishu/main.py的复杂填充逻辑。
    Args:
        template_url (str): 模板文件的URL地址，支持http/https链接
                          示例: "https://example.com/template.docx"
        
        data (dict): 用于填充模板的JSON数据字典，支持以下格式：
                    - 文本占位符: {"label_name": "值", "inline_hint": "值", "blank_0": "值"}
                    - 图片占位符: {"attachments_map": {"image_key": "/path/to/image.png"}}
                    - 结构数据: {"table_0_row_1_col_2": "表格内容", "paragraph_3": "段落内容"}
                    示例: {
                        "project_name": "AI文档生成项目",
                        "attachments_map": {
                            "shiGongTu": "/path/to/construction_drawing.png",
                            "xianChangZhaoPian": "/path/to/site_photo.jpg"
                        }
                    }
        
        output_filename (str, optional): 输出文件名，如果不提供则自动生成
                                       示例: "项目报告.docx"
    
    Returns:
        dict: 包含处理结果的字典
            - success (bool): 是否成功
            - template_url (str): 原始模板URL
            - output_filename (str): 输出文件名
            - minio_url (str): MinIO文件访问地址
            - message (str): 处理结果消息
            - filled_fields_count (int): 填充的字段数量
    
    Raises:
        Exception: 当模板下载、填充或上传过程中发生错误时抛出
    
    Example:
        >>> result = fill_template_from_url(
        ...     template_url="https://example.com/template.docx",
        ...     data={
        ...         "project_name": "测试项目",
        ...         "attachments_map": {"image1": "/path/to/image.png"}
        ...     },
        ...     output_filename="测试报告.docx"
        ... )
        >>> print(result["minio_url"])
        "http://43.139.19.144:9001/templates/测试报告.docx"
    
    Features:
        - 支持多种占位符格式：{{image:key}}、{label_key}、{inline_key}、{blank_key}
        - 自动处理图片附件，将图片附加到文档末尾
        - 智能表格和段落填充
        - 自动上传到MinIO存储
        - 完善的错误处理和日志记录
    """
    try:
        # 1. 下载模板
        template_path = download_template_from_url(template_url)
        
        # 2. 填充数据（使用复杂的填充逻辑）
        filled_doc_path = fill_template_with_data(template_path, data)
        
        # 3. 生成输出文件名
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            output_filename = f"filled_template_{timestamp}.docx"
        elif not output_filename.endswith('.docx'):
            output_filename += '.docx'
        
        # 4. 上传到MinIO
        minio_url = upload_to_local_minio(filled_doc_path, output_filename)
        
        # 5. 清理临时文件
        template_path.unlink(missing_ok=True)
        filled_doc_path.unlink(missing_ok=True)
        
        return {
            "success": True,
            "template_url": template_url,
            "output_filename": output_filename,
            "minio_url": minio_url,
            "message": "模板填充并上传成功",
            "filled_fields_count": len(data) - (1 if 'attachments_map' in data else 0)
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "template_url": template_url
        }


# --- API 端点定义 ---

@app.post(
    "/tools/convert_doc_to_docx",
    summary="转换 .doc 文件为 .docx",
    description="上传一个旧版的 .doc 文件，此工具会使用 LibreOffice 将其转换为现代的 .docx 格式并返回。服务器需要安装 LibreOffice。",
    operation_id="convert_doc_to_docx",
    tags=["文件处理工具"]
)
async def tool_convert_doc_to_docx(
        file: UploadFile = File(..., description="要转换的 .doc 文件。")
):
    """
    接收 .doc 文件，转换为 .docx 并提供下载。
    """
    if not file.filename.lower().endswith('.doc'):
        raise HTTPException(status_code=400, detail="文件格式错误，请上传 .doc 文件。")

    # 安全地保存上传的文件
    session_id = str(uuid.uuid4())
    session_dir = TEMP_DIR / session_id
    session_dir.mkdir()

    original_path = session_dir / file.filename
    try:
        with open(original_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 调用核心转换逻辑
        converted_path = convert_doc_to_docx(original_path)

        # 提供转换后的文件下载
        return FileResponse(
            path=str(converted_path),
            filename=converted_path.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except RuntimeError as e:
        # 捕获 LibreOffice 相关的错误
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {e}")
    except Exception as e:
        logger.error(f"处理 /tools/convert_doc_to_docx 时发生意外错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"处理文件时发生意外错误: {e}")
    finally:
        # 清理临时会话目录
        if session_dir.exists():
            shutil.rmtree(session_dir)


@app.post(
    "/tools/generate_doc_from_json",
    response_model=GenerateDocResponse,
    summary="根据JSON数据生成文档并上传",
    description="接收JSON格式的数据，在服务器端根据该数据动态生成一个简单的Word文档，将其（模拟）上传到对象存储（如MinIO），并返回可访问的URL。",
    operation_id="generate_document_from_json_and_upload",
    tags=["文档生成工具"]
)
async def tool_generate_and_upload(
        request_data: GenerateDocRequest = Body(...)
):
    """
    从JSON数据创建Word文档，上传到MinIO，并返回URL。
    """
    session_id = str(uuid.uuid4())
    session_dir = TEMP_DIR / session_id
    session_dir.mkdir()

    try:
        # 1. 在服务器上根据JSON动态生成一个简单的.docx文件
        doc = Document()
        doc.add_heading(f"数据报告: {request_data.filename}", level=0)
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        for key, value in request_data.data.items():
            doc.add_heading(str(key), level=2)
            # 将各种类型的值转换为字符串进行添加
            if isinstance(value, (dict, list)):
                doc.add_paragraph(json.dumps(value, ensure_ascii=False, indent=2))
            else:
                doc.add_paragraph(str(value))

        # 加上时间戳以避免文件名冲突
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        final_filename = f"{request_data.filename}_{timestamp}.docx"
        temp_doc_path = session_dir / final_filename
        doc.save(temp_doc_path)
        logger.info(f"📄 临时文档已生成: {temp_doc_path}")

        # 2. 上传到MinIO并获取URL
        minio_url = upload_to_local_minio(temp_doc_path, final_filename)

        # 3. 返回包含URL的成功响应
        return GenerateDocResponse(
            filename=final_filename,
            url=minio_url
        )

    except Exception as e:
        logger.error(f"处理 /tools/generate_doc_from_json 时发生意外错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"生成或上传文档时发生意外错误: {e}")
    finally:
        # 4. 清理临时目录
        if session_dir.exists():
            shutil.rmtree(session_dir)
            logger.info(f"🗑️ 已清理临时目录: {session_dir}")


# --- 运行应用 ---
if __name__ == "__main__":
    import uvicorn
    
    print("--- 文档处理工具箱 API ---")
    print("启动 Uvicorn 服务器...")
    print("访问 http://127.0.0.1:8080/docs 查看 FastAPI 自动文档")
    
    uvicorn.run(app, host="0.0.0.0", port=8080)






