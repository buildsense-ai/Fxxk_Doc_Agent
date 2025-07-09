#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FastAPI 应用：文档处理工具箱
提供独立的、可作为 MCP 工具使用的文档处理接口。
"""

import os
import re
import logging
import subprocess
import shutil
from typing import List, Dict
from pathlib import Path
import uuid

import re
import subprocess
from mcp.server.fastmcp import FastMCP
mcp = FastMCP("doc vertorize")
from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel, Field
from docx import Document

# --- 日志配置 ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# --- FastAPI 应用设置 ---
# 为工具API使用专用的标题和描述
app = FastAPI(
    title="文档处理工具箱 API",
    description="提供一系列独立的文档处理工具，例如文件格式转换、模板分析等。可以被 MCP (Model-as-a-Service Consumer Platform) 或其他应用集成。",
    version="1.0.0"
)

# --- 临时文件目录 ---
TEMP_DIR = Path("temp_tool_files")
TEMP_DIR.mkdir(exist_ok=True)


# --- Pydantic 模型定义 (用于请求和响应) ---

class PlaceholderExtractionResponse(BaseModel):
    """占位符提取工具的响应模型"""
    original_filename: str = Field(..., description="上传的原始文件名。")
    placeholders_found: int = Field(..., description="提取到的占位符总数。")
    placeholders: List[str] = Field(..., description="从文档中提取出的占位符键名列表。")


class DocConversionResponse(BaseModel):
    """DOC转换工具的响应模型（用于成功时的JSON消息）"""
    message: str = Field("文件转换成功，请查收下载的文件。", description="操作结果消息。")
    original_filename: str = Field(..., description="上传的原始 .doc 文件名。")
    converted_filename: str = Field(..., description="转换后的 .docx 文件名。")


# --- 核心工具逻辑 (从 main.py 中抽离并优化) ---

def convert_doc_to_docx(doc_path: Path) -> Path:
    """
    使用 LibreOffice 将 .doc 文件安全地转换为 .docx 文件。

    Args:
        doc_path: 输入的 .doc 文件路径对象。

    Returns:
        转换后的 .docx 文件路径对象。

    Raises:
        FileNotFoundError: 如果输入文件不存在。
        RuntimeError: 如果 LibreOffice 未安装或转换失败。
    """
    logger.info(f"🔄 开始 .doc -> .docx 转换，输入: {doc_path}")
    if not doc_path.exists():
        raise FileNotFoundError(f".doc 文件不存在: {doc_path}")

    # 定义输出路径
    output_dir = doc_path.parent
    docx_path = output_dir / f"{doc_path.stem}.docx"

    # 检查 LibreOffice 是否可用
    libreoffice_paths = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
        'libreoffice',  # Linux/Windows (in PATH)
        'soffice',  # Alternative command name
    ]
    libreoffice_cmd = None
    for path in libreoffice_paths:
        try:
            result = subprocess.run([path, '--version'], capture_output=True, text=True, timeout=10)
            if result.returncode == 0:
                libreoffice_cmd = path
                logger.info(f"✅ 找到可用的 LibreOffice: {path}")
                break
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    if not libreoffice_cmd:
        raise RuntimeError("LibreOffice 未安装或无法在系统路径中找到。文件转换失败。")

    # 执行转换命令
    cmd = [
        libreoffice_cmd,
        '--headless',
        '--convert-to', 'docx',
        '--outdir', str(output_dir),
        str(doc_path)
    ]
    logger.info(f"🔧 执行转换命令: {' '.join(cmd)}")

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            logger.error(f"❌ LibreOffice 转换失败。返回码: {result.returncode}\n标准错误: {result.stderr}")
            raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")

        if not docx_path.exists():
            raise RuntimeError(f"转换命令已执行，但输出文件 {docx_path} 未找到。")

        logger.info(f"✅ 转换成功，输出文件: {docx_path}")
        return docx_path

    except subprocess.TimeoutExpired:
        logger.error("❌ LibreOffice 转换超时 (超过30秒)。")
        raise RuntimeError("LibreOffice 转换超时。")
    except Exception as e:
        logger.error(f"❌ 转换过程中发生未知错误: {e}")
        raise


def extract_placeholders_from_template(doc_path: Path) -> List[str]:
    """
    从 Word 文档模板中提取所有被识别的占位符。
    支持下划线、标签等多种形式。

    Args:
        doc_path: .docx 模板文件的路径对象。

    Returns:
        一个包含所有占位符键名的字符串列表。
    """
    logger.info(f"🛠️ 开始从模板中提取占位符: {doc_path}")
    doc = Document(doc_path)
    placeholders = set()
    blank_counter = 0

    def repl_func(match: re.Match) -> str:
        nonlocal blank_counter
        # 这个函数仅用于在分析时生成键名，在工具场景下不实际替换文本
        if match.group(1) is not None:  # "致...（...）" 模式
            if "（签字）" not in match.group(0):
                hint = match.group(2)
                placeholders.add(f"inline_{hint}")
        elif match.group(3) is not None:  # "标签:" 模式
            label = match.group(3).strip()
            placeholders.add(f"label_{label}")
        elif match.group(4) is not None:  # 通用下划线模式
            placeholders.add(f"blank_{blank_counter}")
            blank_counter += 1
        return ""  # 返回值不重要

    pattern = re.compile(
        r"致\s*(__{3,})\s*（([^）]+)）"  # G1: underscore, G2: hint
        r"|([^：\n（(]+?)：\s*$"  # G3: label, avoids '(...):'
        r"|((?:_{4,}[\s\xa0]*)+)"  # G4: general underscore blocks
    )

    # 遍历所有段落和表格单元格
    for para in doc.paragraphs:
        if para.text.strip():
            pattern.sub(repl_func, para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    pattern.sub(repl_func, cell.text)

    logger.info(f"✅ 提取完成，发现 {len(placeholders)} 个唯一占位符。")
    return sorted(list(placeholders))


# --- API 端点定义 ---

@app.post(
    "/tools/convert_doc_to_docx",
    summary="转换 .doc 文件为 .docx",
    description="上传一个旧版的 .doc 文件，此工具会使用 LibreOffice 将其转换为现代的 .docx 格式并返回。服务器需要安装 LibreOffice。",
    operation_id="convert_doc_to_docx",
    tags=["文件处理工具"]
)
async def tool_convert_doc_to_docx(
        file: UploadFile = File(..., description="要转换的 .doc 文件。")
):
    """
    接收 .doc 文件，转换为 .docx 并提供下载。
    """
    if not file.filename.lower().endswith('.doc'):
        raise HTTPException(status_code=400, detail="文件格式错误，请上传 .doc 文件。")

    # 安全地保存上传的文件
    session_id = str(uuid.uuid4())
    session_dir = TEMP_DIR / session_id
    session_dir.mkdir()

    original_path = session_dir / file.filename
    try:
        with open(original_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 调用核心转换逻辑
        converted_path = convert_doc_to_docx(original_path)

        # 提供转换后的文件下载
        return FileResponse(
            path=str(converted_path),
            filename=converted_path.name,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            # headers={"X-Success-Message": json.dumps(success_response.dict())} # 另一种传递元数据的方式
        )
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except RuntimeError as e:
        # 捕获 LibreOffice 相关的错误
        raise HTTPException(status_code=500, detail=f"服务器内部错误: {e}")
    except Exception as e:
        logger.error(f"处理 /tools/convert_doc_to_docx 时发生意外错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"处理文件时发生意外错误: {e}")
    finally:
        # 清理临时会话目录
        if session_dir.exists():
            shutil.rmtree(session_dir)


@app.post(
    "/tools/extract_placeholders",
    response_model=PlaceholderExtractionResponse,
    summary="提取 Word 模板中的占位符",
    description="上传一个 .docx 模板文件，此工具会分析其内容（包括段落和表格），并返回一个包含所有可识别占位符的列表。支持如下划线、'标签:' 等格式。",
    operation_id="extract_placeholders_from_template",
    tags=["模板分析工具"]
)
async def tool_extract_placeholders(
        file: UploadFile = File(..., description="要分析的 .docx 模板文件。")
):
    """
    接收 .docx 文件，分析并返回所有占位符。
    """
    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="文件格式错误，请上传 .docx 文件。")

    # 安全地保存上传的文件
    session_id = str(uuid.uuid4())
    session_dir = TEMP_DIR / session_id
    session_dir.mkdir()

    original_path = session_dir / file.filename
    try:
        with open(original_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # 调用核心提取逻辑
        placeholders = extract_placeholders_from_template(original_path)

        return PlaceholderExtractionResponse(
            original_filename=file.filename,
            placeholders_found=len(placeholders),
            placeholders=placeholders
        )
    except Exception as e:
        logger.error(f"处理 /tools/extract_placeholders 时发生意外错误: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析文件时发生意外错误: {e}")
    finally:
        # 清理临时会话目录
        if session_dir.exists():
            shutil.rmtree(session_dir)






