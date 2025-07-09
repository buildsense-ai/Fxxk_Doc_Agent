# main.py
# -*- coding: utf-8 -*-

import time
import random
import math
from typing import Optional # 用于可选参数
from fastapi import FastAPI, Query, HTTPException, UploadFile, File, Body
from pydantic import BaseModel, Field # 用于数据校验和响应模型
import uvicorn # 用于运行 ASGI 应用
import mysql.connector  # 修正数据库连接导入
import shutil
import requests
import tempfile
import os
import re
import json
import uuid
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches
from minio import Minio
from minio.error import S3Error
import logging

from mcptool.doc_vertorize import process_doc_file
from mcptool.micheal_doc import tool_extract_placeholders,tool_convert_doc_to_docx, convert_doc_to_docx, extract_placeholders_from_template

# 1. 导入 FastApiMCP
from fastapi_mcp import FastApiMCP

# 设置日志
logger = logging.getLogger(__name__)

# 2. 创建 FastAPI 应用实例
app = FastAPI(
    title="简单工具箱 API (MCP Enabled)",
    description="提供 BMI 计算、获取时间戳、加法、反转字符串和随机数等简单工具，并可通过 MCP 协议调用。",
    version="1.0.0",
)

# 导入必要的库
from datetime import date
from decimal import Decimal
from pydantic import BaseModel
from fastapi import Query

# --- 定义 Pydantic 模型 (用于规范化响应体) ---
class ElectricityUsageResponse(BaseModel):
    record_time: str = Field(..., description="用电记录时间（timestamp字符串）")
    account_balance: Decimal = Field(..., description="账户余额")
    electricity_kwh: Decimal = Field(..., description="用电量（千瓦时）")
    message: str = Field(..., description="操作结果信息")

class BMICalculationResponse(BaseModel):
    bmi: float = Field(..., description="计算出的 BMI 指数")
    category: str = Field(..., description="根据 BMI 值判断的健康分类")
    height_m: float = Field(..., description="输入的有效身高（米）")
    weight_kg: float = Field(..., description="输入的有效体重（公斤）")

class TimestampResponse(BaseModel):
    timestamp: float = Field(..., description="当前的 Unix 时间戳 (秒)")
    human_readable_utc: str = Field(..., description="当前时间的可读格式 (UTC)")

class AddResponse(BaseModel):
    x: float
    y: float
    result: float = Field(..., description="两数之和")

class ReverseStringResponse(BaseModel):
    original_text: str
    reversed_text: str = Field(..., description="反转后的字符串")

class RandomNumberResponse(BaseModel):
    min_value: int
    max_value: int
    random_number: int = Field(..., description="在指定范围内生成的随机整数")

class FillTemplateResponse(BaseModel):
    success: bool = Field(..., description="操作是否成功")
    template_url: str = Field(..., description="原始模板URL")
    output_filename: str = Field(..., description="输出文件名")
    minio_url: str = Field(..., description="MinIO文件访问地址")
    message: str = Field(..., description="处理结果消息")
    filled_fields_count: int = Field(..., description="填充的字段数量")

# --- 填充模板的核心功能函数 ---

def download_template_from_url(template_url: str) -> Path:
    """从URL下载模板文件"""
    try:
        response = requests.get(template_url, timeout=30)
        response.raise_for_status()
        
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(response.content)
        temp_file.close()
        
        return Path(temp_file.name)
    except Exception as e:
        raise Exception(f"下载模板失败: {str(e)}")

def fill_template_with_data(template_path: Path, data: dict) -> Path:
    """使用JSON数据填充模板，参考feishu/main.py的复杂填充逻辑"""
    try:
        doc = Document(template_path)
        filled_count = 0
        
        # 1. 准备附件信息
        attachments_map = data.pop('attachments_map', {})
        attachment_ref_map = {}
        ordered_attachments = []
        if attachments_map and isinstance(attachments_map, dict):
            logger.info(f"🖼️  找到 {len(attachments_map)} 个图片附件待处理。")
            ordered_attachments = list(attachments_map.items())
            for i, (key, _) in enumerate(ordered_attachments):
                attachment_ref_map[key.strip()] = i + 1
        else:
            attachments_map = {}

        # 2. 分离文本填充数据
        placeholder_data = {k: v for k, v in data.items() if k.startswith(('label_', 'inline_', 'blank_'))}
        structure_data = {k: v for k, v in data.items() if k.startswith(('table_', 'paragraph_'))}
        
        # 3. 替换所有文本占位符（包括图片引用）
        image_placeholder_pattern = re.compile(r'\{\{image:([^}]+)\}\}')
        text_placeholder_pattern = re.compile(r'\{(label_[^}]+|inline_[^}]+|blank_[^}]+)\}')

        def process_element_text(element):
            nonlocal filled_count
            if '{' not in element.text:
                return

            original_text = element.text
            new_text = ""
            last_end = 0

            # 创建一个组合的正则表达式来查找所有类型的占位符
            combined_pattern = re.compile(f"({image_placeholder_pattern.pattern}|{text_placeholder_pattern.pattern})")
            
            for match in combined_pattern.finditer(original_text):
                new_text += original_text[last_end:match.start()]
                last_end = match.end()
                
                image_key_match = image_placeholder_pattern.match(match.group(0))
                text_key_match = text_placeholder_pattern.match(match.group(0))

                if image_key_match:
                    key = image_key_match.group(1).strip()
                    if key in attachment_ref_map:
                        number = attachment_ref_map[key]
                        replacement = f"（详见附件{number}）"
                        new_text += replacement
                        logger.info(f"   🖼️  图片引用替换: '{match.group(0)}' -> '{replacement}'")
                    else:
                        logger.warning(f"   ⚠️  找到图片占位符 {match.group(0)} 但无匹配图片，已移除。")
                
                elif text_key_match:
                    placeholder_key = text_key_match.group(1)
                    placeholder = f"{{{placeholder_key}}}"

                    if placeholder_key in placeholder_data:
                        value = str(placeholder_data[placeholder_key])
                        new_text += value
                        logger.info(f"   ✏️  占位符填充: {placeholder} -> {value[:50]}...")
                        filled_count += 1
                    else: # 未匹配的文本占位符
                        if placeholder_key.startswith('label_'):
                            logger.info(f"   🔘  移除未匹配标签占位符: {placeholder}")
                            # The replacement is empty string, so we add nothing
                        elif placeholder_key.startswith(('inline_', 'blank_')):
                            original_underscore = "____"  # 默认下划线
                            new_text += original_underscore
                            logger.info(f"   🔘  恢复未匹配占位符: {placeholder} -> '{original_underscore}'")

            new_text += original_text[last_end:]
            
            if new_text != original_text:
                element.text = new_text

        # 遍历段落和表格进行统一替换
        for para in doc.paragraphs:
            process_element_text(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_element_text(cell)

        # 4. 填充原始结构
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    cell_key = f"table_{i}_row_{j}_col_{k}"
                    if cell_key in structure_data:
                        cell.text = str(structure_data[cell_key])
                        logger.info(f"   ✏️  结构填充(表格): {cell_key} -> {str(structure_data[cell_key])[:50]}...")
                        filled_count += 1

        for i, para in enumerate(doc.paragraphs):
            para_key = f"paragraph_{i}"
            if para_key in structure_data:
                # 只有在段落中没有占位符的情况下才进行结构填充
                if not re.search(r'\{.*\}', para.text):
                    para.text = str(structure_data[para_key])
                    logger.info(f"   ✏️  结构填充(段落): {para_key} -> {str(structure_data[para_key])[:50]}...")
                    filled_count += 1

        # 5. 将图片作为附件附加到文档末尾
        if ordered_attachments:
            logger.info("📎 开始在文档末尾附加图片...")
            try:
                doc.add_page_break()
                doc.add_heading('附件列表', level=1)
                
                for i, (key, image_path) in enumerate(ordered_attachments):
                    attachment_counter = i + 1
                    if not image_path or not isinstance(image_path, str) or not os.path.exists(image_path):
                        logger.warning(f"⚠️ 图片路径不存在或无效，跳过附件 '{key}': {image_path}")
                        continue
                    
                    try:
                        heading_text = f"附件 {attachment_counter}: {key}"
                        doc.add_heading(heading_text, level=2)
                        doc.add_picture(image_path, width=Inches(6.0))
                        doc.add_paragraph()
                        logger.info(f"   ✅ 成功附加图片: {heading_text} ({image_path})")
                    except Exception as pic_e:
                        logger.error(f"❌ 附加图片 '{key}' ({image_path}) 时出错: {pic_e}")
            except Exception as e:
                logger.error(f"❌ 处理附件时发生意外错误: {e}")
        
        # 保存填充后的文档
        output_path = template_path.parent / f"filled_{template_path.name}"
        doc.save(output_path)
        logger.info(f"✅ 混合模式填充完成，共填充 {filled_count} 个字段: {output_path}")
        
        return output_path
        
    except Exception as e:
        raise Exception(f"填充模板失败: {str(e)}")

def upload_to_local_minio(file_path: Path, filename: str) -> str:
    """上传文件到MinIO服务器"""
    try:
        # MinIO配置
        minio_client = Minio(
            "43.139.19.144:9000",
            access_key="minioadmin",
            secret_key="minioadmin",
            secure=False
        )
        
        bucket_name = "templates"
        
        # 确保存储桶存在
        if not minio_client.bucket_exists(bucket_name):
            minio_client.make_bucket(bucket_name)
        
        # 上传文件
        minio_client.fput_object(bucket_name, filename, str(file_path))
        
        # 返回访问URL
        return f"http://43.139.19.144:9000/{bucket_name}/{filename}"
        
    except Exception as e:
        raise Exception(f"上传到MinIO失败: {str(e)}")

def fill_template_from_url_impl(template_url: str, data: dict, output_filename: str = None) -> dict:
    """
    从URL获取模板，填充JSON数据，上传到MinIO并返回地址。
    """
    try:
        # 1. 下载模板
        template_path = download_template_from_url(template_url)
        
        # 2. 填充数据（使用复杂的填充逻辑）
        filled_doc_path = fill_template_with_data(template_path, data.copy())
        
        # 3. 生成输出文件名
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            output_filename = f"filled_template_{timestamp}.docx"
        elif not output_filename.endswith('.docx'):
            output_filename += '.docx'
        
        # 4. 上传到MinIO
        minio_url = upload_to_local_minio(filled_doc_path, output_filename)
        
        # 5. 清理临时文件
        template_path.unlink(missing_ok=True)
        filled_doc_path.unlink(missing_ok=True)
        
        return {
            "success": True,
            "template_url": template_url,
            "output_filename": output_filename,
            "minio_url": minio_url,
            "message": "模板填充并上传成功",
            "filled_fields_count": len(data) - (1 if 'attachments_map' in data else 0)
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "template_url": template_url
        }

def template_convert_doc_to_docx(doc_path: str) -> dict:
    """
    使用 LibreOffice 将 .doc 文件转换为 .docx 文件。
    Args:
        doc_path: 输入的 .doc 文件路径。
    Returns:
        包含转换结果的字典。
    """
    doc_path = Path(doc_path)
    if not doc_path.exists():
        return {"error": f"文件不存在: {doc_path}"}

    try:
        # 定义输出路径
        output_dir = doc_path.parent
        docx_path = output_dir / f"{doc_path.stem}.docx"

        # 检查 LibreOffice
        libreoffice_cmd = None
        for cmd in ['libreoffice', 'soffice']:
            try:
                import subprocess
                result = subprocess.run([cmd, '--version'], capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    libreoffice_cmd = cmd
                    break
            except:
                continue

        if not libreoffice_cmd:
            return {"error": "LibreOffice未安装或无法访问"}

        # 执行转换
        cmd = [
            libreoffice_cmd,
            '--headless',
            '--convert-to', 'docx',
            '--outdir', str(output_dir),
            str(doc_path)
        ]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            return {"error": f"转换失败: {result.stderr}"}

        if not docx_path.exists():
            return {"error": "转换命令执行成功但未找到输出文件"}

        return {
            "success": True,
            "original_filename": doc_path.name,
            "converted_filename": docx_path.name,
            "converted_path": str(docx_path)
        }

    except Exception as e:
        return {"error": f"转换过程中发生错误: {str(e)}"}

def generate_doc_from_json(filename: str, data: dict) -> dict:
    """
    根据JSON数据生成Word文档。
    Args:
        filename: 希望生成的文档名（无需后缀）
        data: 用于填充文档的数据字典
    Returns:
        包含生成结果的字典
    """
    try:
        # 创建临时目录
        temp_dir = Path("temp_docx_files")
        temp_dir.mkdir(exist_ok=True)
        session_id = str(uuid.uuid4())
        session_dir = temp_dir / session_id
        session_dir.mkdir()

        # 生成文档
        doc = Document()
        doc.add_heading(f"数据报告: {filename}", level=0)
        doc.add_paragraph(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        for key, value in data.items():
            doc.add_heading(str(key), level=2)
            if isinstance(value, (dict, list)):
                doc.add_paragraph(json.dumps(value, ensure_ascii=False, indent=2))
            else:
                doc.add_paragraph(str(value))

        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        final_filename = f"{filename}_{timestamp}.docx"
        temp_doc_path = session_dir / final_filename
        doc.save(temp_doc_path)

        # 上传到MinIO
        minio_url = upload_to_local_minio(temp_doc_path, final_filename)

        result = {
            "success": True,
            "filename": final_filename,
            "url": minio_url,
            "local_path": str(temp_doc_path)
        }

        return result

    except Exception as e:
        return {"error": f"生成文档时发生错误: {str(e)}"}
    finally:
        # 清理临时目录
        if 'session_dir' in locals() and session_dir.exists():
            shutil.rmtree(session_dir)

# --- API 端点定义 ---

@app.post(
    "/tools/insert_electricity_usage",
    response_model=ElectricityUsageResponse,
    summary="插入电费使用记录",
    description="向数据库插入一条电费使用记录，包含使用日期、账户余额和用电量。",
    operation_id="insert_electricity_usage"
)
async def insert_electricity_usage(
    account_balance: float = Query(..., gt=0, description="账户余额（元）", examples={"example": {"value": 500.50}}),
    electricity_kwh: float = Query(..., gt=0, description="用电量（千瓦时）", examples={"example": {"value": 120.5}})
):
    """插入一条电费使用记录到数据库。record_time自动取当前时间。"""
    import datetime
    record_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        # 建立数据库连接
        conn = mysql.connector.connect(
            host="localhost",
            port=3306,
            user="root",
            password="Aa@831015",
            database="elec_bills"
        )
        cursor = conn.cursor()

        # 执行插入操作
        insert_query = "INSERT INTO electricity_consumption (record_time, account_balance, electricity_kwh) VALUES (%s, %s, %s)"
        cursor.execute(insert_query, (record_time, account_balance, electricity_kwh))
        conn.commit()

        # 关闭连接
        cursor.close()
        conn.close()

        return ElectricityUsageResponse(
            record_time=record_time,
            account_balance=Decimal(str(account_balance)),
            electricity_kwh=Decimal(str(electricity_kwh)),
            message="数据插入成功"
        )

    except mysql.connector.Error as e:
        return ElectricityUsageResponse(
            record_time=record_time,
            account_balance=Decimal(str(account_balance)),
            electricity_kwh=Decimal(str(electricity_kwh)),
            message=f"数据库操作失败: {str(e)}"
        )

# --- 定义 API 端点 (这些将作为 MCP 工具暴露) ---

@app.get(
    "/tools/calculate_bmi", # 建议给工具类接口加上 /tools 前缀
    response_model=BMICalculationResponse,
    summary="计算 BMI 指数", # 简洁的概括，用于 OpenAPI 和 MCP
    description="根据输入的身高（单位：米）和体重（单位：公斤）计算身体质量指数 (BMI) 并给出健康分类。", # 详细描述
    operation_id="calculate_bmi"
)
async def calculate_bmi(
    height: float = Query(..., gt=0, description="身高，单位为米 (m)，必须大于0。", example=1.75),
    weight: float = Query(..., gt=0, description="体重，单位为公斤 (kg)，必须大于0。", example=70.0)
):
    """
    计算身体质量指数 (BMI)。

    需要提供身高（米）和体重（公斤）作为参数。
    """
    # FastAPI 的 Query(gt=0) 已经做了基础校验，这里可以省略显式检查
    bmi = round(weight / (height ** 2), 2) # 保留两位小数
    category = "未知分类"
    # 可以根据需要添加更详细的分类逻辑
    if bmi < 18.5:
        category = "体重过轻"
    elif 18.5 <= bmi < 24:
        category = "体重正常"
    elif 24 <= bmi < 28:
        category = "超重"
    else: # bmi >= 28
        category = "肥胖"

    return BMICalculationResponse(
        bmi=bmi,
        category=category,
        height_m=height,
        weight_kg=weight
    )

@app.get(
    "/tools/current_timestamp",
    response_model=TimestampResponse,
    summary="获取当前 Unix 时间戳",
    description="返回服务器当前的 Unix 时间戳（自 UTC 1970年1月1日午夜以来的秒数）和对应的 UTC 时间字符串。",
    operation_id="get_current_timestamp"
)
async def get_current_timestamp():
    """获取当前时间的 Unix 时间戳。"""
    current_timestamp = time.time()
    # 转换为整数时间戳（如果需要）
    # current_timestamp_int = int(current_timestamp)
    readable_utc = time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime(current_timestamp))
    return TimestampResponse(timestamp=current_timestamp, human_readable_utc=readable_utc)

@app.get(
    "/tools/add_numbers",
    response_model=AddResponse,
    summary="计算两数之和",
    description="输入两个数字 x 和 y，返回它们的和。",
    operation_id="add_numbers"
)
async def add_numbers(
    x: float = Query(..., description="第一个数字", example=5.5),
    y: float = Query(..., description="第二个数字", example=4.5)
):
    """计算两个浮点数的和。"""
    result = x + y
    return AddResponse(x=x, y=y, result=result)

@app.get(
    "/tools/reverse_string",
    response_model=ReverseStringResponse,
    summary="反转字符串（有趣小函数）",
    description="输入一个字符串，将其反转后返回。",
    operation_id="reverse_string"
)
async def reverse_string_tool(
    text: str = Query(..., description="需要反转的文本", example="你好，世界")
):
    """反转给定的输入字符串。"""
    reversed_text = text[::-1]
    return ReverseStringResponse(original_text=text, reversed_text=reversed_text)

@app.get(
    "/tools/random_number",
    response_model=RandomNumberResponse,
    summary="获取指定范围内的随机数",
    description="在指定的最小值和最大值区间内生成一个随机整数。",
    operation_id="get_random_number"
)
async def get_random_number(
    min_value: int = Query(..., description="随机数的最小值", example=1),
    max_value: int = Query(..., description="随机数的最大值", example=100)
):
    """在指定的 [min_value, max_value] 区间内生成一个随机整数。"""
    if min_value > max_value:
        raise HTTPException(status_code=400, detail="最小值 (min_value) 不能大于最大值 (max_value)。")
    random_integer = random.randint(min_value, max_value)
    return RandomNumberResponse(min_value=min_value, max_value=max_value, random_number=random_integer)

@app.post(
    "/tools/upload_and_vectorize",
    summary="上传文档并进行智能结构化处理",
    description="上传一个doc/docx文档，自动提取结构化信息和向量描述。",
    operation_id="upload_and_vectorize",
)
async def upload_and_vectorize(file: UploadFile = File(...)):
    temp_dir = "temp_docx_files"
    import os
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    result = process_doc_file(temp_path)
    return result

@app.post(
    "/tools/extract_placeholders",
    summary="提取文档中的占位符",
    description="上传一个docx文档，自动提取所有占位符。",
    operation_id="extract_placeholders",
)
async def extract_placeholders(file: UploadFile = File(...)):
    temp_dir = "temp_docx_files"
    import os
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    placeholders = extract_placeholders_from_template(Path(temp_path))
    return {
        "original_filename": file.filename,
        "placeholders_found": len(placeholders),
        "placeholders": placeholders
    }

@app.post(
    "/tools/convert_doc_to_docx",
    summary="将doc文档转换为docx格式",
    description="上传一个doc文档，自动转换为docx格式并返回新文件名。",
    operation_id="convert_doc_to_docx",
)
async def convert_doc_to_docx_api(file: UploadFile = File(...)):
    temp_dir = "temp_docx_files"
    import os
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    docx_path = convert_doc_to_docx(Path(temp_path))
    return {
        "original_filename": file.filename,
        "converted_filename": os.path.basename(docx_path)
    }

@app.post(
    "/tools/template_convert_doc",
    summary="模板工具：转换doc为docx",
    description="使用模板工具将doc文件转换为docx格式。",
    operation_id="template_convert_doc"
)
async def template_convert_doc(file: UploadFile = File(...)):
    temp_dir = "temp_docx_files"
    import os
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, file.filename)
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    result = template_convert_doc_to_docx(temp_path)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result

@app.post(
    "/tools/generate_doc",
    summary="根据JSON数据生成Word文档",
    description="接收文档名和JSON数据，生成对应的Word文档。",
    operation_id="generate_doc"
)
async def generate_doc(filename: str = Query(..., description="文档名（无需后缀）"), 
                      data: dict = Body(..., description="用于生成文档的数据")):
    result = generate_doc_from_json(filename, data)
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result

@app.post(
    "/tools/fill_template",
    response_model=FillTemplateResponse,
    summary="从URL获取模板并填充数据",
    description="从指定URL下载模板，使用JSON数据填充，上传到本地MinIO并返回地址。支持复杂的占位符填充逻辑，包括图片处理。",
    operation_id="fill_template"
)
async def fill_template(
    template_url: str = Query(..., description="模板文件的URL地址"),
    data: dict = Body(..., description="用于填充模板的JSON数据"),
    output_filename: str = Query(None, description="输出文件名（可选）")
):
    """
    从URL获取模板并填充数据。
    
    支持多种占位符格式：
    - 图片占位符：{{image:key}}
    - 文本占位符：{label_key}、{inline_key}、{blank_key}
    - 结构数据：table_*、paragraph_*
    """
    result = fill_template_from_url_impl(template_url, data, output_filename)
    if not result.get("success", False):
        raise HTTPException(status_code=400, detail=result.get("error", "未知错误"))
    
    return FillTemplateResponse(
        success=result["success"],
        template_url=result["template_url"],
        output_filename=result["output_filename"],
        minio_url=result["minio_url"],
        message=result["message"],
        filled_fields_count=result["filled_fields_count"]
    )

# --- 3. 集成 fastapi_mcp ---
# 创建 FastApiMCP 实例，把 FastAPI 的 app 对象传给它
mcp = FastApiMCP(app)
# 将 MCP 服务器挂载到 FastAPI 应用

mcp.mount()

# --- 4. Uvicorn 启动入口 (可选，方便直接运行脚本) ---
if __name__ == "__main__":
    print("--- 简单工具箱 API (MCP Enabled) ---")
    print("启动 Uvicorn 服务器...")
    print("访问 http://127.0.0.1:8000/docs 查看 FastAPI 自动文档")
    # 监听 0.0.0.0 使服务可以被外部访问 (如果需要)
    # reload=True 会在代码变动时自动重启服务器，方便开发
    uvicorn.run("main:app", host="0.0.0.0", port=8000)