# 文件名: generator.py
# -*- coding: utf-8 -*-

"""
generator.py

包含项目的核心业务逻辑：状态管理和文档生成状态机。
"""

import json
import os
import time
import uuid
from typing import Dict, Any, Optional

import docx

# --- 从其他模块导入依赖 ---
from config import Config
from services import call_ai_model, search_vectordata
# [新增] 导入新的上传模块
from upload_cloud import upload_to_minio


# TaskState 类的代码保持不变
class TaskState:
    """管理单个生成任务的状态，负责JSON文件的读写和更新。"""

    def __init__(self, task_id: str):
        self.task_id = task_id
        self.filepath = os.path.join(Config.TASKS_DIR, f"task_{self.task_id}.json")
        self.data: Dict[str, Any] = {}
        os.makedirs(Config.TASKS_DIR, exist_ok=True)

    def load(self) -> bool:
        if os.path.exists(self.filepath):
            with open(self.filepath, 'r', encoding='utf-8') as f:
                self.data = json.load(f)
            return True
        return False

    def save(self):
        self.data['lastUpdatedTimestamp'] = time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())
        with open(self.filepath, 'w', encoding='utf-8') as f:
            json.dump(self.data, f, ensure_ascii=False, indent=2)
        print(f"--- 状态已保存: {self.filepath} ---")

    def initialize(self, initial_request: Dict[str, str]):
        # [变更] 新增 docxPublicUrl 字段
        self.data = {
            "taskId": self.task_id, "status": "pending", "progressPercentage": 0,
            "currentStatusMessage": "任务已创建，等待初始化...", "initialRequest": initial_request,
            "creativeBrief": "", "projectName": "", "outline": {}, "finalDocument": "",
            "docxPublicUrl": "", "errorLog": []
        }
        self.save()

    def update_status(self, status: str, message: str, progress: int):
        self.data['status'] = status
        self.data['currentStatusMessage'] = message
        self.data['progressPercentage'] = progress
        self.save()
        print(f"--- 进度更新: {progress}% - {message} ---")

    def log_error(self, stage: str, error_message: str):
        self.data['errorLog'].append({
            "timestamp": time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime()),
            "stage": stage, "message": error_message
        })
        self.update_status("failed", f"在 {stage} 阶段发生错误。", self.data.get('progressPercentage', 0))


class LongDocumentGenerator:
    """负责执行整个长文生成任务的业务流程。"""

    def __init__(self, task_id: Optional[str] = None):
        self.task_id = task_id or str(uuid.uuid4())
        self.state = TaskState(self.task_id)

    def start_new_job(self, chathistory: str, request: str) -> str:
        print(f"启动新任务: {self.task_id}")
        self.state.initialize(initial_request={"chathistory": chathistory, "request": request})
        self.run()
        return self.task_id

    def run(self):
        """执行任务的主循环。"""
        try:
            if not self.state.load():
                print(f"错误：无法加载任务 {self.task_id}。")
                return

            while self.state.data['status'] not in ['completed', 'failed']:
                current_status = self.state.data['status']
                if current_status == 'pending':
                    self._prepare_creative_brief()
                elif current_status == 'brief_prepared':
                    self._generate_initial_outline()
                elif current_status == 'outline_generated':
                    self._refine_outline_cycle()
                elif current_status == 'outline_finalized':
                    self._generate_all_chapters()
                elif current_status == 'chapters_generated':
                    self._assemble_final_document()
                else:
                    print(f"!! 发现未知任务状态: {current_status}，任务中止。")
                    self.state.log_error("run_loop", f"未知状态: {current_status}")
                    break

            if self.state.data['status'] == 'completed':
                print(f"\n任务 {self.task_id} 已成功完成！")

        except Exception as e:
            error_stage = self.state.data.get('status', 'unknown')
            print(f"!! 在 {error_stage} 阶段发生严重错误: {e}")
            self.state.log_error(error_stage, str(e))
            self.state.update_status("failed", "任务因意外错误而终止。", self.state.data.get('progressPercentage', 0))

    def _prepare_creative_brief(self):
        """[已优化] 阶段1: 准备创作指令, 并提取核心项目主题"""
        self.state.update_status("brief_generation", "正在分析聊天记录和用户请求...", 5)
        chathistory = self.state.data['initialRequest']['chathistory']
        request = self.state.data['initialRequest']['request']

        prompt_brief = f"""你是一个任务规划助理。请根据下面的聊天记录和用户的最终请求，提炼并生成一份详尽的、用于指导后续长文写作的“创作指令”（Creative Brief）。
这份指令应该清晰、结构化，并综合所有已知信息。
【聊天记录】
{chathistory}
【用户最终请求】
{request}
请以JSON格式返回你的分析结果，包含一个'creative_brief'字段，其内容是一段详细的文本指令。
重要提示：所有生成的文本内容都必须使用中文。
"""
        response_brief = call_ai_model(prompt_brief, expect_json=True)
        brief = response_brief.get("creative_brief")
        if not brief:
            raise ValueError("AI未能生成有效的创作指令（creative_brief）。")
        self.state.data['creativeBrief'] = brief

        # --- [新增] 提取核心项目名/主题，用于优化后续的向量检索 ---
        self.state.update_status("brief_generation", "正在提炼项目主题以优化检索...", 7)
        prompt_project_name = f"""从以下创作指令中，提取一个简短的核心项目名称或主题（例如，“长沙理工大学羽毛球比赛”或“1号住宅楼结构设计”），用于优化后续的知识库检索。
请以JSON格式返回，只包含一个 'project_name' 字段。
重要提示：项目名称必须使用中文。
创作指令：{brief}
"""
        response_name = call_ai_model(prompt_project_name, expect_json=True)
        project_name = response_name.get("project_name", "")
        self.state.data['projectName'] = project_name
        print(f"--- 提炼出项目主题: {project_name} ---")
        # --- [新增结束] ---

        self.state.data['status'] = 'brief_prepared'
        self.state.save()

    def _generate_initial_outline(self):
        """阶段2: 生成初始大纲"""
        self.state.update_status("outline_generation", "正在生成初始大纲...", 10)
        prompt = f"""根据以下创作指令，为一篇长文档，生成一份结构化的JSON格式大纲。
JSON的根节点应该有一个名为 'chapters' 的列表。
列表中的每个对象都应包含以下三个字段：
1. 'chapterId' (字符串, 例如 "ch_01")
2. 'title' (字符串, 章节的标题)
3. 'key_points' (字符串列表, 这一章的关键论点)
重要提示：所有生成的标题和文本内容都必须使用中文。
指令：{self.state.data['creativeBrief']}
"""
        response = call_ai_model(prompt, expect_json=True)
        chapters = response.get('chapters')
        if chapters is None:
            raise ValueError("AI未能生成有效的大纲（缺少'chapters'字段）。")
        self.state.data['outline'] = {"metadata": {"refinementCycles": 0}, "chapters": chapters}
        self.state.data['status'] = 'outline_generated'
        self.state.save()

    def _refine_outline_cycle(self):
        """阶段3: 大纲评审与精炼循环"""
        self.state.update_status("outline_refinement", "开始大纲评审与精炼...", 15)
        project_name = self.state.data.get('projectName', '')  # 获取项目名称

        for i in range(Config.MAX_REFINEMENT_CYCLES):
            self.state.update_status("outline_refinement", f"第 {i + 1} 轮大纲精炼：AI自我评审...", 15 + i * 5)
            current_outline_str = json.dumps(self.state.data['outline']['chapters'], ensure_ascii=False)
            prompt_critique = f"""你是一位追求完美的资深编辑。请评审以下大纲。
即使它看起来逻辑完整，也请你主动思考：哪些章节可以通过补充外部的、更具体的数据、案例或细节来变得更具深度和说服力？
请以JSON格式返回你的分析。JSON的根节点应包含一个名为 'gaps_identified' 的列表。
列表中的每个对象都应代表一个需要补充信息的章节，并包含以下字段：
1. 'chapterId' (字符串, 对应原大纲中的章节ID)
2. 'title' (字符串, 对应原大纲中的章节标题)
3. 'query_keywords' (字符串列表, 用于检索外部知识的推荐关键词)
如果这个大纲确实完美无缺，才返回一个空的 'gaps_identified' 列表。
重要提示：所有生成的文本内容都必须使用中文。
大纲：{current_outline_str}
"""
            critique_response = call_ai_model(prompt_critique, expect_json=True)
            gaps = critique_response.get('gaps_identified')
            if not gaps:
                print("评审完成，AI认为大纲已足够完善。")
                break

            self.state.update_status("outline_refinement", f"发现信息缺口: {gaps}，正在检索知识...", 16 + i * 5)
            all_knowledge = []
            all_gap_titles = []
            for gap_info in gaps:
                gap_title = gap_info.get('title', '')
                if gap_title:
                    all_gap_titles.append(gap_title)
                query_keywords = gap_info.get('query_keywords', [])
                for keyword in query_keywords:
                    scoped_query = f"{project_name} {keyword}".strip()
                    knowledge_pieces = search_vectordata(query=scoped_query, top_k=Config.SEARCH_API_TOP_K)
                    all_knowledge.extend(knowledge_pieces)

            if not all_knowledge:
                print("未能从向量数据库检索到有效信息，跳过本轮增强。")
                continue

            knowledge_str = "\n\n---\n\n".join(all_knowledge)
            prompt_integrate = f"""请参考以下背景资料，扩充和完善这份大纲，特别是关于'{','.join(all_gap_titles)}'的这些章节。
请返回完整的、更新后的JSON格式大纲。
返回的JSON结构应与原大纲一致（一个包含 'chapters' 列表的根对象）。
重要提示：所有生成的文本内容都必须使用中文。
背景资料：{knowledge_str}
原大纲：{current_outline_str}
"""
            updated_outline_response = call_ai_model(prompt_integrate, expect_json=True)
            self.state.data['outline']['chapters'] = updated_outline_response.get('chapters',
                                                                                  self.state.data['outline'][
                                                                                      'chapters'])
            self.state.data['outline']['metadata']['refinementCycles'] = i + 1
            self.state.save()

        self.state.data['status'] = 'outline_finalized'
        self.state.update_status("outline_finalized", "大纲已最终确定，准备生成正文。", 30)

    def _generate_all_chapters(self):
        """阶段4: 逐章生成内容"""
        chapters = self.state.data['outline']['chapters']
        project_name = self.state.data.get('projectName', '')
        total_chapters = len(chapters)
        if total_chapters == 0:
            print("!! 警告：大纲为空，无法生成任何章节内容。")
            self.state.data['status'] = 'chapters_generated'
            self.state.save()
            return

        for i, chapter in enumerate(chapters):
            progress = 30 + int((i / total_chapters) * 60)
            chapter_title = chapter.get('title', '')
            self.state.update_status("content_generation", f"正在生成第 {i + 1}/{total_chapters} 章: '{chapter_title}'...",
                                     progress)

            print(f"--> 为章节 '{chapter_title}' 检索背景知识...")
            scoped_query = f"{project_name} {chapter_title}".strip()
            knowledge_pieces = search_vectordata(query=scoped_query, top_k=Config.SEARCH_API_TOP_K)

            context = f"这是全文大纲：{json.dumps(self.state.data['outline']['chapters'], ensure_ascii=False)}\n"
            if i > 0:
                context += f"前一章是关于 '{chapters[i - 1].get('title', '')}' 的。\n"
            if knowledge_pieces:
                knowledge_str = "\n\n---\n\n".join(knowledge_pieces)
                context += f"\n在撰写本章时，请重点参考以下背景资料以确保内容的准确性和深度：\n{knowledge_str}\n"

            prompt = f"请根据以下信息，撰写 '{chapter_title}' 这一章的详细内容。请专注于阐述这些要点：{', '.join(chapter.get('key_points', []))}。重要提示：所有你撰写的内容都必须是中文。"
            response = call_ai_model(prompt, context=context)
            chapter['content'] = response.get('text', '')
            self.state.save()

        self.state.data['status'] = 'chapters_generated'
        self.state.save()

    def _assemble_final_document(self):
        """[已更新] 阶段5: 整合与终审，并上传docx"""
        self.state.update_status("assembling", "所有章节已生成，正在整合文档...", 95)
        chapters_for_context = self.state.data['outline']['chapters']
        if not chapters_for_context:
            print("!! 警告：大纲为空，将生成通用的引言和结论。")

        intro_prompt = "根据以下完整大纲，为这篇文档撰写一段精彩的引言。重要提示：引言必须使用中文。"
        intro_context = json.dumps(chapters_for_context, ensure_ascii=False)
        intro_response = call_ai_model(intro_prompt, intro_context)

        conclusion_prompt = "根据以下完整大纲和所有章节内容，为这篇文档撰写一段总结性的结论。重要提示：结论必须使用中文。"
        conclusion_response = call_ai_model(conclusion_prompt, intro_context)

        full_text_parts = [intro_response.get('text', '')]
        for chapter in chapters_for_context:
            full_text_parts.append(f"\n\n## {chapter.get('title', '')}\n\n")
            full_text_parts.append(chapter.get('content', ''))
        full_text_parts.append(f"\n\n## 结论\n\n")
        full_text_parts.append(conclusion_response.get('text', ''))

        self.state.data['finalDocument'] = "".join(full_text_parts)

        # 调用方法创建并上传 .docx 文件
        try:
            self._create_and_upload_docx()
        except Exception as e:
            print(f"!! [警告] 创建或上传 .docx 文件时发生错误: {e}")
            self.state.log_error("docx_handling", str(e))

        self.state.update_status("completed", "文档已成功生成！", 100)

    def _create_and_upload_docx(self):
        """
        [已重命名并更新] 根据 finalDocument 内容生成 .docx 文件并上传。
        """
        print("--- 正在生成 .docx 文档... ---")
        final_text = self.state.data.get('finalDocument', '')
        if not final_text:
            print("!! [警告] finalDocument 为空，无法生成 .docx 文件。")
            return

        doc = docx.Document()
        
        # 设置文档的默认字体为宋体
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # 设置文档默认样式为宋体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)  # 设置字体大小为12磅
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 确保中文字符使用宋体
        
        # 添加主标题并设置为宋体
        title_heading = doc.add_heading(self.state.data.get('projectName', '生成文档'), level=0)
        # 设置标题字体为宋体
        for run in title_heading.runs:
            run.font.name = '宋体'
            run.font.size = Pt(18)  # 主标题使用18磅
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        def set_font_style(element, font_size=12):
            """设置元素的字体为宋体的辅助函数"""
            try:
                if hasattr(element, 'runs'):
                    # 对于段落或标题，遍历所有runs
                    for run in element.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(font_size)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    # 如果没有runs，创建一个
                    if not element.runs:
                        run = element.add_run()
                        run.font.name = '宋体'
                        run.font.size = Pt(font_size)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                elif hasattr(element, 'font'):
                    # 对于单个run
                    element.font.name = '宋体'
                    element.font.size = Pt(font_size)
                    element._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except Exception as e:
                print(f"设置字体时出现错误: {e}")

        for line in final_text.split('\n'):
            stripped_line = line.strip()
            if stripped_line.startswith('## '):
                heading_text = stripped_line.replace('## ', '', 1)
                heading = doc.add_heading(heading_text, level=2)
                set_font_style(heading, 14)  # 二级标题使用14磅
            elif stripped_line:  # 只有非空行才添加段落
                paragraph = doc.add_paragraph(stripped_line)
                set_font_style(paragraph, 12)  # 正文使用12磅

        docx_filename = f"task_{self.state.task_id}.docx"
        docx_filepath = os.path.join(Config.TASKS_DIR, docx_filename)
        doc.save(docx_filepath)
        print(f"--- .docx 文档已成功保存至本地: {docx_filepath} ---")

        # --- [新增] 上传到MinIO ---
        public_url = upload_to_minio(
            file_path=docx_filepath,
            object_name=docx_filename
        )
        if public_url:
            self.state.data['docxPublicUrl'] = public_url
            self.state.save()  # 保存包含URL的最终状态
        # --- [新增结束] ---
