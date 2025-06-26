#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板转换工具：DOC/DOCX → JSON结构化模板
功能：
1. DOC文件转换为DOCX（使用LibreOffice）
2. DOCX模板分析和占位符提取
3. 输出JSON格式的模板结构
4. 为ReAct Agent提供工具接口
"""

import os
import json
import logging
import subprocess
import re
import time
import psutil
from typing import Dict, Any, List, Optional
from docx import Document

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

class TemplateConverter:
    """模板转换器 - 将DOC/DOCX转换为JSON结构"""
    
    def __init__(self):
        """初始化转换器"""
        self.placeholder_originals = {}  # Store original text of placeholders
        logger.info("🔧 模板转换器初始化完成")
    
    def kill_all_libreoffice_processes(self) -> None:
        """杀死所有LibreOffice进程，防止冲突和弹窗"""
        logger.info("🔄 检查并关闭现有LibreOffice进程...")
        
        process_names = ['soffice.exe', 'soffice.bin', 'soffice', 'libreoffice']
        killed_count = 0
        
        try:
            for proc in psutil.process_iter(['pid', 'name']):
                try:
                    if proc.info['name'] and any(name in proc.info['name'].lower() for name in process_names):
                        proc.kill()
                        killed_count += 1
                        logger.info(f"🔪 关闭进程: {proc.info['name']} (PID: {proc.info['pid']})")
                except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
                    pass
            
            if killed_count > 0:
                logger.info(f"✅ 已关闭 {killed_count} 个LibreOffice进程")
                time.sleep(2)  # 等待进程完全关闭
            else:
                logger.info("✅ 没有发现运行中的LibreOffice进程")
        except Exception as e:
            logger.warning(f"⚠️ 清理LibreOffice进程时出现问题: {e}")

    def set_no_popup_environment(self) -> dict:
        """设置防弹窗的环境变量"""
        env = os.environ.copy()
        
        # 设置环境变量禁用GUI
        env.update({
            'DISPLAY': '',  # Linux/Unix禁用显示
            'SAL_USE_VCLPLUGIN': 'svp',  # 使用无头插件
            'SAL_NO_MOUSEGRABS': '1',  # 禁用鼠标捕获
            'SAL_DISABLE_OPENCL': '1',  # 禁用OpenCL
            'SAL_DISABLE_CRASHDUMP': '1',  # 禁用崩溃转储
            'OOO_DISABLE_RECOVERY': '1',  # 禁用恢复功能
        })
        
        return env

    def convert_doc_to_docx(self, doc_path: str) -> str:
        """
        使用LibreOffice将.doc文件转换为.docx文件 - 完全无弹窗版本
        
        Args:
            doc_path: .doc文件路径
            
        Returns:
            转换后的.docx文件路径
        """
        logger.info("🔄 开始DOC到DOCX转换（无弹窗模式）...")
        
        if not os.path.exists(doc_path):
            logger.error(f"❌ DOC文件不存在: {doc_path}")
            raise FileNotFoundError(f"DOC文件不存在: {doc_path}")
        
        # 检查是否已经是DOCX文件
        if doc_path.lower().endswith('.docx'):
            logger.info("📄 文件已经是DOCX格式，无需转换")
            return doc_path
        
        # 生成输出文件名
        docx_path = doc_path.replace('.doc', '_converted.docx')
        
        try:
            # 步骤1: 强制关闭所有LibreOffice进程
            self.kill_all_libreoffice_processes()
            
            # 步骤2: 查找LibreOffice
            logger.info("🔍 检查LibreOffice可用性...")
            
            # 简化的LibreOffice检测
            libreoffice_cmd = r'C:\Program Files\LibreOffice\program\soffice.exe'
            
            if os.path.exists(libreoffice_cmd):
                logger.info(f"✅ 找到LibreOffice: {libreoffice_cmd}")
            else:
                # 尝试其他路径
                alternative_paths = [
                    r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
                ]
                
                libreoffice_cmd = None
                for path in alternative_paths:
                    if os.path.exists(path):
                        libreoffice_cmd = path
                        logger.info(f"✅ 找到LibreOffice: {path}")
                        break
            
            if not libreoffice_cmd:
                logger.error("❌ 未找到LibreOffice")
                raise RuntimeError("LibreOffice未安装或不可用。建议使用DOCX格式替代DOC格式。")
            
            # 步骤3: 设置无弹窗环境
            env = self.set_no_popup_environment()
            
            # 步骤4: 构建简化的无弹窗命令
            cmd = [
                libreoffice_cmd,
                '--headless',                    # 无头模式
                '--invisible',                   # 不可见模式
                '--nodefault',                   # 不使用默认文档
                '--nolockcheck',                 # 不检查文件锁
                '--nologo',                      # 不显示启动画面
                '--norestore',                   # 不恢复文档
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(doc_path),
                doc_path
            ]
            
            logger.info(f"📄 正在转换: {doc_path} -> {docx_path}")
            
            # 步骤5: 删除已存在的输出文件
            if os.path.exists(docx_path):
                os.remove(docx_path)
                logger.info("🗑️ 删除已存在的转换文件")
            
            # 步骤6: 执行转换（简化版）
            creation_flags = 0
            if os.name == 'nt':  # Windows
                creation_flags = subprocess.CREATE_NO_WINDOW
            
            result = subprocess.run(
                cmd, 
                capture_output=True, 
                text=True, 
                timeout=30,  # 减少超时时间
                env=env,
                creationflags=creation_flags
            )
            
            # 步骤7: 检查结果
            if result.returncode != 0:
                logger.error(f"❌ LibreOffice转换失败: {result.stderr}")
                raise RuntimeError(f"LibreOffice转换失败: {result.stderr}")
            
            # 步骤8: 验证输出文件
            expected_docx = doc_path.replace('.doc', '.docx')
            if os.path.exists(expected_docx):
                if expected_docx != docx_path:
                    os.rename(expected_docx, docx_path)
                
                logger.info(f"✅ 转换成功: {docx_path}")
                return docx_path
            else:
                logger.error(f"❌ 转换后的文件未找到: {expected_docx}")
                raise RuntimeError("转换后的文件未找到")
                
        except subprocess.TimeoutExpired:
            logger.error("❌ LibreOffice转换超时")
            raise RuntimeError("LibreOffice转换超时")
        except Exception as e:
            logger.error(f"❌ 转换过程中出错: {e}")
            raise
        finally:
            # 步骤9: 最终清理
            time.sleep(1)
            self.kill_all_libreoffice_processes()
    
    def _preprocess_template_and_extract_placeholders(self, doc_path: str, output_path: str) -> List[str]:
        """
        扩展占位符预处理，以包含通用的下划线字段，并优化替换逻辑
        """
        logger.info("🛠️  开始扩展占位符预处理...")
        
        self.placeholder_originals = {} # Reset for each new template analysis
        doc = Document(doc_path)
        placeholders = set()
        blank_counter = 0 # Counter for generic underscore placeholders
        
        def process_text_and_extract_keys(text: str) -> (str, List[str]):
            nonlocal blank_counter
            found_keys = []

            def repl_func(match):
                nonlocal blank_counter
                # Pattern for '致...': underscore_str in group(1), hint in group(2)
                if match.group(1) is not None:
                    if "（签字）" in match.group(0) or "(签字)" in match.group(0):
                        return match.group(0)
                    
                    underscore_str = match.group(1)
                    hint = match.group(2)
                    placeholder_key = f"inline_{hint}"
                    found_keys.append(placeholder_key)
                    self.placeholder_originals[placeholder_key] = underscore_str
                    replacement = f"致{{{placeholder_key}}}（{hint}）"
                    logger.info(f"   - 发现内联模式: '{match.group(0)}' -> '{replacement}'")
                    return replacement

                # Pattern for 'label:': label in group(3)
                elif match.group(3) is not None:
                    # The regex now prevents matching '（签字）:'
                    label = match.group(3).strip()
                    placeholder_key = f"label_{label}"
                    found_keys.append(placeholder_key)
                    replacement = f"{label}：{{{placeholder_key}}}"
                    logger.info(f"   - 发现标签模式: '{match.group(0)}' -> '{replacement}'")
                    return replacement

                # Pattern for general underscores: underscore_str in group(4)
                elif match.group(4) is not None:
                    underscore_str = match.group(4)
                    placeholder_key = f"blank_{blank_counter}"
                    found_keys.append(placeholder_key)
                    self.placeholder_originals[placeholder_key] = underscore_str
                    replacement = f"{{{placeholder_key}}}"
                    logger.info(f"   - 发现通用下划线模式: '{underscore_str}' -> '{replacement}'")
                    blank_counter += 1
                    return replacement
                
                return match.group(0)

            # Regex updated to handle spaced underscores and avoid capturing signature labels
            pattern = re.compile(
                r"致\s*(__{3,})\s*（([^）]+)）"              # G1: underscore, G2: hint
                r"|([^：\n（(]+?)：\s*$"                    # G3: label, avoids '(...):'
                r"|((?:_{4,}[\s\xa0]*)+)"               # G4: general underscore blocks
            )

            processed_text = pattern.sub(repl_func, text)
            
            return processed_text, found_keys
        
        # --- Process all paragraphs ---
        for para in doc.paragraphs:
            original_text = para.text
            if not original_text.strip():
                continue

            new_text, keys = process_text_and_extract_keys(original_text)
            if new_text != original_text:
                placeholders.update(keys)
                # To preserve formatting, we clear runs and add a new one
                para.clear()
                para.add_run(new_text)
                logger.info(f"   📝 段落更新: '{original_text.strip()}' -> '{new_text.strip()}'")

        # --- Process all tables ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    if not original_text.strip():
                        continue
                        
                    new_text, keys = process_text_and_extract_keys(original_text)
                    if new_text != original_text:
                        placeholders.update(keys)
                        # Reverted to cell.text for simplicity and correctness.
                        # This replaces the content of the first paragraph in the cell.
                        cell.text = new_text
                        logger.info(f"   📋 表格更新: '{original_text.strip()}' -> '{new_text.strip()}'")
        
        doc.save(output_path)
        logger.info(f"✅ 扩展预处理完成. 找到 {len(placeholders)} 个占位符. 新模板: {output_path}")
        return list(placeholders)

    def analyze_template_structure(self, template_path: str) -> Dict[str, str]:
        """
        确定性地分析Word模板，提取带有位置信息的结构。
        
        Args:
            template_path: .docx模板文件路径

        Returns:
            一个字典，其中键是单元格的唯一标识符，值是单元格的文本内容。
        """
        logger.info("🔍 开始确定性模板结构分析...")
        
        try:
            doc = Document(template_path)
            template_structure = {}
            
            logger.info(f"📄 正在读取模板文件: {template_path}")
            
            # 提取表格结构
            for i, table in enumerate(doc.tables):
                for j, row in enumerate(table.rows):
                    for k, cell in enumerate(row.cells):
                        cell_key = f"table_{i}_row_{j}_col_{k}"
                        template_structure[cell_key] = cell.text.strip()
            
            # 提取段落结构（不做特殊处理，保持原始内容）
            for i, para in enumerate(doc.paragraphs):
                para_key = f"paragraph_{i}"
                template_structure[para_key] = para.text.strip()
            
            logger.info(f"✅ 成功提取 {len(template_structure)} 个结构元素。")
            
            # Log a snippet of the extracted structure for verification
            structure_snippet = json.dumps(dict(list(template_structure.items())[:5]), ensure_ascii=False, indent=2)
            logger.info(f"  结构实例:\n{structure_snippet}")

            return template_structure
            
        except Exception as e:
            logger.error(f"❌ 模板结构分析错误: {e}")
            raise

    def convert_template_to_json(self, template_path: str, output_json_path: str = None) -> Dict[str, Any]:
        """
        主函数：将DOC/DOCX模板转换为JSON结构化数据
        
        Args:
            template_path: 模板文件路径（.doc或.docx）
            output_json_path: 输出JSON文件路径（可选）
            
        Returns:
            包含模板结构和占位符的字典
        """
        logger.info("🚀 开始模板转换流程...")
        
        try:
            # Step 1: Convert .doc to .docx if necessary
            if template_path.lower().endswith('.doc'):
                logger.info(f"📄 检测到.doc模板，开始转换: {template_path}")
                try:
                    original_docx_path = self.convert_doc_to_docx(template_path)
                except RuntimeError as e:
                    if "LibreOffice" in str(e):
                        logger.error(f"❌ DOC转换失败: {e}")
                        logger.error("💡 建议: 请直接使用DOCX格式的模板文件，或安装LibreOffice")
                        raise RuntimeError(f"DOC转换失败: {e}\n建议使用DOCX格式的模板文件")
                    else:
                        raise
            else:
                original_docx_path = template_path

            # Step 2: Preprocess template and extract placeholders
            processed_template_path = original_docx_path.replace(".docx", "_processed.docx")
            placeholders = self._preprocess_template_and_extract_placeholders(
                doc_path=original_docx_path,
                output_path=processed_template_path
            )
            
            # Step 3: Analyze processed template structure
            template_structure = self.analyze_template_structure(processed_template_path)

            # Step 4: Create final JSON output
            result = {
                "template_info": {
                    "original_path": template_path,
                    "processed_template_path": processed_template_path,
                    "conversion_timestamp": None  # Could add datetime if needed
                },
                "template_structure": template_structure,
                "placeholders": placeholders,
                "placeholder_originals": self.placeholder_originals
            }
            
            # Step 5: Save to file if requested
            if output_json_path:
                with open(output_json_path, 'w', encoding='utf-8') as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)
                logger.info(f"✅ JSON模板已保存到: {output_json_path}")
            
            logger.info(f"✅ 模板转换完成！")
            logger.info(f"   - 模板结构元素: {len(template_structure)}")
            logger.info(f"   - 占位符数量: {len(placeholders)}")
            
            return result
            
        except Exception as e:
            logger.error(f"❌ 模板转换失败: {e}", exc_info=True)
            raise


class TemplateConversionTool:
    """模板转换工具 - ReAct Agent工具接口"""
    
    name = "template_conversion"
    description = """模板转换工具：将DOC/DOCX模板文件转换为JSON结构化数据
    
功能：
1. 自动转换DOC文件为DOCX格式
2. 智能识别和提取模板占位符
3. 分析模板结构（段落、表格）
4. 输出JSON格式的结构化模板数据

适用场景：
- 处理旧版DOC格式模板文件
- 提取模板中的占位符信息
- 分析模板文档结构
- 为模板填充做准备工作

参数：
- template_path: 模板文件路径（必需）
- output_json_path: 输出JSON文件路径（可选）
- save_processed_template: 是否保存处理后的模板文件（可选，默认True）
"""
    
    def __init__(self, ai_client=None):
        """初始化模板转换工具"""
        self.ai_client = ai_client
        self.converter = TemplateConverter()
        logger.info("🔧 模板转换工具初始化完成")
    
    def execute(self, template_path: str, output_json_path: str = None, save_processed_template: bool = True) -> str:
        """
        执行模板转换
        
        Args:
            template_path: 模板文件路径
            output_json_path: 输出JSON文件路径（可选）
            save_processed_template: 是否保存处理后的模板文件
            
        Returns:
            转换结果的文本描述
        """
        try:
            logger.info(f"🚀 开始执行模板转换: {template_path}")
            
            # 检查文件是否存在
            if not os.path.exists(template_path):
                error_msg = f"❌ 模板文件不存在: {template_path}"
                logger.error(error_msg)
                return error_msg
            
            # 检查文件格式
            if not (template_path.lower().endswith('.doc') or template_path.lower().endswith('.docx')):
                error_msg = f"❌ 不支持的文件格式，仅支持DOC和DOCX格式: {template_path}"
                logger.error(error_msg)
                return error_msg
            
            # 生成默认输出路径
            if not output_json_path:
                base_name = os.path.splitext(template_path)[0]
                output_json_path = f"{base_name}_template_structure.json"
            
            # 执行转换
            result = self.converter.convert_template_to_json(template_path, output_json_path)
            
            # 构建结果报告
            report = f"""✅ 模板转换成功完成！

📊 转换统计：
- 原始模板: {template_path}
- 模板结构元素: {len(result['template_structure'])} 个
- 识别占位符: {len(result['placeholders'])} 个
- JSON输出文件: {output_json_path}

🔍 发现的占位符类型：
"""
            
            # 分析占位符类型
            placeholder_types = {}
            for placeholder in result['placeholders']:
                if placeholder.startswith('inline_'):
                    placeholder_types['内联提示'] = placeholder_types.get('内联提示', 0) + 1
                elif placeholder.startswith('label_'):
                    placeholder_types['标签字段'] = placeholder_types.get('标签字段', 0) + 1
                elif placeholder.startswith('blank_'):
                    placeholder_types['通用下划线'] = placeholder_types.get('通用下划线', 0) + 1
                else:
                    placeholder_types['其他'] = placeholder_types.get('其他', 0) + 1
            
            for ptype, count in placeholder_types.items():
                report += f"- {ptype}: {count} 个\n"
            
            # 添加占位符列表
            if result['placeholders']:
                report += f"\n📝 占位符列表:\n"
                for i, placeholder in enumerate(result['placeholders'][:10], 1):  # 只显示前10个
                    report += f"{i}. {placeholder}\n"
                
                if len(result['placeholders']) > 10:
                    report += f"... 还有 {len(result['placeholders']) - 10} 个占位符\n"
            
            # 添加处理后的模板信息
            if 'processed_template_path' in result['template_info']:
                processed_path = result['template_info']['processed_template_path']
                if save_processed_template and os.path.exists(processed_path):
                    report += f"\n📄 处理后的模板: {processed_path}"
                else:
                    # 如果不保存，删除处理后的模板
                    if os.path.exists(processed_path):
                        os.remove(processed_path)
                        report += f"\n🗑️ 临时处理文件已清理"
            
            report += f"\n\n💡 提示: 可以使用JSON文件进行后续的模板填充操作"
            
            logger.info("✅ 模板转换工具执行完成")
            return report
            
        except Exception as e:
            error_msg = f"❌ 模板转换失败: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return error_msg
    
    def get_help(self) -> str:
        """获取工具帮助信息"""
        return f"""
🔧 {self.name} - 模板转换工具

{self.description}

📋 使用方法:
1. 基本转换: template_conversion(template_path="path/to/template.doc")
2. 指定输出: template_conversion(template_path="template.docx", output_json_path="output.json")
3. 不保存处理文件: template_conversion(template_path="template.doc", save_processed_template=False)

💡 支持的文件格式:
- Microsoft Word DOC文件（需要LibreOffice）
- Microsoft Word DOCX文件

🎯 输出内容:
- JSON格式的模板结构数据
- 识别的占位符列表
- 原始占位符映射关系
- 处理后的模板文件（可选）

⚠️ 注意事项:
- DOC转换需要安装LibreOffice
- 确保模板文件路径正确
- 输出目录需要有写入权限
"""


def main():
    """命令行测试函数"""
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python template_conversion_tool.py <template_path> [output_json_path]")
        return
    
    template_path = sys.argv[1]
    output_json_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(template_path):
        print(f"❌ 模板文件不存在: {template_path}")
        return
    
    # 测试工具
    tool = TemplateConversionTool()
    result = tool.execute(template_path, output_json_path)
    print(result)


if __name__ == "__main__":
    main() 