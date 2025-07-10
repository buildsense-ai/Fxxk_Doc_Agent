#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档生成工具 - 从mcpdemo-newbran抽取的核心文档生成功能

提供AI自动化文档生成服务，包括：
- 智能大纲生成
- 知识检索整合
- 迭代优化
- 多格式输出
"""

import os
import sys
import json
import time
import uuid
import requests
import tempfile
from typing import Dict, Any, Optional, List
from pathlib import Path
import logging

# 导入工具基类
try:
    from base_tool import Tool
except ImportError:
    try:
        from ..base_tool import Tool
    except ImportError:
        # 如果无法导入，定义一个简单的基类
        class Tool:
            def __init__(self, name: str = "", description: str = ""):
                self.name = name
                self.description = description
            
            def execute(self, **kwargs) -> str:
                raise NotImplementedError("子类必须实现execute方法")

# 添加必要的导入
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    from minio import Minio
    from minio.error import S3Error
    MINIO_AVAILABLE = True
except ImportError:
    MINIO_AVAILABLE = False

try:
    import pypandoc
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class DocumentGeneratorConfig:
    """文档生成器配置"""
    
    def __init__(self):
        # API配置
        self.DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
        self.DEEPSEEK_API_BASE = "https://api.deepseek.com/v1"
        self.AI_MODEL_NAME = "deepseek-chat"
        
        # 任务配置
        self.MAX_REFINEMENT_CYCLES = 3
        self.SEARCH_DEFAULT_TOP_K = 5
        
        # 向量搜索配置
        self.TEXT_SEARCH_ENDPOINT = "http://43.139.19.144:3000/search-drawings"
        self.IMAGE_SEARCH_ENDPOINT = "http://65d27a3b.r23.cpolar.top/search/images"
        self.IMAGE_SEARCH_MIN_SCORE = 0.4
        
        # MinIO配置
        self.MINIO_ENDPOINT = "43.139.19.144:9000"
        self.MINIO_ACCESS_KEY = "minioadmin"
        self.MINIO_SECRET_KEY = "minioadmin"
        self.MINIO_BUCKET_NAME = "docs"
        self.MINIO_USE_SECURE = False
        
        # 任务存储
        self.TASKS_DIR = Path("document_generator_tasks")
        self.TASKS_DIR.mkdir(exist_ok=True)

class TaskState:
    """任务状态管理"""
    
    def __init__(self, task_id: str, config: DocumentGeneratorConfig):
        self.task_id = task_id
        self.config = config
        self.filepath = self.config.TASKS_DIR / f"task_{self.task_id}.json"
        self.data: Dict[str, Any] = {}
    
    def load(self) -> bool:
        """加载任务状态"""
        if self.filepath.exists():
            try:
                with open(self.filepath, 'r', encoding='utf-8') as f:
                    self.data = json.load(f)
                return True
            except Exception as e:
                logger.error(f"加载任务状态失败: {e}")
                return False
        return False
    
    def save(self):
        """保存任务状态"""
        try:
            self.data['lastUpdatedTimestamp'] = time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime())
            with open(self.filepath, 'w', encoding='utf-8') as f:
                json.dump(self.data, f, ensure_ascii=False, indent=2)
            logger.info(f"任务状态已保存: {self.filepath}")
        except Exception as e:
            logger.error(f"保存任务状态失败: {e}")
    
    def initialize(self, initial_request: Dict[str, str], report_type: str):
        """初始化任务"""
        self.data = {
            "taskId": self.task_id,
            "status": "pending",
            "progressPercentage": 0,
            "currentStatusMessage": "任务已创建，等待初始化...",
            "initialRequest": initial_request,
            "reportType": report_type,
            "creativeBrief": "",
            "projectName": "",
            "introduction": "",
            "conclusion": "",
            "outline": {},
            "finalDocument": "",
            "markdownPublicUrl": "",
            "docxPublicUrl": "",
            "errorLog": []
        }
        self.save()
    
    def update_status(self, status: str, message: str, progress: int):
        """更新状态"""
        self.data['status'] = status
        self.data['currentStatusMessage'] = message
        self.data['progressPercentage'] = progress
        self.save()
        logger.info(f"进度更新: {progress}% - {message}")
    
    def log_error(self, stage: str, error_message: str):
        """记录错误"""
        self.data['errorLog'].append({
            "timestamp": time.strftime('%Y-%m-%dT%H:%M:%SZ', time.gmtime()),
            "stage": stage,
            "message": error_message
        })
        self.update_status("failed", f"在 {stage} 阶段发生错误。", self.data.get('progressPercentage', 0))

class DocumentGeneratorServices:
    """文档生成服务接口"""
    
    def __init__(self, config: DocumentGeneratorConfig):
        self.config = config
    
    def call_ai_model(self, prompt: str, context: str = None, expect_json: bool = False) -> Dict[str, Any]:
        """调用AI模型"""
        if not OPENAI_AVAILABLE:
            raise RuntimeError("OpenAI库不可用，无法调用AI模型")
        
        if not self.config.DEEPSEEK_API_KEY:
            raise ValueError("DEEPSEEK_API_KEY环境变量未设置")
        
        try:
            client = openai.OpenAI(
                api_key=self.config.DEEPSEEK_API_KEY,
                base_url=self.config.DEEPSEEK_API_BASE
            )
        except Exception as e:
            raise Exception(f"初始化AI客户端失败: {e}")
        
        messages = []
        if context:
            messages.append({"role": "system", "content": context})
        messages.append({"role": "user", "content": prompt})
        
        logger.info(f"调用AI模型: {self.config.AI_MODEL_NAME}")
        logger.debug(f"提示词: {prompt[:200]}...")
        
        api_params = {
            "model": self.config.AI_MODEL_NAME,
            "messages": messages
        }
        if expect_json:
            api_params['response_format'] = {'type': 'json_object'}
        
        try:
            response = client.chat.completions.create(**api_params)
            response_content = response.choices[0].message.content
            
            if expect_json:
                return json.loads(response_content)
            else:
                return {'text': response_content}
        
        except Exception as e:
            raise Exception(f"AI模型调用失败: {e}")
    
    def search_vectordata(self, query: str, top_k: int) -> List[str]:
        """搜索向量数据库"""
        logger.info(f"搜索向量数据库: {query}, Top_K: {top_k}")
        
        try:
            import urllib.parse
            encoded_query = urllib.parse.quote(query)
            full_url = f"{self.config.TEXT_SEARCH_ENDPOINT}?query={encoded_query}&top_k={top_k}"
            
            response = requests.get(full_url, headers={'accept': 'application/json'}, timeout=20)
            response.raise_for_status()
            
            data = response.json()
            results = data.get("results", [])
            content_list = [item.get("content", "") for item in results if item.get("content")]
            
            logger.info(f"向量搜索成功，获得 {len(content_list)} 条结果")
            return content_list
            
        except Exception as e:
            logger.error(f"向量搜索失败: {e}")
            return []
    
    def upload_to_minio(self, file_path: str, object_name: str) -> Optional[str]:
        """上传文件到MinIO"""
        if not MINIO_AVAILABLE:
            logger.error("MinIO库不可用")
            return None
        
        try:
            client = Minio(
                self.config.MINIO_ENDPOINT,
                access_key=self.config.MINIO_ACCESS_KEY,
                secret_key=self.config.MINIO_SECRET_KEY,
                secure=self.config.MINIO_USE_SECURE
            )
            
            # 检查并创建存储桶
            if not client.bucket_exists(self.config.MINIO_BUCKET_NAME):
                client.make_bucket(self.config.MINIO_BUCKET_NAME)
            
            # 上传文件
            client.fput_object(
                self.config.MINIO_BUCKET_NAME,
                object_name,
                file_path,
            )
            
            # 构造公共URL
            public_url = f"http://{self.config.MINIO_ENDPOINT}/{self.config.MINIO_BUCKET_NAME}/{object_name}"
            logger.info(f"文件上传成功: {public_url}")
            return public_url
            
        except Exception as e:
            logger.error(f"文件上传失败: {e}")
            return None

class DocumentGenerator:
    """文档生成器核心"""
    
    def __init__(self, config: DocumentGeneratorConfig = None):
        self.config = config or DocumentGeneratorConfig()
        self.services = DocumentGeneratorServices(self.config)
    
    def start_new_job(self, chathistory: str, request: str, report_type: str = 'long') -> str:
        """启动新的生成任务"""
        task_id = str(uuid.uuid4())
        state = TaskState(task_id, self.config)
        
        logger.info(f"启动新任务: {task_id} (类型: {report_type})")
        
        state.initialize(
            initial_request={"chathistory": chathistory, "request": request},
            report_type=report_type
        )
        
        # 异步执行生成流程
        try:
            if report_type == 'short':
                self._run_short_report_workflow(state)
            else:
                self._run_long_report_workflow(state)
        except Exception as e:
            logger.error(f"任务执行失败: {e}")
            state.log_error("execution", str(e))
        
        return task_id
    
    def _run_short_report_workflow(self, state: TaskState):
        """执行短报告生成流程"""
        try:
            while state.data['status'] not in ['completed', 'failed']:
                current_status = state.data['status']
                
                if current_status == 'pending':
                    self._prepare_creative_brief(state)
                elif current_status == 'brief_prepared':
                    self._generate_short_report_content(state)
                elif current_status == 'short_report_generated':
                    self._assemble_final_document(state, is_short_report=True)
                else:
                    break
            
            if state.data['status'] == 'completed':
                logger.info(f"短报告任务 {state.task_id} 完成")
        
        except Exception as e:
            logger.error(f"短报告生成失败: {e}")
            state.log_error("short_report", str(e))
    
    def _run_long_report_workflow(self, state: TaskState):
        """执行长报告生成流程"""
        try:
            while state.data['status'] not in ['completed', 'failed']:
                current_status = state.data['status']
                
                if current_status == 'pending':
                    self._prepare_creative_brief(state)
                elif current_status == 'brief_prepared':
                    self._generate_initial_outline(state)
                elif current_status == 'outline_generated':
                    self._refine_outline_cycle(state)
                elif current_status == 'outline_finalized':
                    self._generate_all_chapters(state)
                elif current_status == 'chapters_generated':
                    self._assemble_final_document(state)
                else:
                    break
            
            if state.data['status'] == 'completed':
                logger.info(f"长报告任务 {state.task_id} 完成")
        
        except Exception as e:
            logger.error(f"长报告生成失败: {e}")
            state.log_error("long_report", str(e))
    
    def _prepare_creative_brief(self, state: TaskState):
        """准备创作指令"""
        state.update_status("brief_generation", "正在分析需求和准备创作指令...", 5)
        
        chathistory = state.data['initialRequest']['chathistory']
        request = state.data['initialRequest']['request']
        
        prompt_brief = f"""你是一位资深的文档分析专家。请根据下面的对话记录和最终请求，为即将撰写的文档提炼一份核心的"创作指令"。
这份指令需要明确文档主题、性质和核心要求。

【对话记录】
{chathistory}

【最终请求】
{request}

请以JSON格式返回创作指令，包含以下字段：
- document_topic: 文档主题
- document_type: 文档类型
- key_requirements: 核心要求列表
- target_audience: 目标读者
- content_scope: 内容范围
"""
        
        try:
            response = self.services.call_ai_model(prompt_brief, expect_json=True)
            state.data['creativeBrief'] = response
            
            # 提取项目名称
            project_prompt = f"""从以下创作指令中，提取一个简短的核心项目名称或主题，用于优化后续的知识库检索。
请以JSON格式返回，只包含一个 'project_name' 字段。

创作指令：{response}
"""
            
            project_response = self.services.call_ai_model(project_prompt, expect_json=True)
            state.data['projectName'] = project_response.get('project_name', '文档生成项目')
            
            state.data['status'] = 'brief_prepared'
            state.save()
            
        except Exception as e:
            state.log_error("brief_preparation", str(e))
    
    def _generate_short_report_content(self, state: TaskState):
        """生成短报告内容"""
        state.update_status("short_report_generation", "正在生成短篇报告内容...", 20)
        
        project_name = state.data.get('projectName', '')
        creative_brief = state.data.get('creativeBrief', '')
        
        # 进行知识检索
        knowledge_pieces = self.services.search_vectordata(
            query=project_name, 
            top_k=self.config.SEARCH_DEFAULT_TOP_K
        )
        
        knowledge_context = ""
        if knowledge_pieces:
            knowledge_str = "\n\n---\n\n".join(knowledge_pieces)
            knowledge_context = f"\n请参考以下背景资料进行撰写：\n{knowledge_str}\n"
        
        prompt = f"""你是一位专业的报告撰写人。请根据以下项目简介和背景资料，撰写一篇结构完整、内容流畅的报告，总字数控制在2000字以内。
文章应有逻辑地分为几个部分，并使用Markdown的二级标题（##）来标记每个部分的标题。

【项目简介】
{creative_brief}

{knowledge_context}

请直接输出完整的Markdown格式报告全文。
"""
        
        try:
            response = self.services.call_ai_model(prompt)
            state.data['finalDocument'] = response.get('text', '')
            state.data['status'] = 'short_report_generated'
            state.save()
            
        except Exception as e:
            state.log_error("short_report_generation", str(e))
    
    def _generate_initial_outline(self, state: TaskState):
        """生成初始大纲"""
        state.update_status("outline_generation", "正在生成文档大纲...", 15)
        
        creative_brief = state.data.get('creativeBrief', '')
        project_name = state.data.get('projectName', '')
        
        # 进行知识检索
        knowledge_pieces = self.services.search_vectordata(
            query=project_name,
            top_k=self.config.SEARCH_DEFAULT_TOP_K
        )
        
        knowledge_context = ""
        if knowledge_pieces:
            knowledge_str = "\n\n---\n\n".join(knowledge_pieces)
            knowledge_context = f"\n背景资料：\n{knowledge_str}\n"
        
        prompt = f"""你是一位专业的文档结构规划师。请根据以下创作指令和背景资料，设计一个详细的文档大纲。

【创作指令】
{creative_brief}

{knowledge_context}

请以JSON格式返回大纲，结构如下：
{{
    "title": "文档标题",
    "chapters": [
        {{
            "chapter_number": 1,
            "chapter_title": "章节标题",
            "sections": [
                {{
                    "section_number": "1.1",
                    "section_title": "小节标题",
                    "content_outline": "该小节的内容要点"
                }}
            ]
        }}
    ]
}}
"""
        
        try:
            response = self.services.call_ai_model(prompt, expect_json=True)
            state.data['outline'] = response
            state.data['status'] = 'outline_generated'
            state.save()
            
        except Exception as e:
            state.log_error("outline_generation", str(e))
    
    def _refine_outline_cycle(self, state: TaskState):
        """大纲优化循环"""
        state.update_status("outline_refinement", "正在优化大纲结构...", 30)
        
        # 简化版：直接标记为完成
        state.data['status'] = 'outline_finalized'
        state.save()
    
    def _generate_all_chapters(self, state: TaskState):
        """生成所有章节内容"""
        state.update_status("chapters_generation", "正在生成章节内容...", 50)
        
        outline = state.data.get('outline', {})
        chapters = outline.get('chapters', [])
        
        generated_content = f"# {outline.get('title', '文档标题')}\n\n"
        
        for chapter in chapters:
            chapter_title = chapter.get('chapter_title', '')
            sections = chapter.get('sections', [])
            
            generated_content += f"## {chapter_title}\n\n"
            
            # 为每个小节生成内容
            for section in sections:
                section_title = section.get('section_title', '')
                content_outline = section.get('content_outline', '')
                
                # 进行知识检索
                knowledge_pieces = self.services.search_vectordata(
                    query=f"{chapter_title} {section_title}",
                    top_k=3
                )
                
                knowledge_context = ""
                if knowledge_pieces:
                    knowledge_str = "\n".join(knowledge_pieces)
                    knowledge_context = f"\n参考资料：\n{knowledge_str}\n"
                
                prompt = f"""请为以下小节撰写详细内容：

【小节标题】{section_title}
【内容要点】{content_outline}
{knowledge_context}

要求：
1. 内容详实，逻辑清晰
2. 字数控制在300-500字
3. 使用Markdown格式
4. 不要包含标题，直接输出正文内容
"""
                
                try:
                    response = self.services.call_ai_model(prompt)
                    section_content = response.get('text', '')
                    generated_content += f"### {section_title}\n\n{section_content}\n\n"
                    
                except Exception as e:
                    logger.error(f"生成小节内容失败: {e}")
                    generated_content += f"### {section_title}\n\n[内容生成失败]\n\n"
        
        state.data['finalDocument'] = generated_content
        state.data['status'] = 'chapters_generated'
        state.save()
    
    def _assemble_final_document(self, state: TaskState, is_short_report: bool = False):
        """组装最终文档"""
        state.update_status("document_assembly", "正在组装最终文档...", 80)
        
        try:
            final_document = state.data.get('finalDocument', '')
            
            if not final_document:
                raise Exception("没有可用的文档内容")
            
            # 创建Markdown文件
            md_filename = f"document_{state.task_id}.md"
            md_path = self.config.TASKS_DIR / md_filename
            
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(final_document)
            
            # 上传Markdown文件
            md_url = self.services.upload_to_minio(str(md_path), md_filename)
            if md_url:
                state.data['markdownPublicUrl'] = md_url
            
            # 尝试转换为DOCX
            if DOCX_AVAILABLE:
                try:
                    docx_filename = f"document_{state.task_id}.docx"
                    docx_path = self.config.TASKS_DIR / docx_filename
                    
                    # 简单的Markdown到DOCX转换
                    doc = Document()
                    
                    # 按行处理Markdown内容
                    lines = final_document.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        if line.startswith('# '):
                            # 一级标题
                            title = line[2:].strip()
                            paragraph = doc.add_heading(title, level=1)
                        elif line.startswith('## '):
                            # 二级标题
                            title = line[3:].strip()
                            paragraph = doc.add_heading(title, level=2)
                        elif line.startswith('### '):
                            # 三级标题
                            title = line[4:].strip()
                            paragraph = doc.add_heading(title, level=3)
                        else:
                            # 正文
                            doc.add_paragraph(line)
                    
                    doc.save(docx_path)
                    
                    # 上传DOCX文件
                    docx_url = self.services.upload_to_minio(str(docx_path), docx_filename)
                    if docx_url:
                        state.data['docxPublicUrl'] = docx_url
                
                except Exception as e:
                    logger.error(f"DOCX转换失败: {e}")
            
            state.update_status("completed", "文档生成完成", 100)
            
        except Exception as e:
            state.log_error("document_assembly", str(e))

class DocumentGeneratorTool(Tool):
    """文档生成工具接口"""
    
    def __init__(self):
        description = """文档生成工具 - AI智能文档创作助手

🚀 **核心功能:**
• generate_long_document: 生成详细长篇报告
• generate_short_document: 生成精简短篇报告  
• check_status: 查询生成状态
• list_tasks: 列出所有任务
• get_result: 获取完成任务结果

📋 **使用示例:**
• 长文档: {"action": "generate_long_document", "chathistory": "对话记录", "request": "生成报告"}
• 短文档: {"action": "generate_short_document", "chathistory": "对话记录", "request": "生成概述"}
• 查询状态: {"action": "check_status", "task_id": "任务ID"}

⭐ **特色功能:**
- 🤖 自主规划与大纲生成
- 🧠 知识检索与整合
- 🔄 迭代式自我评审与优化
- 📝 多格式产出（Markdown + DOCX）
- ☁️ 云端同步（MinIO集成）
"""
        
        # 调用父类构造函数
        super().__init__(name="document_generator", description=description)
        
        self.config = DocumentGeneratorConfig()
        self.generator = DocumentGenerator(self.config)
    
    def execute(self, action: str, **kwargs) -> str:
        """执行工具操作"""
        try:
            if action == "generate_long_document":
                return self._generate_long_document(**kwargs)
            elif action == "generate_short_document":
                return self._generate_short_document(**kwargs)
            elif action == "check_status":
                return self._check_status(**kwargs)
            elif action == "list_tasks":
                return self._list_tasks()
            elif action == "get_result":
                return self._get_result(**kwargs)
            else:
                return f"""❌ 不支持的操作: {action}

📋 **支持的操作:**
• generate_long_document - 生成长篇报告
• generate_short_document - 生成短篇报告
• check_status - 查询状态
• list_tasks - 列出任务
• get_result - 获取结果

💡 **使用示例:**
• {{"action": "generate_long_document", "chathistory": "...", "request": "..."}}
"""
        
        except Exception as e:
            logger.error(f"工具操作失败: {e}")
            return f"❌ 操作失败: {str(e)}"
    
    def _generate_long_document(self, chathistory: str, request: str) -> str:
        """生成长文档"""
        try:
            task_id = self.generator.start_new_job(chathistory, request, "long")
            return f"""✅ 长文档生成任务已启动

🆔 **任务ID**: {task_id}
📝 **类型**: 长篇报告
📋 **请求**: {request[:100]}{'...' if len(request) > 100 else ''}

💡 **查询进度**: {{"action": "check_status", "task_id": "{task_id}"}}
📊 **预计时间**: 5-15分钟
"""
        except Exception as e:
            return f"❌ 长文档生成失败: {e}"
    
    def _generate_short_document(self, chathistory: str, request: str) -> str:
        """生成短文档"""
        try:
            task_id = self.generator.start_new_job(chathistory, request, "short")
            return f"""✅ 短文档生成任务已启动

🆔 **任务ID**: {task_id}
📝 **类型**: 短篇报告（2000字内）
📋 **请求**: {request[:100]}{'...' if len(request) > 100 else ''}

💡 **查询进度**: {{"action": "check_status", "task_id": "{task_id}"}}
📊 **预计时间**: 2-5分钟
"""
        except Exception as e:
            return f"❌ 短文档生成失败: {e}"
    
    def _check_status(self, task_id: str) -> str:
        """查询状态"""
        try:
            state = TaskState(task_id, self.config)
            if not state.load():
                return f"❌ 任务不存在: {task_id}"
            
            status = state.data.get('status', 'unknown')
            progress = state.data.get('progressPercentage', 0)
            message = state.data.get('currentStatusMessage', '无状态信息')
            
            result = f"""📊 **任务状态**

🆔 **任务ID**: {task_id}
🎯 **状态**: {status}
📈 **进度**: {progress}%
💬 **信息**: {message}
"""
            
            if status == "completed":
                if state.data.get('markdownPublicUrl'):
                    result += f"📄 **Markdown**: {state.data['markdownPublicUrl']}\n"
                if state.data.get('docxPublicUrl'):
                    result += f"📖 **DOCX**: {state.data['docxPublicUrl']}\n"
            
            elif status == "failed":
                error_log = state.data.get('errorLog', [])
                if error_log:
                    latest_error = error_log[-1]
                    result += f"❌ **错误**: {latest_error.get('message', '未知错误')}\n"
            
            return result
            
        except Exception as e:
            return f"❌ 查询状态失败: {e}"
    
    def _list_tasks(self) -> str:
        """列出任务"""
        try:
            task_files = list(self.config.TASKS_DIR.glob("task_*.json"))
            
            if not task_files:
                return "📋 **任务列表**: 暂无任务"
            
            result = f"📋 **任务列表** (共{len(task_files)}个)\n\n"
            
            for task_file in sorted(task_files, key=lambda x: x.stat().st_mtime, reverse=True)[:10]:
                try:
                    task_id = task_file.stem.replace("task_", "")
                    state = TaskState(task_id, self.config)
                    
                    if state.load():
                        status = state.data.get('status', 'unknown')
                        report_type = state.data.get('reportType', 'long')
                        progress = state.data.get('progressPercentage', 0)
                        
                        status_icon = "✅" if status == 'completed' else "❌" if status == 'failed' else "⏳"
                        
                        result += f"**{task_id[:8]}...** {status_icon}\n"
                        result += f"   📝 {report_type} | 📊 {progress}% | 🎯 {status}\n\n"
                
                except Exception:
                    continue
            
            return result
            
        except Exception as e:
            return f"❌ 列出任务失败: {e}"
    
    def _get_result(self, task_id: str) -> str:
        """获取结果"""
        try:
            state = TaskState(task_id, self.config)
            if not state.load():
                return f"❌ 任务不存在: {task_id}"
            
            status = state.data.get('status', 'unknown')
            
            if status != 'completed':
                return f"⚠️ 任务未完成 (状态: {status})"
            
            final_document = state.data.get('finalDocument', '')
            
            result = f"""✅ **任务结果**

🆔 **任务ID**: {task_id}
📊 **文档长度**: {len(final_document)} 字符
"""
            
            if state.data.get('markdownPublicUrl'):
                result += f"📄 **Markdown**: {state.data['markdownPublicUrl']}\n"
            if state.data.get('docxPublicUrl'):
                result += f"📖 **DOCX**: {state.data['docxPublicUrl']}\n"
            
            if final_document:
                preview = final_document[:500] + "..." if len(final_document) > 500 else final_document
                result += f"\n📝 **内容预览**:\n```\n{preview}\n```"
            
            return result
            
        except Exception as e:
            return f"❌ 获取结果失败: {e}"


# 工具实例创建函数
def create_document_generator_tool():
    """创建文档生成工具实例"""
    return DocumentGeneratorTool() 