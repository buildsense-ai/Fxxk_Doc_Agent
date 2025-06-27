#!/usr/bin/env python3
"""
高级长文档生成工具
基于long_generator项目封装的专业长文档生成器
"""

import json
import os
import time
import uuid
import logging
from typing import Dict, Any, Optional, List
from datetime import datetime
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

# 导入必要的库
try:
    import openai
    import requests
    import urllib.parse
    from minio import Minio
    from minio.error import S3Error
except ImportError as e:
    logging.warning(f"部分依赖库未安装: {e}")

from .tools import Tool

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AdvancedLongDocumentConfig:
    """高级长文档生成器配置"""
    
    def __init__(self):
        # 任务状态文件存储目录
        self.TASKS_DIR = "long_document_tasks"
        
        # DeepSeek API配置
        self.DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY") or os.getenv("OPENROUTER_API_KEY")
        self.DEEPSEEK_API_BASE = "https://api.deepseek.com/v1"
        self.AI_MODEL_NAME = "deepseek-chat"
        
        # 大纲精炼的最大循环次数
        self.MAX_REFINEMENT_CYCLES = 3
        
        # 向量搜索API配置
        self.SEARCH_API_BASE = "http://43.139.19.144:3000"
        self.SEARCH_API_TOP_K = 5
        
        # MinIO云存储配置
        self.MINIO_ENDPOINT = "43.139.19.144:9000"
        self.MINIO_ACCESS_KEY = "minioadmin"
        self.MINIO_SECRET_KEY = "minioadmin"
        self.MINIO_BUCKET_NAME = "docs"
        self.MINIO_USE_SECURE = False
        
        # 确保任务目录存在
        os.makedirs(self.TASKS_DIR, exist_ok=True)

class TaskState:
    """任务状态管理器"""
    
    def __init__(self, task_id: str, config: AdvancedLongDocumentConfig):
        self.task_id = task_id
        self.config = config
        self.filepath = os.path.join(config.TASKS_DIR, f"task_{task_id}.json")
        self.data: Dict[str, Any] = {}
    
    def load(self) -> bool:
        """加载任务状态"""
        if os.path.exists(self.filepath):
            try:
                with open(self.filepath, 'r', encoding='utf-8') as f:
                    self.data = json.load(f)
                return True
            except Exception as e:
                logger.error(f"加载任务状态失败: {e}")
        return False
    
    def save(self):
        """保存任务状态"""
        try:
            self.data['lastUpdatedTimestamp'] = datetime.now().isoformat()
            with open(self.filepath, 'w', encoding='utf-8') as f:
                json.dump(self.data, f, ensure_ascii=False, indent=2)
            logger.info(f"任务状态已保存: {self.filepath}")
        except Exception as e:
            logger.error(f"保存任务状态失败: {e}")
    
    def initialize(self, initial_request: Dict[str, str]):
        """初始化任务状态"""
        self.data = {
            "taskId": self.task_id,
            "status": "pending",
            "progressPercentage": 0,
            "currentStatusMessage": "任务已创建，等待初始化...",
            "initialRequest": initial_request,
            "creativeBrief": "",
            "projectName": "",
            "outline": {},
            "finalDocument": "",
            "docxPublicUrl": "",
            "docxLocalPath": "",
            "errorLog": []
        }
        self.save()
    
    def update_status(self, status: str, message: str, progress: int):
        """更新任务状态"""
        self.data['status'] = status
        self.data['currentStatusMessage'] = message
        self.data['progressPercentage'] = progress
        self.save()
        logger.info(f"进度更新: {progress}% - {message}")
    
    def log_error(self, stage: str, error_message: str):
        """记录错误"""
        self.data['errorLog'].append({
            "timestamp": datetime.now().isoformat(),
            "stage": stage,
            "message": error_message
        })
        self.update_status("failed", f"在 {stage} 阶段发生错误。", self.data.get('progressPercentage', 0))

class AIService:
    """AI服务封装"""
    
    def __init__(self, config: AdvancedLongDocumentConfig):
        self.config = config
    
    def call_ai_model(self, prompt: str, context: str = None, expect_json: bool = False) -> Dict[str, Any]:
        """调用AI模型"""
        if not self.config.DEEPSEEK_API_KEY:
            raise ValueError("未设置DEEPSEEK_API_KEY或OPENROUTER_API_KEY环境变量")
        
        try:
            # 创建httpx客户端，禁用代理以避免连接问题，增加超时时间
            import httpx
            http_client = httpx.Client(
                timeout=120.0,  # 增加到2分钟超时
                limits=httpx.Limits(max_keepalive_connections=5, max_connections=10)
            )
            client = openai.OpenAI(
                api_key=self.config.DEEPSEEK_API_KEY,
                base_url=self.config.DEEPSEEK_API_BASE,
                http_client=http_client,
                timeout=120.0  # OpenAI客户端也设置2分钟超时
            )
        except Exception as e:
            raise Exception(f"初始化AI客户端失败: {e}")
        
        messages = []
        if context:
            messages.append({"role": "system", "content": context})
        messages.append({"role": "user", "content": prompt})
        
        logger.info(f"调用AI模型: {self.config.AI_MODEL_NAME}")
        
        api_params = {
            "model": self.config.AI_MODEL_NAME,
            "messages": messages
        }
        if expect_json:
            api_params['response_format'] = {'type': 'json_object'}
        
        try:
            response = client.chat.completions.create(**api_params)
            response_content = response.choices[0].message.content
            
            if expect_json:
                return json.loads(response_content)
            else:
                return {'text': response_content}
                
        except Exception as e:
            raise Exception(f"AI模型调用失败: {e}")
    
    def search_vectordata(self, query: str, top_k: int = None) -> List[str]:
        """搜索向量数据库"""
        if top_k is None:
            top_k = self.config.SEARCH_API_TOP_K
            
        logger.info(f"搜索向量数据库: {query}")
        
        try:
            base_url = f"{self.config.SEARCH_API_BASE}/search-drawings"
            encoded_query = urllib.parse.quote(query)
            full_url = f"{base_url}?query={encoded_query}&top_k={top_k}"
            
            headers = {'accept': 'application/json'}
            # 禁用代理以避免连接问题
            proxies = {'http': '', 'https': ''}
            response = requests.get(full_url, headers=headers, timeout=20, proxies=proxies)
            response.raise_for_status()
            
            data = response.json()
            results = data.get("results", [])
            content_list = [item.get("content", "") for item in results if item.get("content")]
            
            logger.info(f"向量搜索成功，获得 {len(content_list)} 条结果")
            return content_list
            
        except Exception as e:
            logger.warning(f"向量搜索失败: {e}")
            return []

class DocumentService:
    """文档服务"""
    
    def __init__(self, config: AdvancedLongDocumentConfig):
        self.config = config
    
    def create_docx(self, final_text: str, project_name: str, task_id: str) -> str:
        """创建DOCX文档"""
        try:
            doc = docx.Document()
            
            # 设置默认字体为宋体
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加主标题
            title_heading = doc.add_heading(project_name or '生成文档', level=0)
            for run in title_heading.runs:
                run.font.name = '宋体'
                run.font.size = Pt(18)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            def set_font_style(element, font_size=12):
                """设置字体样式"""
                try:
                    if hasattr(element, 'runs'):
                        for run in element.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(font_size)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        if not element.runs:
                            run = element.add_run()
                            run.font.name = '宋体'
                            run.font.size = Pt(font_size)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                except Exception as e:
                    logger.warning(f"设置字体样式失败: {e}")
            
            # 添加内容
            for line in final_text.split('\n'):
                stripped_line = line.strip()
                if stripped_line.startswith('## '):
                    heading_text = stripped_line.replace('## ', '', 1)
                    heading = doc.add_heading(heading_text, level=2)
                    set_font_style(heading, 14)
                elif stripped_line:
                    paragraph = doc.add_paragraph(stripped_line)
                    set_font_style(paragraph, 12)
            
            # 保存文档
            docx_filename = f"task_{task_id}.docx"
            docx_filepath = os.path.join(self.config.TASKS_DIR, docx_filename)
            doc.save(docx_filepath)
            
            logger.info(f"DOCX文档已保存: {docx_filepath}")
            return docx_filepath
            
        except Exception as e:
            logger.error(f"创建DOCX失败: {e}")
            raise
    
    def upload_to_minio(self, file_path: str, object_name: str) -> Optional[str]:
        """上传文件到MinIO"""
        try:
            client = Minio(
                self.config.MINIO_ENDPOINT,
                access_key=self.config.MINIO_ACCESS_KEY,
                secret_key=self.config.MINIO_SECRET_KEY,
                secure=self.config.MINIO_USE_SECURE
            )
            
            bucket_name = self.config.MINIO_BUCKET_NAME
            if not client.bucket_exists(bucket_name):
                client.make_bucket(bucket_name)
                logger.info(f"创建存储桶: {bucket_name}")
            
            client.fput_object(bucket_name, object_name, file_path)
            
            public_url = f"http://{self.config.MINIO_ENDPOINT}/{bucket_name}/{object_name}"
            logger.info(f"文件上传成功: {public_url}")
            return public_url
            
        except Exception as e:
            logger.warning(f"上传到MinIO失败: {e}")
            return None

class AdvancedLongDocumentGenerator:
    """高级长文档生成器"""
    
    def __init__(self, task_id: Optional[str] = None):
        self.task_id = task_id or str(uuid.uuid4())
        self.config = AdvancedLongDocumentConfig()
        self.state = TaskState(self.task_id, self.config)
        self.ai_service = AIService(self.config)
        self.doc_service = DocumentService(self.config)
    
    def generate_document(self, request: str, chathistory: str = "") -> Dict[str, Any]:
        """生成长文档的主要方法"""
        try:
            logger.info(f"开始生成长文档，任务ID: {self.task_id}")
            
            # 初始化任务
            self.state.initialize({
                "chathistory": chathistory,
                "request": request
            })
            
            # 执行生成流程
            self._run_generation_pipeline()
            
            # 返回结果
            return {
                "success": True,
                "task_id": self.task_id,
                "status": self.state.data.get("status"),
                "progress": self.state.data.get("progressPercentage"),
                "message": self.state.data.get("currentStatusMessage"),
                "docx_url": self.state.data.get("docxPublicUrl"),
                "docx_path": self.state.data.get("docxLocalPath"),
                "final_document": self.state.data.get("finalDocument", "")[:500] + "..." if len(self.state.data.get("finalDocument", "")) > 500 else self.state.data.get("finalDocument", "")
            }
            
        except Exception as e:
            logger.error(f"生成文档失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "task_id": self.task_id
            }
    
    def _run_generation_pipeline(self):
        """执行生成流程"""
        try:
            # 阶段1: 准备创作指令
            self._prepare_creative_brief()
            
            # 阶段2: 生成初始大纲
            self._generate_initial_outline()
            
            # 阶段3: 精炼大纲
            self._refine_outline()
            
            # 阶段4: 生成章节内容
            self._generate_chapters()
            
            # 阶段5: 组装最终文档
            self._assemble_final_document()
            
        except Exception as e:
            self.state.log_error("pipeline", str(e))
            raise
    
    def _prepare_creative_brief(self):
        """准备创作指令"""
        self.state.update_status("brief_generation", "正在分析需求并生成创作指令...", 5)
        
        chathistory = self.state.data['initialRequest']['chathistory']
        request = self.state.data['initialRequest']['request']
        
        prompt = f"""你是一个任务规划助理。请根据下面的聊天记录和用户的最终请求，提炼并生成一份详尽的、用于指导后续长文写作的"创作指令"（Creative Brief）。
这份指令应该清晰、结构化，并综合所有已知信息。

【聊天记录】
{chathistory}

【用户最终请求】
{request}

请以JSON格式返回你的分析结果，包含一个'creative_brief'字段，其内容是一段详细的文本指令。
重要提示：所有生成的文本内容都必须使用中文。"""
        
        response = self.ai_service.call_ai_model(prompt, expect_json=True)
        brief = response.get("creative_brief")
        if not brief:
            raise ValueError("AI未能生成有效的创作指令")
        
        self.state.data['creativeBrief'] = brief
        
        # 提取项目名称
        project_prompt = f"""从以下创作指令中，提取一个简短的核心项目名称或主题（例如，"长沙理工大学羽毛球比赛"或"1号住宅楼结构设计"），用于优化后续的知识库检索。
请以JSON格式返回，只包含一个 'project_name' 字段。
重要提示：项目名称必须使用中文。

创作指令：{brief}"""
        
        name_response = self.ai_service.call_ai_model(project_prompt, expect_json=True)
        project_name = name_response.get("project_name", "")
        self.state.data['projectName'] = project_name
        
        self.state.data['status'] = 'brief_prepared'
        self.state.save()
        
        logger.info(f"创作指令已生成，项目名称: {project_name}")
    
    def _generate_initial_outline(self):
        """生成初始大纲"""
        self.state.update_status("outline_generation", "正在生成文档大纲...", 15)
        
        prompt = f"""根据以下创作指令，为一篇长文档生成一份结构化的JSON格式大纲。
JSON的根节点应该有一个名为 'chapters' 的列表。
列表中的每个对象都应包含以下三个字段：
1. 'chapterId' (字符串, 例如 "ch_01")
2. 'title' (字符串, 章节的标题)
3. 'key_points' (字符串列表, 这一章的关键论点)

重要提示：所有生成的标题和文本内容都必须使用中文。

指令：{self.state.data['creativeBrief']}"""
        
        response = self.ai_service.call_ai_model(prompt, expect_json=True)
        chapters = response.get('chapters')
        if chapters is None:
            raise ValueError("AI未能生成有效的大纲")
        
        self.state.data['outline'] = {
            "metadata": {"refinementCycles": 0},
            "chapters": chapters
        }
        self.state.data['status'] = 'outline_generated'
        self.state.save()
        
        logger.info(f"大纲已生成，包含 {len(chapters)} 个章节")
    
    def _refine_outline(self):
        """精炼大纲"""
        self.state.update_status("outline_refinement", "开始大纲精炼...", 25)
        
        project_name = self.state.data.get('projectName', '')
        
        for i in range(self.config.MAX_REFINEMENT_CYCLES):
            self.state.update_status("outline_refinement", f"第 {i + 1} 轮大纲精炼...", 25 + i * 5)
            
            current_outline = json.dumps(self.state.data['outline']['chapters'], ensure_ascii=False)
            
            # AI评审大纲
            critique_prompt = f"""你是一位追求完美的资深编辑。请评审以下大纲。
即使它看起来逻辑完整，也请你主动思考：哪些章节可以通过补充外部的、更具体的数据、案例或细节来变得更具深度和说服力？

请以JSON格式返回你的分析。JSON的根节点应包含一个名为 'gaps_identified' 的列表。
列表中的每个对象都应代表一个需要补充信息的章节，并包含以下字段：
1. 'chapterId' (字符串, 对应原大纲中的章节ID)
2. 'title' (字符串, 对应原大纲中的章节标题)
3. 'query_keywords' (字符串列表, 用于检索外部知识的推荐关键词)

如果这个大纲确实完美无缺，才返回一个空的 'gaps_identified' 列表。
重要提示：所有生成的文本内容都必须使用中文。

大纲：{current_outline}"""
            
            critique_response = self.ai_service.call_ai_model(critique_prompt, expect_json=True)
            gaps = critique_response.get('gaps_identified', [])
            
            if not gaps:
                logger.info("大纲精炼完成，AI认为大纲已足够完善")
                break
            
            # 检索外部知识
            all_knowledge = []
            all_gap_titles = []
            
            for gap_info in gaps:
                gap_title = gap_info.get('title', '')
                if gap_title:
                    all_gap_titles.append(gap_title)
                
                query_keywords = gap_info.get('query_keywords', [])
                for keyword in query_keywords:
                    scoped_query = f"{project_name} {keyword}".strip()
                    knowledge_pieces = self.ai_service.search_vectordata(scoped_query)
                    all_knowledge.extend(knowledge_pieces)
            
            if not all_knowledge:
                logger.info("未检索到有效知识，跳过本轮精炼")
                continue
            
            # 整合知识到大纲
            knowledge_str = "\n\n---\n\n".join(all_knowledge)
            integrate_prompt = f"""请参考以下背景资料，扩充和完善这份大纲，特别是关于'{','.join(all_gap_titles)}'的这些章节。
请返回完整的、更新后的JSON格式大纲。
返回的JSON结构应与原大纲一致（一个包含 'chapters' 列表的根对象）。
重要提示：所有生成的文本内容都必须使用中文。

背景资料：{knowledge_str}

原大纲：{current_outline}"""
            
            updated_response = self.ai_service.call_ai_model(integrate_prompt, expect_json=True)
            updated_chapters = updated_response.get('chapters', self.state.data['outline']['chapters'])
            
            self.state.data['outline']['chapters'] = updated_chapters
            self.state.data['outline']['metadata']['refinementCycles'] = i + 1
            self.state.save()
        
        self.state.data['status'] = 'outline_finalized'
        self.state.update_status("outline_finalized", "大纲已最终确定", 40)
    
    def _generate_chapters(self):
        """生成章节内容"""
        chapters = self.state.data['outline']['chapters']
        project_name = self.state.data.get('projectName', '')
        total_chapters = len(chapters)
        
        if total_chapters == 0:
            logger.warning("大纲为空，无法生成章节内容")
            self.state.data['status'] = 'chapters_generated'
            self.state.save()
            return
        
        for i, chapter in enumerate(chapters):
            progress = 40 + int((i / total_chapters) * 40)
            chapter_title = chapter.get('title', '')
            
            self.state.update_status("content_generation", 
                                   f"正在生成第 {i + 1}/{total_chapters} 章: '{chapter_title}'...", 
                                   progress)
            
            # 检索章节相关知识
            scoped_query = f"{project_name} {chapter_title}".strip()
            knowledge_pieces = self.ai_service.search_vectordata(scoped_query)
            
            # 构建上下文
            context = f"这是全文大纲：{json.dumps(self.state.data['outline']['chapters'], ensure_ascii=False)}\n"
            if i > 0:
                context += f"前一章是关于 '{chapters[i - 1].get('title', '')}' 的。\n"
            
            if knowledge_pieces:
                knowledge_str = "\n\n---\n\n".join(knowledge_pieces)
                context += f"\n在撰写本章时，请重点参考以下背景资料以确保内容的准确性和深度：\n{knowledge_str}\n"
            
            # 生成章节内容
            prompt = f"""请根据以下信息，撰写 '{chapter_title}' 这一章的详细内容。
请专注于阐述这些要点：{', '.join(chapter.get('key_points', []))}。
重要提示：所有你撰写的内容都必须是中文。"""
            
            response = self.ai_service.call_ai_model(prompt, context=context)
            chapter['content'] = response.get('text', '')
            self.state.save()
        
        self.state.data['status'] = 'chapters_generated'
        self.state.save()
        
        logger.info(f"所有 {total_chapters} 个章节内容已生成")
    
    def _assemble_final_document(self):
        """组装最终文档"""
        self.state.update_status("assembling", "正在组装最终文档...", 85)
        
        chapters = self.state.data['outline']['chapters']
        
        # 生成引言
        intro_prompt = "根据以下完整大纲，为这篇文档撰写一段精彩的引言。重要提示：引言必须使用中文。"
        intro_context = json.dumps(chapters, ensure_ascii=False)
        intro_response = self.ai_service.call_ai_model(intro_prompt, intro_context)
        
        # 生成结论
        conclusion_prompt = "根据以下完整大纲和所有章节内容，为这篇文档撰写一段总结性的结论。重要提示：结论必须使用中文。"
        conclusion_response = self.ai_service.call_ai_model(conclusion_prompt, intro_context)
        
        # 组装完整文档
        full_text_parts = [intro_response.get('text', '')]
        
        for chapter in chapters:
            full_text_parts.append(f"\n\n## {chapter.get('title', '')}\n\n")
            full_text_parts.append(chapter.get('content', ''))
        
        full_text_parts.append(f"\n\n## 结论\n\n")
        full_text_parts.append(conclusion_response.get('text', ''))
        
        self.state.data['finalDocument'] = "".join(full_text_parts)
        
        # 创建并上传DOCX文件
        try:
            self.state.update_status("docx_creation", "正在创建DOCX文档...", 90)
            
            project_name = self.state.data.get('projectName', '生成文档')
            docx_path = self.doc_service.create_docx(
                self.state.data['finalDocument'], 
                project_name, 
                self.task_id
            )
            
            self.state.data['docxLocalPath'] = docx_path
            
            # 尝试上传到云存储
            docx_filename = f"task_{self.task_id}.docx"
            public_url = self.doc_service.upload_to_minio(docx_path, docx_filename)
            if public_url:
                self.state.data['docxPublicUrl'] = public_url
            
        except Exception as e:
            logger.warning(f"创建或上传DOCX失败: {e}")
            self.state.log_error("docx_handling", str(e))
        
        self.state.update_status("completed", "文档生成完成！", 100)
        logger.info(f"任务 {self.task_id} 完成")

class AdvancedLongDocumentGeneratorTool(Tool):
    """高级长文档生成工具"""
    
    def __init__(self):
        super().__init__(
            name="advanced_long_document_generator",
            description="🚀 高级长文档生成工具 - 基于AI的专业长篇文档智能生成器！\n\n"
                       "✨ **核心功能：**\n"
                       "- 🧠 AI驱动的创作指令分析\n"
                       "- 📋 智能大纲生成与精炼\n"
                       "- 🔍 向量知识库检索增强\n"
                       "- ✍️ 逐章节内容生成\n"
                       "- 📄 专业DOCX文档输出\n"
                       "- ☁️ 云存储自动上传\n\n"
                       "🎯 **适用场景：**\n"
                       "- 技术报告、研究报告、项目报告\n"
                       "- 施工方案、设计方案、实施方案\n"
                       "- 学术论文、调研分析、可行性研究\n"
                       "- 产品文档、技术规范、操作手册\n\n"
                       "📊 **智能特性：**\n"
                       "- 多轮大纲精炼优化\n"
                       "- 外部知识库智能检索\n"
                       "- 任务状态实时跟踪\n"
                       "- 专业格式文档输出\n\n"
                       "🔧 **使用方法：**\n"
                       "generate_document(request='生成需求描述', chathistory='可选的聊天历史')"
        )
    
    def execute(self, action: str = "generate", request: str = "", chathistory: str = "", **kwargs) -> str:
        """
        执行高级长文档生成
        
        Args:
            action: 操作类型，默认为"generate"
            request: 生成需求描述
            chathistory: 可选的聊天历史上下文
            **kwargs: 其他参数
        
        Returns:
            生成结果的JSON字符串
        """
        try:
            if action == "generate":
                if not request:
                    return json.dumps({
                        "success": False,
                        "error": "请提供生成需求描述（request参数）"
                    }, ensure_ascii=False)
                
                # 创建生成器实例
                generator = AdvancedLongDocumentGenerator()
                
                # 执行生成
                result = generator.generate_document(request, chathistory)
                
                return json.dumps(result, ensure_ascii=False, indent=2)
            
            elif action == "status":
                # 查询任务状态
                task_id = kwargs.get("task_id")
                if not task_id:
                    return json.dumps({
                        "success": False,
                        "error": "请提供任务ID（task_id参数）"
                    }, ensure_ascii=False)
                
                config = AdvancedLongDocumentConfig()
                state = TaskState(task_id, config)
                
                if state.load():
                    return json.dumps({
                        "success": True,
                        "task_id": task_id,
                        "status": state.data.get("status"),
                        "progress": state.data.get("progressPercentage"),
                        "message": state.data.get("currentStatusMessage"),
                        "docx_url": state.data.get("docxPublicUrl"),
                        "docx_path": state.data.get("docxLocalPath")
                    }, ensure_ascii=False)
                else:
                    return json.dumps({
                        "success": False,
                        "error": f"未找到任务: {task_id}"
                    }, ensure_ascii=False)
            
            else:
                return json.dumps({
                    "success": False,
                    "error": f"不支持的操作: {action}。支持的操作: generate, status"
                }, ensure_ascii=False)
                
        except Exception as e:
            logger.error(f"高级长文档生成工具执行失败: {e}")
            return json.dumps({
                "success": False,
                "error": f"执行失败: {str(e)}"
            }, ensure_ascii=False)

# 创建工具实例
def create_advanced_long_document_generator_tool() -> AdvancedLongDocumentGeneratorTool:
    """创建高级长文档生成工具实例"""
    return AdvancedLongDocumentGeneratorTool() 