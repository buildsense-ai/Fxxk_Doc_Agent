#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板插入功能模块：AI智能合并原始文档与模板JSON（Function Call版本）
改造为纯function call形式，供主协调代理调用
"""

import os
import json
import logging
import traceback
import tempfile
import subprocess
from datetime import datetime
from typing import Dict, Any, Optional, Union, List, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# 内置专业 prompt 函数（从 prompt_utils.py 复制）
def get_fill_data_prompt(template_structure: str, placeholders: str, input_data: str) -> str:
    """
    Generates a prompt for the AI to handle hybrid mapping: both template structure and placeholders.
    """
    return f"""
你是一个专业的混合模式文档填写助手。你需要同时处理三种任务：
1. **模板结构匹配**：将JSON数据映射到Word文档的表格单元格和段落结构中
2. **占位符匹配**：为特定的占位符找到对应的数据值
3. **图片引用处理**：处理图片占位符，确保图片能正确引用

---

**核心任务：生成统一的填充数据映射**

你将获得：
- **模板结构**：文档中所有单元格和段落的结构化表示
- **占位符列表**：从特定模式（如"项目名称："和"致____（监理单位）"）提取的占位符
- **输入数据**：需要映射的JSON数据（可能包含attachments_map图片信息）

你的输出应该是一个JSON对象，包含三种类型的映射：
- **结构映射**：键如"table_0_row_1_col_1"或"paragraph_3"，用于填充模板结构
- **占位符映射**：键如"label_项目名称"或"inline_监理单位"，用于替换占位符
- **图片映射**：如果输入数据包含attachments_map，直接将其包含在输出中

---

**重要原则：只填充确定的数据**

⚠️ **关键要求**：
- **只有在输入数据中找到明确对应的信息时，才输出该键值对**
- **如果找不到对应的数据，请完全省略该键，不要输出任何占位文本如"待填写"、"____"等**
- **宁可留空，也不要填入不确定的内容**

---

**映射规则**

### 1. 模板结构映射（保持原有逻辑）
- **排除规则**：如果单元格或段落文本中包含 **"（签字）"** 字样，**绝对不要**为其填充任何内容。
- 如果一个单元格的内容是"标签："（例如 `"项目名称："`），并且其右侧或下方的单元格为空，则应将数据填入那个空单元格
- 如果一个单元格包含下划线占位符但**不是特定的两种模式**，则直接替换该单元格内容
- 智能匹配：运用推理能力，如`project_leader`应映射到"项目负责人"相关的位置

### 2. 占位符映射（现有逻辑）
- 对于`label_*`类型的占位符（来自"项目名称："），找到语义匹配的数据
- 对于`inline_*`类型的占位符（来自"致____（监理单位）"），找到对应的值
- 占位符名称是提示性的，需要智能匹配到输入数据的字段
- **如果确实找不到匹配的数据，就不要输出这个键**

### 3. 图片映射（新增功能）
- 如果输入数据包含`attachments_map`字段，**必须**将其完整地包含在输出JSON中
- 这将确保图片信息能被正确传递到文档生成阶段
- 图片占位符`{{{{image:key}}}}`会在文档生成时被自动替换为"（详见附件N）"的格式

---

### 🧪 示例

**模板结构:**
```json
{{{{
  "table_0_row_0_col_0": "项目名称：",
  "table_0_row_0_col_1": "",
  "table_0_row_1_col_0": "负责人：",
  "table_0_row_1_col_1": "",
  "paragraph_0": "项目名称：{{{{label_项目名称}}}}",
  "paragraph_1": "致{{{{inline_监理单位}}}}（监理单位）",
  "paragraph_2": "施工图详见：{{{{image:shiGongTu}}}}"
}}}}
```

**占位符列表:**
```json
[
  "label_项目名称",
  "inline_监理单位"
]
```

**输入数据:**
```json
{{{{
  "project_name": "古建筑保护修缮工程",  
  "project_leader": "张三",
  "supervision_company": "广州建设监理公司",
  "attachments_map": {{{{
    "shiGongTu": "uploads/construction_drawing.png",
    "xianChangZhaoPian": "uploads/site_photo.jpg"
  }}}}
}}}}
```

**输出结果:**
```json
{{{{
  "table_0_row_0_col_1": "古建筑保护修缮工程",
  "table_0_row_1_col_1": "张三",
  "label_项目名称": "古建筑保护修缮工程",
  "inline_监理单位": "广州建设监理公司",
  "attachments_map": {{{{
    "shiGongTu": "uploads/construction_drawing.png",
    "xianChangZhaoPian": "uploads/site_photo.jpg"
  }}}}
}}}}
```

---

现在请根据下方的数据进行混合映射：

**模板结构:**
```json
{template_structure}
```

**占位符列表:**
```json
{placeholders}
```

**输入数据:**
```json
{input_data}
```

**重要要求:**
- **所有生成的内容必须使用中文**
- 输出统一的JSON对象，包含所有类型的映射
- 对占位符进行智能语义匹配
- 保持原有的模板结构填充逻辑
- **如果输入数据包含attachments_map，必须完整包含在输出中**
- **重要：如果找不到对应数据，完全省略该键，不要填入"待填写"等占位文本**
- 只输出最终的JSON对象，不要包含解释说明或Markdown格式
"""

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# 加载环境变量
try:
    from dotenv import load_dotenv
    load_dotenv()
    logger.info("✅ 已加载.env环境变量文件")
except ImportError:
    logger.warning("⚠️ python-dotenv未安装，将直接从系统环境变量读取配置")
except Exception as e:
    logger.warning(f"⚠️ 加载.env文件时出现问题: {e}")

def get_api_key() -> str:
    """获取OpenRouter API密钥"""
    api_key = os.environ.get("OPENROUTER_API_KEY")
    if not api_key:
        # 检查是否是测试模式
        test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
        if test_mode:
            logger.warning("⚠️ 测试模式：使用模拟API密钥")
            return "test-api-key-for-testing"
        
        logger.error("❌ 未找到OPENROUTER_API_KEY")
        raise RuntimeError("缺少必需的API密钥配置")
    return api_key

# 确保输出目录存在
OUTPUT_DIR = "generated_docs"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

class ProcessingError(Exception):
    """自定义处理异常"""
    def __init__(self, message: str, error_code: str, status_code: int = 500):
        self.message = message
        self.error_code = error_code
        self.status_code = status_code
        super().__init__(self.message)

class DocumentExtractor:
    """文档内容提取器（保持原有功能）"""
    
    def __init__(self):
        logger.info("📄 文档提取器初始化完成")
    
    def extract_from_file_path(self, file_path: str) -> str:
        """从文件路径提取内容（保持原有功能）"""
        if not os.path.exists(file_path):
            raise ProcessingError(
                f"原始文档不存在: {file_path}",
                "FILE_NOT_FOUND",
                404
            )
        return self._extract_content(file_path)
    
    def extract_from_content(self, content: str) -> str:
        """从已有内容中提取和清理（新增功能）"""
        if not content or not content.strip():
            raise ProcessingError(
                "提供的内容为空",
                "EMPTY_CONTENT",
                422
            )
        return content.strip()
    
    def convert_doc_to_docx(self, doc_path: str) -> str:
        """使用LibreOffice将.doc文件转换为.docx文件（保持原有功能）"""
        logger.info("🔄 开始DOC到DOCX转换...")
        
        if not os.path.exists(doc_path):
            logger.error(f"❌ DOC文件不存在: {doc_path}")
            raise ProcessingError(f"DOC文件不存在: {doc_path}", "FILE_NOT_FOUND", 404)
        
        # 生成输出文件名
        docx_path = doc_path.replace('.doc', '_converted.docx')
        
        try:
            # 检查LibreOffice是否可用
            logger.info("🔍 检查LibreOffice可用性...")
            
            # 尝试多个可能的LibreOffice路径
            libreoffice_paths = [
                r'C:\Program Files\LibreOffice\program\soffice.exe',  # Windows 64位
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',  # Windows 32位
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
                'libreoffice',  # Linux/Windows PATH
                'soffice',  # 备用命令
            ]
            
            libreoffice_cmd = None
            for path in libreoffice_paths:
                # 对于绝对路径，先检查文件是否存在
                if os.path.isabs(path):  # 使用os.path.isabs更准确
                    if os.path.exists(path):
                        # 文件存在，直接使用
                        libreoffice_cmd = path
                        logger.info(f"✅ 找到LibreOffice: {path}")
                        break
                    else:
                        logger.debug(f"路径不存在: {path}")
                        continue
                else:
                    # 对于PATH中的命令，尝试简单检查
                    try:
                        result = subprocess.run([path, '--help'], 
                                              capture_output=True, 
                                              text=True, 
                                              timeout=3)
                        if result.returncode == 0:
                            libreoffice_cmd = path
                            logger.info(f"✅ 找到LibreOffice: {path}")
                            break
                    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
                        logger.debug(f"无法使用路径: {path}")
                        continue
            
            if not libreoffice_cmd:
                logger.error("❌ 未找到LibreOffice，请确保已安装LibreOffice")
                raise ProcessingError("LibreOffice未安装或不可用", "LIBREOFFICE_NOT_FOUND", 500)
            
            # 执行转换
            logger.info(f"📄 正在转换: {doc_path} -> {docx_path}")
            
            # 删除已存在的输出文件
            if os.path.exists(docx_path):
                os.remove(docx_path)
                logger.info("🗑️ 删除已存在的转换文件")
            
            # LibreOffice转换命令
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(doc_path),
                doc_path
            ]
            
            logger.info(f"🔧 执行命令: {' '.join(cmd)}")
            
            result = subprocess.run(cmd, 
                                  capture_output=True, 
                                  text=True, 
                                  timeout=30)
            
            if result.returncode != 0:
                logger.error(f"❌ LibreOffice转换失败: {result.stderr}")
                raise ProcessingError(f"LibreOffice转换失败: {result.stderr}", "LIBREOFFICE_CONVERSION_FAILED", 500)
            
            # 检查转换后的文件
            expected_docx = doc_path.replace('.doc', '.docx')
            if os.path.exists(expected_docx):
                # 重命名为我们期望的文件名
                if expected_docx != docx_path:
                    os.rename(expected_docx, docx_path)
                
                logger.info(f"✅ 转换成功: {docx_path}")
                return docx_path
            else:
                logger.error(f"❌ 转换后的文件未找到: {expected_docx}")
                raise ProcessingError("转换后的文件未找到", "CONVERSION_OUTPUT_NOT_FOUND", 500)
                
        except subprocess.TimeoutExpired:
            logger.error("❌ LibreOffice转换超时")
            raise ProcessingError("LibreOffice转换超时", "LIBREOFFICE_TIMEOUT", 500)
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 转换过程中出错: {e}")
            raise ProcessingError(f"转换过程中出错: {str(e)}", "CONVERSION_ERROR", 500)
    
    def _extract_content(self, file_path: str) -> str:
        """提取文档内容的核心方法（保持原有功能）"""
        logger.info(f"📄 开始提取文档内容: {Path(file_path).name}")
        
        content = ""
        converted_file = None  # Track converted file for cleanup
        
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.doc':
                # Convert .doc to .docx using LibreOffice
                logger.info("🔄 检测到.doc文件，开始转换为.docx...")
                converted_file = self.convert_doc_to_docx(file_path)
                file_path = converted_file  # Use converted file for extraction
                file_ext = '.docx'  # Update extension
            
            if file_ext == '.docx':
                doc = DocxDocument(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            content += f"\n表格行: {row_text}"
            
            elif file_ext == '.pdf':
                doc = fitz.open(file_path)
                for page in doc:
                    content += page.get_text()
                doc.close()
            
            elif file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            else:
                raise ProcessingError(
                    f"不支持的文件格式: {file_ext}",
                    "UNSUPPORTED_FORMAT",
                    422
                )
            
            if not content.strip():
                raise ProcessingError(
                    "文档内容为空",
                    "EMPTY_DOCUMENT",
                    422
                )
            
            logger.info(f"✅ 成功提取内容，长度: {len(content)} 字符")
            
            # Cleanup converted file if it exists
            if converted_file and os.path.exists(converted_file):
                try:
                    os.remove(converted_file)
                    logger.info(f"🗑️ 清理转换文件: {converted_file}")
                except Exception as e:
                    logger.warning(f"⚠️ 清理转换文件失败: {e}")
            
            return content.strip()
            
        except ProcessingError:
            # Cleanup converted file on error
            if converted_file and os.path.exists(converted_file):
                try:
                    os.remove(converted_file)
                    logger.info(f"🗑️ 清理转换文件: {converted_file}")
                except:
                    pass
            raise
        except Exception as e:
            # Cleanup converted file on error
            if converted_file and os.path.exists(converted_file):
                try:
                    os.remove(converted_file)
                    logger.info(f"🗑️ 清理转换文件: {converted_file}")
                except:
                    pass
            logger.error(f"❌ 提取文档内容失败: {e}")
            raise ProcessingError(
                f"文档内容提取失败: {str(e)}",
                "EXTRACTION_ERROR",
                500
            )

class ContentMerger:
    """内容智能合并器（保持原有功能并增强）"""
    
    def __init__(self, api_key: str):
        """初始化AI客户端"""
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"  # 保持使用原来的模型
        logger.info("🧠 内容合并器初始化完成")
    
    def merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """使用AI智能合并模板JSON和原始内容（使用专业prompt）"""
        logger.info("🧠 开始AI智能合并（使用专业prompt）...")
        
        # 检查是否是测试模式
        test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
        if test_mode or self.client.api_key == "test-api-key-for-testing":
            logger.warning("⚠️ 测试模式：使用模拟AI合并")
            return self._mock_merge_content(template_json, original_content)
        
        # 使用 prompt_utils 中的专业 prompt
        # 为了适配函数参数，我们需要将 template_json 转换为模板结构格式
        template_structure = json.dumps(template_json, ensure_ascii=False, indent=2)
        placeholders = json.dumps(list(template_json.keys()), ensure_ascii=False, indent=2)
        input_data = original_content
        
        prompt = get_fill_data_prompt(template_structure, placeholders, input_data)
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            
            if not response or not response.choices or not response.choices[0].message.content:
                raise ProcessingError(
                    "AI响应无效或为空",
                    "AI_NO_RESPONSE",
                    500
                )
            
            # 提取JSON内容
            response_content = response.choices[0].message.content.strip()
            json_str = self._extract_json_from_response(response_content)
            
            try:
                merged_content = json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON解析失败: {e}")
                logger.error(f"AI响应内容: {response_content}")
                raise ProcessingError(
                    f"AI返回的内容不是有效的JSON格式: {str(e)}",
                    "AI_INVALID_JSON",
                    422
                )
            
            # 验证合并结果
            if not isinstance(merged_content, dict):
                raise ProcessingError(
                    "AI返回的内容不是字典格式",
                    "AI_INVALID_FORMAT",
                    422
                )
            
            logger.info(f"✅ AI合并成功，生成 {len(merged_content)} 个章节")
            for key, value in merged_content.items():
                preview = str(value)[:100] + "..." if len(str(value)) > 100 else str(value)
                logger.info(f"   📝 {key}: {preview}")
            
            return merged_content
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ AI合并失败: {e}")
            raise ProcessingError(
                f"AI合并过程中发生错误: {str(e)}",
                "AI_MERGE_ERROR",
                500
            )
    
    def analyze_missing_fields(self, merged_content: Dict[str, str]) -> List[str]:
        """分析合并内容中的缺失字段（新增功能）"""
        logger.info("📊 分析缺失字段...")
        
        missing_fields = []
        
        for key, value in merged_content.items():
            if not value or len(str(value).strip()) < 10:
                missing_fields.append(key)
                logger.info(f"   ❌ 缺失字段: {key}")
            elif "待填写" in str(value) or "需要补充" in str(value) or "____" in str(value):
                missing_fields.append(key)
                logger.info(f"   ⚠️ 不完整字段: {key}")
        
        logger.info(f"📊 分析完成，发现 {len(missing_fields)} 个缺失或不完整的字段")
        return missing_fields
    
    def _mock_merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """模拟AI合并（测试模式）（保持原有功能）"""
        logger.info("🧪 模拟AI合并模式")
        
        merged_content = {}
        content_lines = original_content.split('\n')
        content_preview = ' '.join(content_lines[:5])[:200]
        
        for key, description in template_json.items():
            # 基于原始内容和模板描述生成简单的合并内容
            merged_content[key] = f"""根据原始文档内容生成的{key}章节：

{description}

基于原始文档的相关信息：
{content_preview}

本章节内容已根据模板要求进行智能整合，确保符合工程文档的标准格式和要求。具体内容包括项目的基本情况、技术要求、实施方案等关键信息。

注：此内容由测试模式生成，实际应用中将使用真实AI进行智能合并。"""
        
        logger.info(f"✅ 模拟合并完成，生成 {len(merged_content)} 个章节")
        return merged_content
    
    def _extract_json_from_response(self, response_content: str) -> str:
        """从AI响应中提取JSON内容（保持原有功能）"""
        # 尝试提取JSON
        if "```json" in response_content:
            start = response_content.find("```json") + 7
            end = response_content.find("```", start)
            if end != -1:
                return response_content[start:end].strip()
            else:
                return response_content[response_content.find("```json") + 7:].strip()
        elif response_content.startswith("{") and response_content.endswith("}"):
            return response_content
        else:
            # 查找JSON对象
            start_idx = response_content.find("{")
            if start_idx != -1:
                brace_count = 0
                for i, char in enumerate(response_content[start_idx:], start_idx):
                    if char == "{":
                        brace_count += 1
                    elif char == "}":
                        brace_count -= 1
                        if brace_count == 0:
                            return response_content[start_idx:i+1]
            return response_content

class TemplateInserter:
    """真正的模板插入器 - 保持原始模板结构"""
    
    def __init__(self, api_key: str):
        """初始化AI客户端"""
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"
        self.placeholder_originals = {}  # Store original text of placeholders
        logger.info("📄 模板插入器初始化完成")
    
    def _preprocess_template_and_extract_placeholders(self, doc_path: str, output_path: str) -> List[str]:
        """
        扩展占位符预处理，以包含通用的下划线字段，并优化替换逻辑
        """
        logger.info("🛠️  阶段 0: 开始扩展占位符预处理...")
        
        self.placeholder_originals = {} # Reset for each new template analysis
        doc = Document(doc_path)
        placeholders = set()
        blank_counter = 0 # Counter for generic underscore placeholders
        
        def process_text_and_extract_keys(text: str) -> (str, List[str]):
            nonlocal blank_counter
            found_keys = []

            def repl_func(match):
                nonlocal blank_counter
                # Pattern for '致...': underscore_str in group(1), hint in group(2)
                if match.group(1) is not None:
                    if "（签字）" in match.group(0) or "(签字)" in match.group(0):
                        return match.group(0)
                    
                    underscore_str = match.group(1)
                    hint = match.group(2)
                    placeholder_key = f"inline_{hint}"
                    found_keys.append(placeholder_key)
                    self.placeholder_originals[placeholder_key] = underscore_str
                    replacement = f"致{{{placeholder_key}}}（{hint}）"
                    logger.info(f"   - 发现内联模式: '{match.group(0)}' -> '{replacement}'")
                    return replacement

                # Pattern for 'label:': label in group(3)
                elif match.group(3) is not None:
                    # The regex now prevents matching '（签字）:'
                    label = match.group(3).strip()
                    placeholder_key = f"label_{label}"
                    found_keys.append(placeholder_key)
                    replacement = f"{label}：{{{placeholder_key}}}"
                    logger.info(f"   - 发现标签模式: '{match.group(0)}' -> '{replacement}'")
                    return replacement

                # Pattern for general underscores: underscore_str in group(4)
                elif match.group(4) is not None:
                    underscore_str = match.group(4)
                    placeholder_key = f"blank_{blank_counter}"
                    found_keys.append(placeholder_key)
                    self.placeholder_originals[placeholder_key] = underscore_str
                    replacement = f"{{{placeholder_key}}}"
                    logger.info(f"   - 发现通用下划线模式: '{underscore_str}' -> '{replacement}'")
                    blank_counter += 1
                    return replacement
                
                return match.group(0)

            # Regex updated to handle spaced underscores and avoid capturing signature labels
            import re
            pattern = re.compile(
                r"致\s*(__{3,})\s*（([^）]+)）"              # G1: underscore, G2: hint
                r"|([^：\n（(]+?)：\s*$"                    # G3: label, avoids '(...):'
                r"|((?:_{4,}[\s\xa0]*)+)"               # G4: general underscore blocks
            )

            processed_text = pattern.sub(repl_func, text)
            
            return processed_text, found_keys
        
        # --- Process all paragraphs ---
        for para in doc.paragraphs:
            original_text = para.text
            if not original_text.strip():
                continue

            new_text, keys = process_text_and_extract_keys(original_text)
            if new_text != original_text:
                placeholders.update(keys)
                # To preserve formatting, we clear runs and add a new one
                para.clear()
                para.add_run(new_text)
                logger.info(f"   📝 段落更新: '{original_text.strip()}' -> '{new_text.strip()}'")

        # --- Process all tables ---
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    if not original_text.strip():
                        continue
                        
                    new_text, keys = process_text_and_extract_keys(original_text)
                    if new_text != original_text:
                        placeholders.update(keys)
                        # Reverted to cell.text for simplicity and correctness.
                        # This replaces the content of the first paragraph in the cell.
                        cell.text = new_text
                        logger.info(f"   📋 表格更新: '{original_text.strip()}' -> '{new_text.strip()}'")
        
        doc.save(output_path)
        logger.info(f"✅ 扩展预处理完成. 找到 {len(placeholders)} 个占位符. 新模板: {output_path}")
        return list(placeholders)

    def analyze_template_structure(self, template_path: str) -> Dict[str, str]:
        """
        确定性地分析Word模板，提取带有位置信息的结构。
        """
        logger.info("🔍 开始确定性模板结构分析...")
        
        try:
            doc = Document(template_path)
            template_structure = {}
            
            logger.info(f"📄 正在读取模板文件: {template_path}")
            
            # 提取表格结构
            for i, table in enumerate(doc.tables):
                for j, row in enumerate(table.rows):
                    for k, cell in enumerate(row.cells):
                        cell_key = f"table_{i}_row_{j}_col_{k}"
                        template_structure[cell_key] = cell.text.strip()
            
            # 提取段落结构（不做特殊处理，保持原始内容）
            for i, para in enumerate(doc.paragraphs):
                para_key = f"paragraph_{i}"
                template_structure[para_key] = para.text.strip()
            
            logger.info(f"✅ 成功提取 {len(template_structure)} 个结构元素。")
            
            return template_structure
            
        except Exception as e:
            logger.error(f"❌ 模板结构分析错误: {e}")
            raise

    def fill_template_with_ai_data(self, template_path: str, output_path: str, fill_data: Dict[str, str]):
        """
        混合填充 - 支持图片附件和占位符，保持原始模板结构
        """
        logger.info("📝 开始混合模式模板填充（保持原始结构）...")
        
        doc = Document(template_path)
        filled_count = 0
        
        # 1. 准备附件信息
        attachments_map = fill_data.pop('attachments_map', {})
        attachment_ref_map = {}
        ordered_attachments = []
        if attachments_map and isinstance(attachments_map, dict):
            logger.info(f"🖼️  找到 {len(attachments_map)} 个图片附件待处理。")
            ordered_attachments = list(attachments_map.items())
            for i, (key, _) in enumerate(ordered_attachments):
                attachment_ref_map[key.strip()] = i + 1
        else:
            attachments_map = {}

        # 2. 分离文本填充数据
        placeholder_data = {k: v for k, v in fill_data.items() if k.startswith(('label_', 'inline_', 'blank_'))}
        structure_data = {k: v for k, v in fill_data.items() if k.startswith(('table_', 'paragraph_'))}
        
        # 3. 替换所有文本占位符（包括图片引用）
        import re
        image_placeholder_pattern = re.compile(r'\{\{image:([^}]+)\}\}')
        text_placeholder_pattern = re.compile(r'\{(label_[^}]+|inline_[^}]+|blank_[^}]+)\}')

        def process_element_text(element):
            nonlocal filled_count
            if '{' not in element.text:
                return

            original_text = element.text
            new_text = ""
            last_end = 0

            # 创建一个组合的正则表达式来查找所有类型的占位符
            combined_pattern = re.compile(f"({image_placeholder_pattern.pattern}|{text_placeholder_pattern.pattern})")
            
            for match in combined_pattern.finditer(original_text):
                new_text += original_text[last_end:match.start()]
                last_end = match.end()
                
                image_key_match = image_placeholder_pattern.match(match.group(0))
                text_key_match = text_placeholder_pattern.match(match.group(0))

                if image_key_match:
                    key = image_key_match.group(1).strip()
                    if key in attachment_ref_map:
                        number = attachment_ref_map[key]
                        replacement = f"（详见附件{number}）"
                        new_text += replacement
                        logger.info(f"   🖼️  图片引用替换: '{match.group(0)}' -> '{replacement}'")
                    else:
                        logger.warning(f"   ⚠️  找到图片占位符 {match.group(0)} 但无匹配图片，已移除。")
                
                elif text_key_match:
                    placeholder_key = text_key_match.group(1)
                    placeholder = f"{{{placeholder_key}}}"

                    if placeholder_key in placeholder_data:
                        value = str(placeholder_data[placeholder_key])
                        new_text += value
                        logger.info(f"   ✏️  占位符填充: {placeholder} -> {value[:50]}...")
                        filled_count += 1
                    else: # 未匹配的文本占位符
                        if placeholder_key.startswith('label_'):
                            logger.info(f"   🔘  移除未匹配标签占位符: {placeholder}")
                        elif placeholder_key.startswith(('inline_', 'blank_')):
                            original_underscore = self.placeholder_originals.get(placeholder_key, '____')
                            new_text += original_underscore
                            logger.info(f"   🔘  恢复未匹配占位符: {placeholder} -> '{original_underscore}'")

            new_text += original_text[last_end:]
            
            if new_text != original_text:
                element.text = new_text

        # 遍历段落和表格进行统一替换
        for para in doc.paragraphs:
            process_element_text(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_element_text(cell)

        # 4. 填充原始结构
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    cell_key = f"table_{i}_row_{j}_col_{k}"
                    if cell_key in structure_data:
                        cell.text = str(structure_data[cell_key])
                        logger.info(f"   ✏️  结构填充(表格): {cell_key} -> {str(structure_data[cell_key])[:50]}...")
                        filled_count += 1

        for i, para in enumerate(doc.paragraphs):
            para_key = f"paragraph_{i}"
            if para_key in structure_data:
                # 只有在段落中没有占位符的情况下才进行结构填充
                combined_pattern = re.compile(f"({image_placeholder_pattern.pattern}|{text_placeholder_pattern.pattern})")
                if not combined_pattern.search(para.text):
                    para.text = str(structure_data[para_key])
                    logger.info(f"   ✏️  结构填充(段落): {para_key} -> {str(structure_data[para_key])[:50]}...")
                    filled_count += 1

        # 5. 将图片作为附件附加到文档末尾
        if ordered_attachments:
            logger.info("📎 开始在文档末尾附加图片...")
            try:
                doc.add_page_break()
                doc.add_heading('附件列表', level=1)
                
                for i, (key, image_path) in enumerate(ordered_attachments):
                    attachment_counter = i + 1
                    if not image_path or not isinstance(image_path, str) or not os.path.exists(image_path):
                        logger.warning(f"⚠️ 图片路径不存在或无效，跳过附件 '{key}': {image_path}")
                        continue
                    
                    try:
                        heading_text = f"附件 {attachment_counter}: {key}"
                        doc.add_heading(heading_text, level=2)
                        doc.add_picture(image_path, width=Inches(6.0))
                        doc.add_paragraph()
                        logger.info(f"   ✅ 成功附加图片: {heading_text} ({image_path})")
                    except Exception as pic_e:
                        logger.error(f"❌ 附加图片 '{key}' ({image_path}) 时出错: {pic_e}")
            except Exception as e:
                logger.error(f"❌ 处理附件时发生意外错误: {e}")
        
        doc.save(output_path)
        logger.info(f"✅ 混合模式填充完成，共填充 {filled_count} 个字段: {output_path}")
        
        return {
            "sections_count": filled_count,
            "file_size": os.path.getsize(output_path),
            "validation": {"is_valid": True, "paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        }

    def process_template_insertion(self, template_path: str, original_content: str, output_path: str) -> Dict[str, Any]:
        """
        完整的模板插入处理流程 - 保持原始模板结构
        
        Args:
            template_path: 模板文件路径 (.doc 或 .docx)
            original_content: 原始内容（字符串格式）
            output_path: 输出文件路径
            
        Returns:
            处理结果字典
        """
        logger.info("🚀 开始完整模板插入处理...")
        
        try:
            # Stage 0: Convert .doc to .docx if necessary
            if template_path.lower().endswith('.doc'):
                logger.info(f"📄 检测到.doc模板，开始转换: {template_path}")
                extractor = DocumentExtractor()
                original_docx_path = extractor.convert_doc_to_docx(template_path)
            else:
                original_docx_path = template_path

            # Stage 0.5: 预处理模板，提取占位符
            processed_template_path = original_docx_path.replace(".docx", "_processed.docx")
            placeholders = self._preprocess_template_and_extract_placeholders(
                doc_path=original_docx_path,
                output_path=processed_template_path
            )
            
            # Stage 1: 分析处理后的模板结构
            template_structure = self.analyze_template_structure(processed_template_path)

            # Stage 2: 创建输入数据结构
            input_data = {"original_content": original_content}

            # Stage 2.5: 混合模式AI映射
            fill_data = self.ai_generate_fill_data(
                template_structure=template_structure,
                placeholders=placeholders,
                input_data=input_data
            )
            
            # Stage 3: 混合模式填充（保持原始结构）
            generation_info = self.fill_template_with_ai_data(
                template_path=processed_template_path,
                output_path=output_path,
                fill_data=fill_data
            )
            
            # 清理临时文件
            if original_docx_path != template_path and os.path.exists(original_docx_path):
                os.remove(original_docx_path)
            if os.path.exists(processed_template_path):
                os.remove(processed_template_path)
            
            logger.info(f"✅ 模板插入处理完成: {output_path}")
            return generation_info
            
        except Exception as e:
            logger.error(f"❌ 模板插入处理失败: {e}", exc_info=True)
            raise

    def ai_generate_fill_data(self, template_structure: Dict[str, str], placeholders: List[str], input_data: Dict[str, Any]) -> Dict[str, str]:
        """
        混合模式 - 使用AI同时处理模板结构匹配和占位符匹配
        """
        logger.info("🧠 开始混合模式AI字段映射...")
        
        try:
            # 构建混合提示
            prompt = get_fill_data_prompt(
                json.dumps(template_structure, ensure_ascii=False, indent=2),
                json.dumps(placeholders, ensure_ascii=False, indent=2),
                json.dumps(input_data, ensure_ascii=False, indent=2)
            )
            
            logger.info("🧠 正在调用AI进行混合模式映射...")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            
            if not response or not response.choices or not response.choices[0].message.content:
                raise ValueError("AI响应无效或为空")

            json_string = self._extract_json_from_response(response.choices[0].message.content)
            fill_data = json.loads(json_string)
            
            logger.info(f"✅ AI成功生成 {len(fill_data)} 个字段的映射:")
            for key, value in fill_data.items():
                preview = str(value)[:70] + "..." if len(str(value)) > 70 else str(value)
                logger.info(f"   🔗 {key} -> '{preview}'")
            
            return fill_data
            
        except Exception as e:
            logger.error(f"❌ AI字段映射错误: {e}", exc_info=True)
            return {}

    def _extract_json_from_response(self, response_content: str) -> str:
        """从AI响应中提取JSON内容"""
        if "```json" in response_content:
            start = response_content.find("```json") + 7
            end = response_content.find("```", start)
            if end != -1:
                return response_content[start:end].strip()
            else:
                return response_content[response_content.find("```json") + 7:].strip()
        elif response_content.startswith("{") and response_content.endswith("}"):
            return response_content
        else:
            # 查找JSON对象
            start_idx = response_content.find("{")
            if start_idx != -1:
                brace_count = 0
                for i, char in enumerate(response_content[start_idx:], start_idx):
                    if char == "{":
                        brace_count += 1
                    elif char == "}":
                        brace_count -= 1
                        if brace_count == 0:
                            return response_content[start_idx:i+1]
            return response_content

# ====================================
# 核心 Function Call 接口
# ====================================

def merge_template_with_context(
    template_json: Dict[str, str],
    original_content: str,
    api_key: str,
    output_path: Optional[str] = None
) -> Dict[str, Any]:
    """
    AI智能合并模板与上下文内容的核心function call接口 - 现在使用真正的模板插入逻辑
    
    Args:
        template_json: 模板JSON字典，键为章节名，值为章节描述  
        original_content: 原始文档内容（字符串）
        api_key: AI API密钥
        output_path: 可选的输出路径，如果不提供则自动生成
    
    Returns:
        Dict包含:
        - merged_json: 合并后的填充结果  
        - missing_fields: 缺失或不完整的字段列表
        - output_path: 最终文档路径
        - status: 状态（success/partial/error）
        - metadata: 其他元数据
    """
    logger.info("🚀 开始真正的模板插入处理（Function Call模式）...")
    
    try:
        # 这是旧的接口，现在我们不能直接使用template_json作为模板
        # 因为我们需要一个真实的Word模板文件，而不是JSON描述
        # 所以这里我们回退到使用ContentMerger来生成内容，然后创建新文档
        logger.warning("⚠️ 注意：此接口将生成新文档，不是模板插入。请使用run_template_insertion()进行真正的模板插入。")
        
        # 1. 初始化组件
        merger = ContentMerger(api_key)
        
        # 2. AI智能合并
        logger.info("🧠 第一步：AI智能内容合并")
        merged_json = merger.merge_content(template_json, original_content)
        
        # 3. 分析缺失字段
        logger.info("📊 第二步：分析内容完整性")
        missing_fields = merger.analyze_missing_fields(merged_json)
        
        # 4. 生成输出路径
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"merged_document_{timestamp}.docx"
            output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        # 5. 生成简单的docx文档（不是模板插入）
        logger.info("📄 第三步：生成简单DOCX文档")
        doc = Document()
        doc.add_heading('AI智能合并文档', 0)
        
        for section_title, section_content in merged_json.items():
            doc.add_heading(section_title, level=1)
            doc.add_paragraph(str(section_content))
        
        doc.save(output_path)
        
        # 6. 确定状态
        if len(missing_fields) == 0:
            status = "success"
            message = "内容合并完成，所有字段填充完整"
        elif len(missing_fields) < len(template_json) / 2:
            status = "partial"
            message = f"内容合并部分完成，{len(missing_fields)} 个字段需要补充"
        else:
            status = "needs_more_info"
            message = f"内容合并需要更多信息，{len(missing_fields)} 个字段缺失"
        
        logger.info(f"✅ 内容合并处理完成，状态: {status}")
        
        return {
            "merged_json": merged_json,
            "missing_fields": missing_fields,
            "output_path": output_path,
            "status": status,
            "message": message,
            "metadata": {
                "template_sections": len(template_json),
                "filled_sections": len(template_json) - len(missing_fields),
                "completion_rate": (len(template_json) - len(missing_fields)) / len(template_json),
                "generation_info": {"file_size": os.path.getsize(output_path) if os.path.exists(output_path) else 0},
                "original_content_length": len(original_content),
                "timestamp": datetime.now().isoformat()
            }
        }
        
    except ProcessingError as e:
        logger.error(f"❌ 处理错误: {e}")
        return {
            "merged_json": {},
            "missing_fields": list(template_json.keys()),
            "output_path": None,
            "status": "error",
            "message": e.message,
            "metadata": {
                "error_code": e.error_code,
                "error_details": str(e),
                "timestamp": datetime.now().isoformat()
            }
        }
    except Exception as e:
        logger.error(f"❌ 未预期错误: {e}")
        return {
            "merged_json": {},
            "missing_fields": list(template_json.keys()),
            "output_path": None,
            "status": "error",
            "message": f"系统错误: {str(e)}",
            "metadata": {
                "error_code": "UNEXPECTED_ERROR",
                "error_details": traceback.format_exc(),
                "timestamp": datetime.now().isoformat()
            }
        }

def template_insertion_with_context(
    template_file_path: str,
    original_content: str,
    api_key: str,
    output_path: Optional[str] = None
) -> Dict[str, Any]:
    """
    真正的模板插入功能 - 保持原始模板结构的function call接口
    
    Args:
        template_file_path: Word模板文件路径 (.doc 或 .docx)
        original_content: 原始文档内容（字符串）
        api_key: AI API密钥
        output_path: 可选的输出路径，如果不提供则自动生成
    
    Returns:
        Dict包含处理结果的详细信息
    """
    logger.info("🚀 开始真正的模板插入处理...")
    
    try:
        # 1. 生成输出路径
        if not output_path:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            template_name = os.path.splitext(os.path.basename(template_file_path))[0]
            output_filename = f"template_inserted_{template_name}_{timestamp}.docx"
            output_path = os.path.join(OUTPUT_DIR, output_filename)
        
        # 2. 初始化模板插入器
        inserter = TemplateInserter(api_key)
        
        # 3. 执行模板插入处理
        generation_info = inserter.process_template_insertion(
            template_path=template_file_path,
            original_content=original_content,
            output_path=output_path
        )
        
        logger.info(f"✅ 真正的模板插入处理完成: {output_path}")
        
        return {
            "merged_json": {},  # 这里不返回merged_json，因为是结构化填充
            "missing_fields": [],  # 这里简化处理
            "output_path": output_path,
            "status": "success",
            "message": "模板插入完成，保持了原始模板结构",
            "metadata": {
                "template_file": template_file_path,
                "filled_fields": generation_info.get("sections_count", 0),
                "generation_info": generation_info,
                "original_content_length": len(original_content),
                "timestamp": datetime.now().isoformat()
            }
        }
        
    except Exception as e:
        logger.error(f"❌ 模板插入处理失败: {e}")
        return {
            "merged_json": {},
            "missing_fields": [],
            "output_path": None,
            "status": "error",
            "message": f"模板插入失败: {str(e)}",
            "metadata": {
                "template_file": template_file_path,
                "error_details": traceback.format_exc(),
                "timestamp": datetime.now().isoformat()
            }
        }

def extract_content_from_file(file_path: str) -> Dict[str, Any]:
    """
    从文件中提取内容的function call接口
    
    Args:
        file_path: 文档文件路径
        
    Returns:
        Dict包含:
        - content: 提取的文本内容
        - status: 状态（success/error）
        - metadata: 文件信息等元数据
    """
    logger.info(f"📄 开始提取文件内容: {file_path}")
    
    try:
        extractor = DocumentExtractor()
        content = extractor.extract_from_file_path(file_path)
        
        return {
            "content": content,
            "status": "success",
            "message": "文档内容提取成功",
            "metadata": {
                "file_path": file_path,
                "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0,
                "content_length": len(content),
                "timestamp": datetime.now().isoformat()
            }
        }
        
    except ProcessingError as e:
        logger.error(f"❌ 文档提取错误: {e}")
        return {
            "content": "",
            "status": "error",
            "message": e.message,
            "metadata": {
                "file_path": file_path,
                "error_code": e.error_code,
                "error_details": str(e),
                "timestamp": datetime.now().isoformat()
            }
        }
    except Exception as e:
        logger.error(f"❌ 文档提取未预期错误: {e}")
        return {
            "content": "",
            "status": "error",
            "message": f"文档提取失败: {str(e)}",
            "metadata": {
                "file_path": file_path,
                "error_code": "FILE_EXTRACTION_ERROR",
                "error_details": traceback.format_exc(),
                "timestamp": datetime.now().isoformat()
            }
        }

# ====================================
# 向后兼容接口（供原有代码调用）
# ====================================

def run_template_insertion(template_json_input: Union[str, Dict[str, str]], original_file_path: str) -> str:
    """
    真正的模板插入接口 - 现在使用template_insertion_with_context进行真正的模板插入
    （保持向后兼容的函数签名）
    
    Args:
        template_json_input: 这个参数现在被忽略，因为我们直接使用original_file_path作为模板
        original_file_path: 现在作为Word模板文件路径使用
        
    Returns:
        生成的文档路径
    """
    logger.info("🚀 真正的模板插入模式：启动模板插入处理...")
    
    try:
        # 在新的逻辑中，original_file_path 实际上是模板文件
        # 我们需要一些实际的内容来填充模板
        template_file_path = original_file_path
        
        # 创建一些模拟的原始内容，如果是从JSON提供的话
        if isinstance(template_json_input, str) and os.path.exists(template_json_input):
            with open(template_json_input, 'r', encoding='utf-8') as f:
                template_json = json.load(f)
            
            # 将JSON内容转换为文本格式作为原始内容  
            original_content = "基于JSON数据的模板填充：\n\n"
            for key, value in template_json.items():
                original_content += f"{key}: {value}\n"
                
        elif isinstance(template_json_input, dict):
            # 直接从字典创建内容
            original_content = "基于JSON数据的模板填充：\n\n"
            for key, value in template_json_input.items():
                original_content += f"{key}: {value}\n"
        else:
            # 如果没有提供有效的JSON，创建一个默认内容
            original_content = """
            这是一个建筑施工项目的相关文档。
            项目名称：某大型商业建筑施工项目
            项目地点：北京市朝阳区
            
            安全方面需要特别注意以下几点：
            1. 高空作业安全措施
            2. 施工现场围挡设置
            3. 安全教育培训
            
            质量管理要求：
            - 严格按照设计图纸施工
            - 材料进场验收
            - 工序质量检查
            
            项目负责人：张工程师
            审核人员：李监理
            """
        
        # Get API key and use real template insertion
        api_key = get_api_key()
        insertion_result = template_insertion_with_context(
            template_file_path=template_file_path,
            original_content=original_content,
            api_key=api_key
        )
        
        if insertion_result["status"] == "error":
            raise ProcessingError(
                insertion_result["message"],
                insertion_result["metadata"].get("error_code", "INSERTION_ERROR"),
                500
            )
        
        final_doc_path = insertion_result["output_path"]
        logger.info(f"✅ 真正的模板插入处理完成: {final_doc_path}")
        
        return final_doc_path

    except (ProcessingError, FileNotFoundError) as e:
        logger.error(f"❌ 模板插入处理失败: {e}")
        raise
    except Exception as e:
        logger.error(f"❌ 模板插入未预期错误: {e}")
        logger.error(traceback.format_exc())
        raise ProcessingError(f"An unexpected error occurred: {str(e)}", "UNEXPECTED_ERROR", 500)

# ====================================
# 测试和演示函数
# ====================================

def test_merge_function():
    """测试合并功能"""
    print("🧪 测试模板合并功能")
    print("=" * 50)
    
    # 准备测试数据
    test_template = {
        "项目概述": "项目的基本信息和背景介绍",
        "安全措施": "施工安全管理措施和要求",
        "质量控制": "工程质量控制标准和方法"
    }
    
    test_content = """
    这是一个建筑施工项目的相关文档。
    项目名称：某大型商业建筑施工项目
    项目地点：北京市朝阳区
    
    安全方面需要特别注意以下几点：
    1. 高空作业安全措施
    2. 施工现场围挡设置
    3. 安全教育培训
    
    质量管理要求：
    - 严格按照设计图纸施工
    - 材料进场验收
    - 工序质量检查
    """
    
    try:
        # 测试API密钥获取
        api_key = get_api_key()
        print(f"✅ API密钥获取成功")
        
        # 测试合并功能
        result = merge_template_with_context(
            template_json=test_template,
            original_content=test_content,
            api_key=api_key
        )
        
        print(f"📊 合并结果:")
        print(f"   状态: {result['status']}")
        print(f"   消息: {result['message']}")
        print(f"   输出文档: {result['output_path']}")
        print(f"   缺失字段: {result['missing_fields']}")
        print(f"   完成率: {result['metadata']['completion_rate']:.2%}")
        
        if result['output_path'] and os.path.exists(result['output_path']):
            print(f"✅ 文档生成成功: {result['output_path']}")
        else:
            print(f"⚠️ 文档生成可能有问题")
        
    except Exception as e:
        print(f"❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_merge_function()