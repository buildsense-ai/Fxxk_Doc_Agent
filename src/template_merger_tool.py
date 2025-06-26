#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板合并工具：AI智能合并原始文档与模板，支持模板插入功能
"""

import os
import json
import logging
import traceback
import tempfile
import subprocess
from datetime import datetime
from typing import Dict, Any, Optional, Union, List, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# 加载环境变量
try:
    from dotenv import load_dotenv
    load_dotenv()
    logger.info("✅ 已加载.env环境变量文件")
except ImportError:
    logger.warning("⚠️ python-dotenv未安装，将直接从系统环境变量读取配置")
except Exception as e:
    logger.warning(f"⚠️ 加载.env文件时出现问题: {e}")

def get_fill_data_prompt(template_structure: str, placeholders: str, input_data: str) -> str:
    """
    Generates a prompt for the AI to handle hybrid mapping: both template structure and placeholders.
    """
    return f"""
你是一个专业的混合模式文档填写助手。你需要同时处理三种任务：
1. **模板结构匹配**：将JSON数据映射到Word文档的表格单元格和段落结构中
2. **占位符匹配**：为特定的占位符找到对应的数据值
3. **图片引用处理**：处理图片占位符，确保图片能正确引用

---

**核心任务：生成统一的填充数据映射**

你将获得：
- **模板结构**：文档中所有单元格和段落的结构化表示
- **占位符列表**：从特定模式（如"项目名称："和"致____（监理单位）"）提取的占位符
- **输入数据**：需要映射的JSON数据（可能包含attachments_map图片信息）

你的输出应该是一个JSON对象，包含三种类型的映射：
- **结构映射**：键如"table_0_row_1_col_1"或"paragraph_3"，用于填充模板结构
- **占位符映射**：键如"label_项目名称"或"inline_监理单位"，用于替换占位符
- **图片映射**：如果输入数据包含attachments_map，直接将其包含在输出中

---

**重要原则：只填充确定的数据**

⚠️ **关键要求**：
- **只有在输入数据中找到明确对应的信息时，才输出该键值对**
- **如果找不到对应的数据，请完全省略该键，不要输出任何占位文本如"待填写"、"____"等**
- **宁可留空，也不要填入不确定的内容**

现在请根据下方的数据进行混合映射：

**模板结构:**
```json
{template_structure}
```

**占位符列表:**
```json
{placeholders}
```

**输入数据:**
```json
{input_data}
```

**重要要求:**
- **所有生成的内容必须使用中文**
- 输出统一的JSON对象，包含所有类型的映射
- 对占位符进行智能语义匹配
- 保持原有的模板结构填充逻辑
- **如果输入数据包含attachments_map，必须完整包含在输出中**
- **重要：如果找不到对应数据，完全省略该键，不要填入"待填写"等占位文本**
- 只输出最终的JSON对象，不要包含解释说明或Markdown格式
"""

class ProcessingError(Exception):
    """自定义处理异常"""
    def __init__(self, message: str, error_code: str, status_code: int = 500):
        self.message = message
        self.error_code = error_code
        self.status_code = status_code
        super().__init__(self.message)

class DocumentExtractor:
    """文档内容提取器"""
    
    def __init__(self):
        logger.info("📄 文档提取器初始化完成")
    
    def extract_from_file_path(self, file_path: str) -> str:
        """从文件路径提取内容"""
        if not os.path.exists(file_path):
            raise ProcessingError(
                f"原始文档不存在: {file_path}",
                "FILE_NOT_FOUND",
                404
            )
        return self._extract_content(file_path)
    
    def convert_doc_to_docx(self, doc_path: str) -> str:
        """使用LibreOffice将.doc文件转换为.docx文件"""
        logger.info("🔄 开始DOC到DOCX转换...")
        
        if not os.path.exists(doc_path):
            logger.error(f"❌ DOC文件不存在: {doc_path}")
            raise ProcessingError(f"DOC文件不存在: {doc_path}", "FILE_NOT_FOUND", 404)
        
        # 生成输出文件名
        docx_path = doc_path.replace('.doc', '_converted.docx')
        
        try:
            # 检查LibreOffice是否可用
            logger.info("🔍 检查LibreOffice可用性...")
            
            # 尝试多个可能的LibreOffice路径
            libreoffice_paths = [
                r'C:\Program Files\LibreOffice\program\soffice.exe',  # Windows 64位
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',  # Windows 32位
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
                'libreoffice',  # Linux/Windows PATH
                'soffice',  # 备用命令
            ]
            
            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run([path, '--version'], 
                                          capture_output=True, 
                                          text=True, 
                                          timeout=10)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        logger.info(f"✅ 找到LibreOffice: {path}")
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            if not libreoffice_cmd:
                logger.error("❌ 未找到LibreOffice，请确保已安装LibreOffice")
                raise ProcessingError("LibreOffice未安装或不可用", "LIBREOFFICE_NOT_FOUND", 500)
            
            # 执行转换
            logger.info(f"📄 正在转换: {doc_path} -> {docx_path}")
            
            # 删除已存在的输出文件
            if os.path.exists(docx_path):
                os.remove(docx_path)
                logger.info("🗑️ 删除已存在的转换文件")
            
            # LibreOffice转换命令
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(doc_path),
                doc_path
            ]
            
            logger.info(f"🔧 执行命令: {' '.join(cmd)}")
            
            result = subprocess.run(cmd, 
                                  capture_output=True, 
                                  text=True, 
                                  timeout=30)
            
            if result.returncode != 0:
                logger.error(f"❌ LibreOffice转换失败: {result.stderr}")
                raise ProcessingError(f"LibreOffice转换失败: {result.stderr}", "LIBREOFFICE_CONVERSION_FAILED", 500)
            
            # 检查转换后的文件
            expected_docx = doc_path.replace('.doc', '.docx')
            if os.path.exists(expected_docx):
                # 重命名为我们期望的文件名
                if expected_docx != docx_path:
                    os.rename(expected_docx, docx_path)
                
                logger.info(f"✅ 转换成功: {docx_path}")
                return docx_path
            else:
                logger.error(f"❌ 转换后的文件未找到: {expected_docx}")
                raise ProcessingError("转换后的文件未找到", "CONVERSION_OUTPUT_NOT_FOUND", 500)
                
        except subprocess.TimeoutExpired:
            logger.error("❌ LibreOffice转换超时")
            raise ProcessingError("LibreOffice转换超时", "LIBREOFFICE_TIMEOUT", 500)
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ 转换过程中出错: {e}")
            raise ProcessingError(f"转换过程中出错: {str(e)}", "CONVERSION_ERROR", 500)
    
    def _extract_content(self, file_path: str) -> str:
        """提取文档内容的核心方法"""
        logger.info(f"📄 开始提取文档内容: {Path(file_path).name}")
        
        content = ""
        converted_file = None
        
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext == '.doc':
                # Convert .doc to .docx using LibreOffice
                logger.info("🔄 检测到.doc文件，开始转换为.docx...")
                converted_file = self.convert_doc_to_docx(file_path)
                file_path = converted_file
                file_ext = '.docx'
            
            if file_ext == '.docx':
                doc = DocxDocument(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            content += f"\n表格行: {row_text}"
            
            elif file_ext == '.pdf':
                doc = fitz.open(file_path)
                for page in doc:
                    content += page.get_text()
                doc.close()
            
            elif file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            else:
                raise ProcessingError(
                    f"不支持的文件格式: {file_ext}",
                    "UNSUPPORTED_FORMAT",
                    422
                )
            
            if not content.strip():
                raise ProcessingError(
                    "文档内容为空",
                    "EMPTY_DOCUMENT",
                    422
                )
            
            logger.info(f"✅ 成功提取内容，长度: {len(content)} 字符")
            
            # Cleanup converted file if it exists
            if converted_file and os.path.exists(converted_file):
                try:
                    os.remove(converted_file)
                    logger.info(f"🗑️ 清理转换文件: {converted_file}")
                except Exception as e:
                    logger.warning(f"⚠️ 清理转换文件失败: {e}")
            
            return content.strip()
            
        except ProcessingError:
            # Cleanup converted file on error
            if converted_file and os.path.exists(converted_file):
                try:
                    os.remove(converted_file)
                except:
                    pass
            raise
        except Exception as e:
            # Cleanup converted file on error
            if converted_file and os.path.exists(converted_file):
                try:
                    os.remove(converted_file)
                except:
                    pass
            logger.error(f"❌ 提取文档内容失败: {e}")
            raise ProcessingError(
                f"文档内容提取失败: {str(e)}",
                "EXTRACTION_ERROR",
                500
            )

class ContentMerger:
    """内容智能合并器"""
    
    def __init__(self, api_key: str):
        """初始化AI客户端"""
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"
        logger.info("🧠 内容合并器初始化完成")
    
    def merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """使用AI智能合并模板JSON和原始内容"""
        logger.info("🧠 开始AI智能合并...")
        
        # 检查是否是测试模式
        test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
        if test_mode or self.client.api_key == "test-api-key-for-testing":
            logger.warning("⚠️ 测试模式：使用模拟AI合并")
            return self._mock_merge_content(template_json, original_content)
        
        # 使用专业 prompt
        template_structure = json.dumps(template_json, ensure_ascii=False, indent=2)
        placeholders = json.dumps(list(template_json.keys()), ensure_ascii=False, indent=2)
        input_data = original_content
        
        prompt = get_fill_data_prompt(template_structure, placeholders, input_data)
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            
            if not response or not response.choices or not response.choices[0].message.content:
                raise ProcessingError(
                    "AI响应无效或为空",
                    "AI_NO_RESPONSE",
                    500
                )
            
            # 提取JSON内容
            response_content = response.choices[0].message.content.strip()
            json_str = self._extract_json_from_response(response_content)
            
            try:
                merged_content = json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON解析失败: {e}")
                logger.error(f"AI响应内容: {response_content}")
                raise ProcessingError(
                    f"AI返回的内容不是有效的JSON格式: {str(e)}",
                    "AI_INVALID_JSON",
                    422
                )
            
            # 验证合并结果
            if not isinstance(merged_content, dict):
                raise ProcessingError(
                    "AI返回的内容不是字典格式",
                    "AI_INVALID_FORMAT",
                    422
                )
            
            logger.info(f"✅ AI合并成功，生成 {len(merged_content)} 个章节")
            return merged_content
            
        except ProcessingError:
            raise
        except Exception as e:
            logger.error(f"❌ AI合并失败: {e}")
            raise ProcessingError(
                f"AI合并过程中发生错误: {str(e)}",
                "AI_MERGE_ERROR",
                500
            )
    
    def _mock_merge_content(self, template_json: Dict[str, str], original_content: str) -> Dict[str, str]:
        """模拟AI合并（测试模式）"""
        logger.info("🧪 模拟AI合并模式")
        
        merged_content = {}
        content_lines = original_content.split('\n')
        content_preview = ' '.join(content_lines[:5])[:200]
        
        for key, description in template_json.items():
            merged_content[key] = f"""根据原始文档内容生成的{key}章节：

{description}

基于原始文档的相关信息：
{content_preview}

本章节内容已根据模板要求进行智能整合，确保符合工程文档的标准格式和要求。"""
        
        logger.info(f"✅ 模拟合并完成，生成 {len(merged_content)} 个章节")
        return merged_content
    
    def _extract_json_from_response(self, response_content: str) -> str:
        """从AI响应中提取JSON内容"""
        # 尝试提取JSON
        if "```json" in response_content:
            start = response_content.find("```json") + 7
            end = response_content.find("```", start)
            if end != -1:
                return response_content[start:end].strip()
            else:
                return response_content[response_content.find("```json") + 7:].strip()
        elif response_content.startswith("{") and response_content.endswith("}"):
            return response_content
        else:
            # 查找JSON对象
            start_idx = response_content.find("{")
            if start_idx != -1:
                brace_count = 0
                for i, char in enumerate(response_content[start_idx:], start_idx):
                    if char == "{":
                        brace_count += 1
                    elif char == "}":
                        brace_count -= 1
                        if brace_count == 0:
                            return response_content[start_idx:i+1]
            return response_content

class TemplateMergerTool:
    """模板合并工具 - AI智能合并原始文档与模板"""
    
    def __init__(self, deepseek_client=None):
        """初始化工具"""
        self.name = "template_merger"
        self.description = """📋 模板合并工具 - AI智能合并原始文档与模板JSON

🎯 **核心功能：**
- 📄 智能内容提取：支持DOC/DOCX/PDF/TXT格式文档
- 🧠 AI智能合并：将原始文档内容与模板JSON智能合并
- 📝 结构化生成：生成符合模板要求的结构化文档
- 🔄 格式转换：自动处理DOC到DOCX的转换

💡 **适用场景：**
- 基于模板生成标准化文档
- 将非结构化内容转换为结构化文档
- 批量文档处理和标准化
- 内容智能整合和重组

⚙️ **使用方法：**
Action Input: {"action": "merge", "template_json": {"章节1": "描述"}, "content_source": "文件路径或内容文本"}"""
        
        # 确保输出目录存在
        self.output_dir = "generated_docs"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        logger.info("📋 模板合并工具初始化完成")
    
    def get_tool_info(self) -> Dict[str, Any]:
        """获取工具信息"""
        return {
            "name": self.name,
            "description": self.description,
            "parameters": {
                "action": {
                    "type": "string",
                    "description": "操作类型",
                    "enum": ["merge", "extract", "analyze"]
                },
                "template_json": {
                    "type": "object",
                    "description": "模板JSON结构，键为章节名，值为章节描述"
                },
                "content_source": {
                    "type": "string",
                    "description": "内容源：可以是文件路径或直接的文本内容"
                },
                "output_path": {
                    "type": "string",
                    "description": "可选的输出路径，如果不提供则自动生成"
                }
            },
            "required": ["action"]
        }
    
    def execute(self, **kwargs) -> str:
        """执行工具操作"""
        action = kwargs.get("action", "").lower()
        
        try:
            if action == "merge":
                return self._merge_template(kwargs)
            elif action == "extract":
                return self._extract_content(kwargs)
            elif action == "analyze":
                return self._analyze_template(kwargs)
            else:
                return f"❌ 不支持的操作: {action}。支持的操作: merge, extract, analyze"
                
        except Exception as e:
            error_msg = f"❌ 模板合并工具执行失败: {str(e)}"
            logger.error(error_msg)
            return error_msg
    
    def _get_api_key(self) -> str:
        """获取OpenRouter API密钥"""
        api_key = os.environ.get("OPENROUTER_API_KEY")
        if not api_key:
            # 检查是否是测试模式
            test_mode = os.environ.get("TEST_MODE", "false").lower() == "true"
            if test_mode:
                logger.warning("⚠️ 测试模式：使用模拟API密钥")
                return "test-api-key-for-testing"
            
            logger.error("❌ 未找到OPENROUTER_API_KEY")
            raise RuntimeError("缺少必需的API密钥配置")
        return api_key
    
    def _merge_template(self, params: Dict[str, Any]) -> str:
        """执行模板合并操作"""
        logger.info("🚀 开始模板合并操作...")
        
        # 获取参数
        template_json = params.get("template_json")
        content_source = params.get("content_source")
        output_path = params.get("output_path")
        
        if not template_json:
            return "❌ 缺少template_json参数"
        
        if not content_source:
            return "❌ 缺少content_source参数"
        
        try:
            # 获取API密钥
            api_key = self._get_api_key()
            
            # 提取内容
            if os.path.exists(content_source):
                # 从文件提取内容
                extractor = DocumentExtractor()
                original_content = extractor.extract_from_file_path(content_source)
                logger.info("✅ 从文件提取内容成功")
            else:
                # 直接使用提供的内容
                original_content = content_source
                logger.info("✅ 使用直接提供的内容")
            
            # 执行AI合并
            merger = ContentMerger(api_key)
            merged_content = merger.merge_content(template_json, original_content)
            
            # 生成输出路径
            if not output_path:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"merged_document_{timestamp}.docx"
                output_path = os.path.join(self.output_dir, output_filename)
            
            # 生成DOCX文档
            doc = Document()
            doc.add_heading('AI智能合并文档', 0)
            
            for section_title, section_content in merged_content.items():
                doc.add_heading(str(section_title), level=1)
                doc.add_paragraph(str(section_content))
            
            doc.save(output_path)
            
            # 分析缺失字段
            missing_fields = []
            for key, value in merged_content.items():
                if not value or len(str(value).strip()) < 10:
                    missing_fields.append(key)
                elif "待填写" in str(value) or "需要补充" in str(value) or "____" in str(value):
                    missing_fields.append(key)
            
            # 生成结果报告
            completion_rate = (len(template_json) - len(missing_fields)) / len(template_json) if template_json else 0
            
            result = f"""✅ 模板合并完成！

📄 **生成文档**: {output_path}
📊 **统计信息**:
   - 模板章节: {len(template_json)} 个
   - 已填充章节: {len(template_json) - len(missing_fields)} 个
   - 完成率: {completion_rate:.1%}
   - 文档大小: {os.path.getsize(output_path) if os.path.exists(output_path) else 0} 字节

📋 **章节详情**:"""
            
            for key, value in merged_content.items():
                status = "✅" if key not in missing_fields else "⚠️"
                preview = str(value)[:50] + "..." if len(str(value)) > 50 else str(value)
                result += f"\n   {status} {key}: {preview}"
            
            if missing_fields:
                result += f"\n\n⚠️ **需要补充的章节**: {', '.join(missing_fields)}"
            
            return result
            
        except Exception as e:
            logger.error(f"❌ 模板合并失败: {e}")
            return f"❌ 模板合并失败: {str(e)}"
    
    def _extract_content(self, params: Dict[str, Any]) -> str:
        """执行内容提取操作"""
        logger.info("📄 开始内容提取操作...")
        
        content_source = params.get("content_source")
        if not content_source:
            return "❌ 缺少content_source参数"
        
        if not os.path.exists(content_source):
            return f"❌ 文件不存在: {content_source}"
        
        try:
            extractor = DocumentExtractor()
            content = extractor.extract_from_file_path(content_source)
            
            result = f"""✅ 内容提取完成！

📄 **源文件**: {content_source}
📊 **统计信息**:
   - 内容长度: {len(content)} 字符
   - 文件大小: {os.path.getsize(content_source)} 字节

📝 **内容预览**:
{content[:500]}{'...' if len(content) > 500 else ''}"""
            
            return result
            
        except Exception as e:
            logger.error(f"❌ 内容提取失败: {e}")
            return f"❌ 内容提取失败: {str(e)}"
    
    def _analyze_template(self, params: Dict[str, Any]) -> str:
        """分析模板结构"""
        logger.info("🔍 开始模板分析操作...")
        
        template_json = params.get("template_json")
        if not template_json:
            return "❌ 缺少template_json参数"
        
        try:
            result = f"""📋 模板结构分析结果：

📊 **基本统计**:
   - 章节总数: {len(template_json)} 个
   - 模板类型: {"复杂模板" if len(template_json) > 5 else "简单模板"}

📝 **章节详情**:"""
            
            for i, (key, value) in enumerate(template_json.items(), 1):
                description_length = len(str(value))
                complexity = "详细" if description_length > 100 else "简单" if description_length > 20 else "基础"
                result += f"\n   {i}. **{key}** ({complexity}): {str(value)[:100]}{'...' if len(str(value)) > 100 else ''}"
            
            # 提供使用建议
            result += f"\n\n💡 **使用建议**:"
            if len(template_json) > 10:
                result += "\n   - 模板较复杂，建议分批处理"
            if any(len(str(v)) < 10 for v in template_json.values()):
                result += "\n   - 部分章节描述较简单，可能需要更详细的说明"
            
            result += "\n   - 准备好原始内容后，使用merge操作进行智能合并"
            
            return result
            
        except Exception as e:
            logger.error(f"❌ 模板分析失败: {e}")
            return f"❌ 模板分析失败: {str(e)}" 