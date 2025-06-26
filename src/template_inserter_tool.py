#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板填充工具：AI智能填充Word文档
功能：
1. 接收JSON格式的模板结构和内容数据
2. 使用AI智能映射和匹配内容
3. 生成填充完成的DOCX文档
"""

import os
import json
import logging
import re
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches
from openai import OpenAI

# Load environment variables from .env file if it exists
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    # python-dotenv not installed, skip .env file loading
    pass

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

class TemplateInserter:
    """模板填充器 - 使用AI智能填充模板内容"""
    
    def __init__(self, api_key: str):
        """
        初始化模板填充器
        
        Args:
            api_key: OpenRouter API密钥
        """
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"
        logger.info("🤖 AI模板填充器初始化完成")
    
    def _extract_json_from_response(self, response_content: str) -> str:
        """
        从AI响应中提取JSON字符串
        处理多种格式（markdown代码块、纯JSON等）
        """
        if not response_content or not response_content.strip():
            raise ValueError("AI response content is empty")
        
        content = response_content.strip()
        
        # Try to extract from markdown code block
        if "```json" in content:
            try:
                start = content.find("```json") + 7
                end = content.find("```", start)
                if end != -1:
                    json_str = content[start:end].strip()
                    if json_str:
                        return json_str
            except Exception:
                pass
        
        # Try to extract from single backticks
        if content.startswith("`") and content.endswith("`"):
            json_str = content.strip("`").strip()
            if json_str:
                return json_str
        
        # Try to find JSON object boundaries
        start_idx = content.find("{")
        if start_idx != -1:
            # Find the matching closing brace
            brace_count = 0
            for i, char in enumerate(content[start_idx:], start_idx):
                if char == "{":
                    brace_count += 1
                elif char == "}":
                    brace_count -= 1
                    if brace_count == 0:
                        json_str = content[start_idx:i+1]
                        # Validate it's proper JSON
                        try:
                            json.loads(json_str)
                            return json_str
                        except json.JSONDecodeError:
                            continue
        
        # If all else fails, try the content as-is
        try:
            json.loads(content)
            return content
        except json.JSONDecodeError:
            raise ValueError(f"Could not extract valid JSON from AI response: {content[:200]}...")

    def ai_generate_fill_mapping(
        self, 
        template_structure: Dict[str, str], 
        placeholders: List[str], 
        input_data: Dict[str, Any]
    ) -> Dict[str, str]:
        """
        使用AI生成填充映射数据
        
        Args:
            template_structure: 模板结构字典
            placeholders: 占位符列表
            input_data: 输入数据
            
        Returns:
            填充映射字典
        """
        logger.info("🧠 开始AI智能字段映射...")
        
        try:
            # 构建AI提示
            prompt = self._get_fill_mapping_prompt(
                json.dumps(template_structure, ensure_ascii=False, indent=2),
                json.dumps(placeholders, ensure_ascii=False, indent=2),
                json.dumps(input_data, ensure_ascii=False, indent=2)
            )
            
            logger.info("🧠 正在调用AI进行智能字段映射...")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            
            if not response or not response.choices or not response.choices[0].message.content:
                raise ValueError("AI响应无效或为空")

            json_string = self._extract_json_from_response(response.choices[0].message.content)
            fill_data = json.loads(json_string)
            
            logger.info(f"✅ AI成功生成 {len(fill_data)} 个字段的映射:")
            for key, value in fill_data.items():
                preview = str(value)[:70] + "..." if len(str(value)) > 70 else str(value)
                logger.info(f"   🔗 {key} -> '{preview}'")
            
            return fill_data
            
        except Exception as e:
            logger.error(f"❌ AI字段映射错误: {e}", exc_info=True)
            return {}

    def fill_template_with_data(
        self, 
        processed_template_path: str, 
        output_path: str, 
        fill_data: Dict[str, str],
        placeholder_originals: Dict[str, str]
    ):
        """
        使用填充数据填充模板
        
        Args:
            processed_template_path: 预处理后的模板路径
            output_path: 输出文档路径
            fill_data: 填充数据映射
            placeholder_originals: 原始占位符文本映射
        """
        logger.info("📝 开始模板内容填充...")
        
        doc = Document(processed_template_path)
        filled_count = 0
        
        # 1. 准备附件信息
        attachments_map = fill_data.pop('attachments_map', {})
        attachment_ref_map = {}
        ordered_attachments = []
        if attachments_map and isinstance(attachments_map, dict):
            logger.info(f"🖼️  找到 {len(attachments_map)} 个图片附件待处理。")
            ordered_attachments = list(attachments_map.items())
            for i, (key, _) in enumerate(ordered_attachments):
                attachment_ref_map[key.strip()] = i + 1
        else:
            attachments_map = {}

        # 2. 分离文本填充数据
        placeholder_data = {k: v for k, v in fill_data.items() if k.startswith(('label_', 'inline_', 'blank_'))}
        structure_data = {k: v for k, v in fill_data.items() if k.startswith(('table_', 'paragraph_'))}
        
        # 3. 替换所有文本占位符（包括图片引用）
        image_placeholder_pattern = re.compile(r'\{\{image:([^}]+)\}\}')
        text_placeholder_pattern = re.compile(r'\{(label_[^}]+|inline_[^}]+|blank_[^}]+)\}')

        def process_element_text(element):
            nonlocal filled_count
            if '{' not in element.text:
                return

            original_text = element.text
            new_text = ""
            last_end = 0

            # 创建一个组合的正则表达式来查找所有类型的占位符
            combined_pattern = re.compile(f"({image_placeholder_pattern.pattern}|{text_placeholder_pattern.pattern})")
            
            for match in combined_pattern.finditer(original_text):
                new_text += original_text[last_end:match.start()]
                last_end = match.end()
                
                image_key_match = image_placeholder_pattern.match(match.group(0))
                text_key_match = text_placeholder_pattern.match(match.group(0))

                if image_key_match:
                    key = image_key_match.group(1).strip()
                    if key in attachment_ref_map:
                        number = attachment_ref_map[key]
                        replacement = f"（详见附件{number}）"
                        new_text += replacement
                        logger.info(f"   🖼️  图片引用替换: '{match.group(0)}' -> '{replacement}'")
                    else:
                        logger.warning(f"   ⚠️  找到图片占位符 {match.group(0)} 但无匹配图片，已移除。")
                
                elif text_key_match:
                    placeholder_key = text_key_match.group(1)
                    placeholder = f"{{{placeholder_key}}}"

                    if placeholder_key in placeholder_data:
                        value = str(placeholder_data[placeholder_key])
                        new_text += value
                        logger.info(f"   ✏️  占位符填充: {placeholder} -> {value[:50]}...")
                        filled_count += 1
                    else: # 未匹配的文本占位符
                        if placeholder_key.startswith('label_'):
                            logger.info(f"   🔘  移除未匹配标签占位符: {placeholder}")
                            # The replacement is empty string, so we add nothing
                        elif placeholder_key.startswith(('inline_', 'blank_')):
                            original_underscore = placeholder_originals.get(placeholder_key, '____')
                            new_text += original_underscore
                            logger.info(f"   🔘  恢复未匹配占位符: {placeholder} -> '{original_underscore}'")

            new_text += original_text[last_end:]
            
            if new_text != original_text:
                element.text = new_text

        # 遍历段落和表格进行统一替换
        for para in doc.paragraphs:
            process_element_text(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_element_text(cell)

        # 4. 填充原始结构
        combined_pattern = re.compile(f"({image_placeholder_pattern.pattern}|{text_placeholder_pattern.pattern})")
        
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    cell_key = f"table_{i}_row_{j}_col_{k}"
                    if cell_key in structure_data:
                        cell.text = str(structure_data[cell_key])
                        logger.info(f"   ✏️  结构填充(表格): {cell_key} -> {str(structure_data[cell_key])[:50]}...")
                        filled_count += 1

        for i, para in enumerate(doc.paragraphs):
            para_key = f"paragraph_{i}"
            if para_key in structure_data:
                # 只有在段落中没有占位符的情况下才进行结构填充
                if not combined_pattern.search(para.text):
                    para.text = str(structure_data[para_key])
                    logger.info(f"   ✏️  结构填充(段落): {para_key} -> {str(structure_data[para_key])[:50]}...")
                    filled_count += 1

        # 5. 将图片作为附件附加到文档末尾
        if ordered_attachments:
            logger.info("📎 开始在文档末尾附加图片...")
            try:
                doc.add_page_break()
                doc.add_heading('附件列表', level=1)
                
                for i, (key, image_path) in enumerate(ordered_attachments):
                    attachment_counter = i + 1
                    if not image_path or not isinstance(image_path, str) or not os.path.exists(image_path):
                        logger.warning(f"⚠️ 图片路径不存在或无效，跳过附件 '{key}': {image_path}")
                        continue
                    
                    try:
                        heading_text = f"附件 {attachment_counter}: {key}"
                        doc.add_heading(heading_text, level=2)
                        doc.add_picture(image_path, width=Inches(6.0))
                        doc.add_paragraph()
                        logger.info(f"   ✅ 成功附加图片: {heading_text} ({image_path})")
                    except Exception as pic_e:
                        logger.error(f"❌ 附加图片 '{key}' ({image_path}) 时出错: {pic_e}")
            except Exception as e:
                logger.error(f"❌ 处理附件时发生意外错误: {e}")
        
        doc.save(output_path)
        logger.info(f"✅ 模板填充完成，共填充 {filled_count} 个字段: {output_path}")

    def insert_content_to_template(
        self, 
        json_template_path: str, 
        content_data: Dict[str, Any], 
        output_path: str
    ) -> bool:
        """
        主函数：将内容插入到JSON格式的模板中
        
        Args:
            json_template_path: JSON模板文件路径
            content_data: 要插入的内容数据
            output_path: 输出DOCX文件路径
            
        Returns:
            是否成功
        """
        logger.info("🚀 开始智能模板填充流程...")
        
        try:
            # Step 1: 加载JSON模板
            logger.info(f"📂 加载JSON模板: {json_template_path}")
            with open(json_template_path, 'r', encoding='utf-8') as f:
                template_data = json.load(f)
            
            # 验证模板结构
            required_keys = ['template_structure', 'placeholders', 'placeholder_originals']
            for key in required_keys:
                if key not in template_data:
                    raise ValueError(f"JSON模板缺少必需的键: {key}")
            
            template_structure = template_data['template_structure']
            placeholders = template_data['placeholders']
            placeholder_originals = template_data['placeholder_originals']
            processed_template_path = template_data['template_info']['processed_template_path']
            
            logger.info(f"✅ 模板加载成功:")
            logger.info(f"   - 结构元素: {len(template_structure)}")
            logger.info(f"   - 占位符: {len(placeholders)}")
            
            # Step 2: AI智能映射
            fill_mapping = self.ai_generate_fill_mapping(
                template_structure=template_structure,
                placeholders=placeholders,
                input_data=content_data
            )
            
            if not fill_mapping:
                logger.warning("⚠️  AI未能生成有效的填充映射，将使用空映射")
            
            # Step 3: 填充模板
            self.fill_template_with_data(
                processed_template_path=processed_template_path,
                output_path=output_path,
                fill_data=fill_mapping,
                placeholder_originals=placeholder_originals
            )
            
            logger.info(f"✅ 智能模板填充完成: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ 模板填充失败: {e}", exc_info=True)
            return False
    
    def _get_fill_mapping_prompt(self, template_structure: str, placeholders: str, input_data: str) -> str:
        """
        生成AI填充映射提示
        """
        return f"""
你是一个专业的混合模式文档填写助手。你需要同时处理三种任务：
1. **模板结构匹配**：将JSON数据映射到Word文档的表格单元格和段落结构中
2. **占位符匹配**：为特定的占位符找到对应的数据值
3. **图片引用处理**：处理图片占位符，确保图片能正确引用

---

**核心任务：生成统一的填充数据映射**

你将获得：
- **模板结构**：文档中所有单元格和段落的结构化表示
- **占位符列表**：从特定模式（如"项目名称："和"致____（监理单位）"）提取的占位符
- **输入数据**：需要映射的JSON数据（可能包含attachments_map图片信息）

你的输出应该是一个JSON对象，包含三种类型的映射：
- **结构映射**：键如"table_0_row_1_col_1"或"paragraph_3"，用于填充模板结构
- **占位符映射**：键如"label_项目名称"或"inline_监理单位"，用于替换占位符
- **图片映射**：如果输入数据包含attachments_map，直接将其包含在输出中

---

**重要原则：只填充确定的数据**

⚠️ **关键要求**：
- **只有在输入数据中找到明确对应的信息时，才输出该键值对**
- **如果找不到对应的数据，请完全省略该键，不要输出任何占位文本如"待填写"、"____"等**
- **宁可留空，也不要填入不确定的内容**

---

**映射规则**

### 1. 模板结构映射（保持原有逻辑）
- **排除规则**：如果单元格或段落文本中包含 **"（签字）"** 字样，**绝对不要**为其填充任何内容。
- 如果一个单元格的内容是"标签："（例如 `"项目名称："`），并且其右侧或下方的单元格为空，则应将数据填入那个空单元格
- 如果一个单元格包含下划线占位符但**不是特定的两种模式**，则直接替换该单元格内容
- 智能匹配：运用推理能力，如`project_leader`应映射到"项目负责人"相关的位置

### 2. 占位符映射（现有逻辑）
- 对于`label_*`类型的占位符（来自"项目名称："），找到语义匹配的数据
- 对于`inline_*`类型的占位符（来自"致____（监理单位）"），找到对应的值
- 占位符名称是提示性的，需要智能匹配到输入数据的字段
- **如果确实找不到匹配的数据，就不要输出这个键**

### 3. 图片映射（新增功能）
- 如果输入数据包含`attachments_map`字段，**必须**将其完整地包含在输出JSON中
- 这将确保图片信息能被正确传递到文档生成阶段
- 图片占位符`{{image:key}}`会在文档生成时被自动替换为"（详见附件N）"的格式

---

现在请根据下方的数据进行混合映射：

**模板结构:**
```json
{template_structure}
```

**占位符列表:**
```json
{placeholders}
```

**输入数据:**
```json
{input_data}
```

**重要要求:**
- **所有生成的内容必须使用中文**
- 输出统一的JSON对象，包含所有类型的映射
- 对占位符进行智能语义匹配
- 保持原有的模板结构填充逻辑
- **如果输入数据包含attachments_map，必须完整包含在输出中**
- **重要：如果找不到对应数据，完全省略该键，不要填入"待填写"等占位文本**
- 只输出最终的JSON对象，不要包含解释说明或Markdown格式
"""


class TemplateInserterTool:
    """模板填充工具 - ReAct Agent工具封装"""
    
    def __init__(self):
        self.name = "template_inserter"
        self.description = "AI智能模板填充工具 - 将JSON内容数据智能填充到模板中生成Word文档"
        
        # 获取API密钥
        self.api_key = os.environ.get("OPENROUTER_API_KEY")
        if not self.api_key:
            logger.warning("⚠️  未找到OPENROUTER_API_KEY环境变量，AI功能将不可用")
            self.inserter = None
        else:
            self.inserter = TemplateInserter(self.api_key)
    
    def execute(self, **kwargs) -> str:
        """
        执行模板填充操作
        
        参数:
        - json_template_path: JSON模板文件路径 (必需)
        - content_data: 内容数据字典或JSON文件路径 (必需)
        - output_path: 输出DOCX文件路径 (可选，默认自动生成)
        
        返回:
        执行结果描述
        """
        try:
            # 检查AI功能是否可用
            if not self.inserter:
                return "❌ 错误: 未配置OPENROUTER_API_KEY环境变量，无法使用AI填充功能"
            
            # 获取参数
            json_template_path = kwargs.get('json_template_path')
            content_data = kwargs.get('content_data')
            output_path = kwargs.get('output_path')
            
            # 验证必需参数
            if not json_template_path:
                return "❌ 错误: 缺少必需参数 json_template_path"
            
            if not content_data:
                return "❌ 错误: 缺少必需参数 content_data"
            
            # 检查JSON模板文件是否存在
            if not os.path.exists(json_template_path):
                return f"❌ 错误: JSON模板文件不存在: {json_template_path}"
            
            # 处理content_data参数
            if isinstance(content_data, str):
                # 如果是字符串，可能是文件路径或JSON字符串
                if os.path.exists(content_data):
                    # 是文件路径，加载JSON文件
                    try:
                        with open(content_data, 'r', encoding='utf-8') as f:
                            content_data = json.load(f)
                    except Exception as e:
                        return f"❌ 错误: 加载内容数据文件失败: {e}"
                else:
                    # 尝试解析为JSON字符串
                    try:
                        content_data = json.loads(content_data)
                    except json.JSONDecodeError as e:
                        return f"❌ 错误: 无效的JSON内容数据: {e}"
            elif not isinstance(content_data, dict):
                return "❌ 错误: content_data必须是字典、JSON字符串或JSON文件路径"
            
            # 生成输出路径（如果未提供）
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = os.path.splitext(os.path.basename(json_template_path))[0]
                output_path = f"generated_documents/filled_{base_name}_{timestamp}.docx"
                
                # 确保输出目录存在
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # 执行模板填充
            logger.info(f"🚀 开始执行模板填充:")
            logger.info(f"   - JSON模板: {json_template_path}")
            logger.info(f"   - 内容数据字段数: {len(content_data)}")
            logger.info(f"   - 输出路径: {output_path}")
            
            success = self.inserter.insert_content_to_template(
                json_template_path=json_template_path,
                content_data=content_data,
                output_path=output_path
            )
            
            if success:
                # 检查输出文件大小
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    file_size_kb = file_size / 1024
                    
                    return f"""✅ 模板填充成功完成！

📄 **输出文件信息:**
- 文件路径: {output_path}
- 文件大小: {file_size_kb:.1f} KB
- 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

🎯 **填充统计:**
- JSON模板: {os.path.basename(json_template_path)}
- 内容数据字段: {len(content_data)} 个
- AI智能映射: 已完成

💡 **使用提示:**
可以直接打开生成的DOCX文件查看填充结果。如需要调整，可以修改内容数据后重新执行填充。"""
                else:
                    return f"⚠️  填充过程显示成功，但输出文件不存在: {output_path}"
            else:
                return "❌ 模板填充失败，请查看日志了解详细错误信息"
                
        except Exception as e:
            logger.error(f"❌ 模板填充工具执行错误: {e}", exc_info=True)
            return f"❌ 工具执行失败: {str(e)}"
    
    def get_usage_guide(self) -> str:
        """获取工具使用指南"""
        return """
🔧 **模板填充工具使用指南**

📋 **功能说明:**
- 使用AI智能将JSON内容数据填充到模板中
- 支持复杂的模板结构映射和占位符替换
- 自动处理图片附件和文档格式

⚙️  **参数说明:**

**必需参数:**
- `json_template_path`: JSON模板文件路径
  - 由模板转换工具生成的JSON格式模板文件
  - 包含模板结构、占位符等信息

- `content_data`: 内容数据
  - 可以是字典对象
  - 可以是JSON文件路径
  - 可以是JSON字符串

**可选参数:**
- `output_path`: 输出DOCX文件路径
  - 如未指定，将自动生成到generated_documents目录

💡 **使用示例:**

```python
# 方式1: 使用字典数据
tool.execute(
    json_template_path="templates/template.json",
    content_data={
        "project_name": "滨海新城商业综合体",
        "location": "滨海新区",
        "attachments_map": {
            "项目效果图": "images/effect.jpg"
        }
    }
)

# 方式2: 使用JSON文件
tool.execute(
    json_template_path="templates/template.json",
    content_data="data/project_data.json",
    output_path="output/filled_document.docx"
)
```

🎯 **AI智能特性:**
- 自动匹配字段语义（如project_name → 项目名称）
- 智能处理表格和段落结构
- 支持图片附件自动引用
- 只填充确定的数据，避免错误占位

⚠️  **注意事项:**
- 需要配置OPENROUTER_API_KEY环境变量
- JSON模板必须由模板转换工具生成
- 图片路径必须是有效的本地文件路径
"""


def main():
    """命令行测试函数"""
    import sys
    
    if len(sys.argv) < 3:
        print("用法: python template_inserter_tool.py <json_template> <content_data> [output_docx]")
        print("示例: python template_inserter_tool.py template.json content.json output.docx")
        return
    
    json_template_path = sys.argv[1]
    content_data = sys.argv[2]
    output_path = sys.argv[3] if len(sys.argv) > 3 else None
    
    # 创建工具实例并执行
    tool = TemplateInserterTool()
    result = tool.execute(
        json_template_path=json_template_path,
        content_data=content_data,
        output_path=output_path
    )
    
    print(result)


if __name__ == "__main__":
    main()