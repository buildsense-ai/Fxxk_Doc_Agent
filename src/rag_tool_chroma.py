"""
ChromaDB RAG工具 - 重新设计版
核心功能：文档embedding处理、智能搜索、基于模板字段的内容填充
"""
import os
import json
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
import hashlib

try:
    import chromadb
    from chromadb.config import Settings
    import fitz  # PyMuPDF for PDF
    from docx import Document as DocxDocument
except ImportError as e:
    print(f"警告: RAG工具依赖未安装: {e}")
    print("请安装: pip install chromadb PyMuPDF python-docx")

try:
    from .tools import Tool
except ImportError:
    from tools import Tool

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentExtractor:
    """文档内容提取器"""
    
    def extract_content(self, file_path: str) -> str:
        """提取文档内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            elif file_ext in ['.txt', '.md']:
                return self._extract_from_text(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
        except Exception as e:
            raise RuntimeError(f"文档内容提取失败: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """从PDF提取内容"""
        content = ""
        doc = fitz.open(file_path)
        for page in doc:
            content += page.get_text()
        doc.close()
        return content.strip()
    
    def _extract_from_docx(self, file_path: str) -> str:
        """从DOCX提取内容"""
        doc = DocxDocument(file_path)
        content = []
        
        # 提取段落内容
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    content.append(f"表格行: {row_text}")
        
        return "\n".join(content)
    
    def _extract_from_text(self, file_path: str) -> str:
        """从文本文件提取内容"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()

class ChromaVectorStore:
    """ChromaDB向量存储"""
    
    def __init__(self, storage_dir: str):
        self.storage_dir = storage_dir
        self.client = chromadb.PersistentClient(
            path=storage_dir,
            settings=Settings(allow_reset=True, anonymized_telemetry=False)
        )
        self.collection = self.client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
    
    def add_document(self, doc_id: str, content: str, metadata: Dict[str, Any]):
        """添加文档到向量库"""
        # 将长文档分块
        chunks = self._split_content(content)
        
        ids = []
        documents = []
        metadatas = []
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}_chunk_{i}"
            ids.append(chunk_id)
            documents.append(chunk)
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": i,
                "total_chunks": len(chunks)
            })
            metadatas.append(chunk_metadata)
        
        self.collection.add(
            ids=ids,
            documents=documents,
            metadatas=metadatas
        )
        
        return len(chunks)
    
    def search_documents(self, query: str, n_results: int = 5) -> List[Dict[str, Any]]:
        """搜索相关文档"""
        results = self.collection.query(
            query_texts=[query],
            n_results=n_results
        )
        
        formatted_results = []
        if results['documents'] and len(results['documents']) > 0:
            for i in range(len(results['documents'][0])):
                formatted_results.append({
                    'content': results['documents'][0][i],
                    'metadata': results['metadatas'][0][i],
                    'distance': results['distances'][0][i]
                })
        
        return formatted_results
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """获取所有文档信息"""
        try:
            results = self.collection.get()
            documents = []
            
            # 按文档ID分组
            doc_groups = {}
            for i, doc_id in enumerate(results['ids']):
                base_id = doc_id.split('_chunk_')[0]
                if base_id not in doc_groups:
                    doc_groups[base_id] = {
                        'id': base_id,
                        'metadata': results['metadatas'][i],
                        'chunks': 0
                    }
                doc_groups[base_id]['chunks'] += 1
            
            return list(doc_groups.values())
        except Exception as e:
            logger.error(f"获取文档列表失败: {e}")
            return []
    
    def clear_all(self):
        """清空所有文档"""
        self.client.delete_collection("documents")
        self.collection = self.client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
    
    def _split_content(self, content: str, chunk_size: int = 1000) -> List[str]:
        """将内容分块"""
        if len(content) <= chunk_size:
            return [content]
        
        chunks = []
        sentences = content.split('。')
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk + sentence + '。') <= chunk_size:
                current_chunk += sentence + '。'
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + '。'
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [content]

class TemplateFieldProcessor:
    """模板字段处理器 - 核心智能填充功能"""
    
    def __init__(self, deepseek_client=None):
        self.deepseek_client = deepseek_client
    
    def fill_template_fields(self, template_fields_json: Dict[str, str], 
                           vector_store: ChromaVectorStore) -> Dict[str, str]:
        """
        基于模板字段JSON进行智能搜索和内容填充
        
        Args:
            template_fields_json: 模板字段JSON，格式为 {"字段名": "字段描述或要求"}
            vector_store: 向量存储实例
            
        Returns:
            填充好的字段JSON，格式为 {"字段名": "填充的具体内容"}
        """
        logger.info(f"🔍 开始基于模板字段进行智能填充，共 {len(template_fields_json)} 个字段")
        
        filled_fields = {}
        
        for field_name, field_requirement in template_fields_json.items():
            logger.info(f"📝 处理字段: {field_name}")
            
            # 1. 基于字段要求搜索相关内容
            search_results = vector_store.search_documents(
                query=f"{field_name} {field_requirement}",
                n_results=3
            )
            
            # 2. 提取搜索到的内容
            relevant_content = []
            for result in search_results:
                relevant_content.append(result['content'])
            
            # 3. 使用AI生成字段内容（如果有AI客户端）
            if self.deepseek_client and relevant_content:
                filled_content = self._generate_field_content_with_ai(
                    field_name, field_requirement, relevant_content
                )
            else:
                # 基础模式：直接使用搜索到的最相关内容
                filled_content = self._generate_field_content_basic(
                    field_name, field_requirement, relevant_content
                )
            
            filled_fields[field_name] = filled_content
            logger.info(f"✅ 字段 {field_name} 填充完成，内容长度: {len(filled_content)} 字符")
        
        logger.info(f"🎉 所有字段填充完成！")
        return filled_fields
    
    def _generate_field_content_with_ai(self, field_name: str, field_requirement: str, 
                                       relevant_content: List[str]) -> str:
        """使用AI生成字段内容"""
        content_text = "\n".join(relevant_content)
        
        prompt = f"""
你是一个专业的文档处理助手。请根据以下信息为字段生成合适的内容：

字段名称：{field_name}
字段要求：{field_requirement}

相关资料内容：
{content_text}

任务要求：
1. 基于相关资料内容，为该字段生成专业、准确的内容
2. 内容应该符合字段要求和描述
3. 保持内容的专业性和完整性
4. 如果资料内容不足，请基于字段要求进行合理补充
5. 内容长度适中，重点突出

请直接返回该字段的具体内容，不要包含解释文字。
"""
        
        try:
            response = self.deepseek_client.chat([{"role": "user", "content": prompt}])
            return response.strip() if response else self._generate_field_content_basic(
                field_name, field_requirement, relevant_content
            )
        except Exception as e:
            logger.warning(f"AI生成失败，使用基础模式: {e}")
            return self._generate_field_content_basic(
                field_name, field_requirement, relevant_content
            )
    
    def _generate_field_content_basic(self, field_name: str, field_requirement: str, 
                                     relevant_content: List[str]) -> str:
        """基础模式生成字段内容"""
        if not relevant_content:
            return f"[{field_name}]：{field_requirement}（待补充具体内容）"
        
        # 选择最相关的内容作为基础
        base_content = relevant_content[0]
        
        # 简单的内容处理
        if len(base_content) > 200:
            # 截取前200字符作为摘要
            summary = base_content[:200] + "..."
            return f"{summary}\n\n（基于相关资料整理，如需详细信息请参考原始文档）"
        else:
            return base_content

class RAGTool(Tool):
    """重新设计的RAG工具 - 核心三功能"""
    
    def __init__(self, storage_dir: str = "rag_storage", deepseek_client=None):
        super().__init__(
            name="rag_tool",
            description="文档处理和智能检索工具。支持文档embedding处理、基于模板字段的智能搜索和内容填充。"
        )
        
        self.storage_dir = storage_dir
        os.makedirs(storage_dir, exist_ok=True)
        
        # 初始化组件
        self.extractor = DocumentExtractor()
        self.vector_store = ChromaVectorStore(storage_dir)
        self.field_processor = TemplateFieldProcessor(deepseek_client)
    
    def execute(self, action: str, **kwargs) -> str:
        """执行RAG操作"""
        try:
            if action == "upload":
                return self._upload_document(**kwargs)
            elif action == "search":
                return self._search_documents(**kwargs)
            elif action == "fill_fields":
                return self._fill_template_fields(**kwargs)
            elif action == "list":
                return self._list_documents()
            elif action == "clear":
                return self._clear_documents()
            else:
                return f"❌ 不支持的操作: {action}\n支持的操作: upload, search, fill_fields, list, clear"
        
        except Exception as e:
            logger.error(f"RAG操作失败: {e}")
            return f"❌ 操作失败: {str(e)}"
    
    def _upload_document(self, file_path: str, filename: str = "") -> str:
        """上传文档进行embedding处理"""
        try:
            if not filename:
                filename = os.path.basename(file_path)
            
            logger.info(f"📄 开始处理文档: {filename}")
            
            # 1. 提取文档内容
            content = self.extractor.extract_content(file_path)
            
            if not content.strip():
                return f"❌ 文档内容为空: {filename}"
            
            # 2. 生成文档ID
            doc_id = hashlib.md5(f"{filename}_{datetime.now().isoformat()}".encode()).hexdigest()
            
            # 3. 添加到向量库
            metadata = {
                "filename": filename,
                "file_path": file_path,
                "upload_time": datetime.now().isoformat(),
                "content_length": len(content)
            }
            
            chunks_count = self.vector_store.add_document(doc_id, content, metadata)
            
            result = f"✅ 文档上传成功！\n\n"
            result += f"📄 文件名: {filename}\n"
            result += f"🆔 文档ID: {doc_id}\n"
            result += f"📊 内容长度: {len(content)} 字符\n"
            result += f"🔢 分块数量: {chunks_count} 个\n"
            result += f"⏰ 上传时间: {metadata['upload_time']}\n"
            result += f"🔍 已完成embedding处理，可用于智能搜索"
            
            logger.info(f"✅ 文档 {filename} 处理完成")
            return result
            
        except Exception as e:
            logger.error(f"文档上传失败: {e}")
            return f"❌ 文档上传失败: {str(e)}"
    
    def _search_documents(self, query: str) -> str:
        """搜索相关文档内容"""
        try:
            logger.info(f"🔍 搜索查询: {query}")
            
            results = self.vector_store.search_documents(query, n_results=5)
            
            if not results:
                return f"🔍 搜索查询: {query}\n\n❌ 未找到相关内容"
            
            result = f"🔍 搜索查询: {query}\n\n"
            result += f"📋 找到 {len(results)} 条相关内容：\n\n"
            
            for i, doc in enumerate(results, 1):
                result += f"📄 结果 {i}:\n"
                result += f"   📝 内容: {doc['content'][:200]}{'...' if len(doc['content']) > 200 else ''}\n"
                result += f"   📁 来源: {doc['metadata'].get('filename', '未知')}\n"
                result += f"   🎯 相似度: {1 - doc['distance']:.3f}\n\n"
            
            return result
            
        except Exception as e:
            logger.error(f"搜索失败: {e}")
            return f"❌ 搜索失败: {str(e)}"
    
    def _fill_template_fields(self, template_fields_json: Dict[str, str]) -> str:
        """基于模板字段JSON进行智能填充 - 核心功能"""
        try:
            logger.info("🎯 开始模板字段智能填充")
            
            # 验证输入
            if not isinstance(template_fields_json, dict):
                return "❌ template_fields_json 必须是字典格式"
            
            if not template_fields_json:
                return "❌ template_fields_json 不能为空"
            
            # 执行智能填充
            filled_fields = self.field_processor.fill_template_fields(
                template_fields_json, self.vector_store
            )
            
            # 格式化返回结果
            result = f"✅ 模板字段智能填充完成！\n\n"
            result += f"📋 输入字段: {len(template_fields_json)} 个\n"
            result += f"📝 填充字段: {len(filled_fields)} 个\n\n"
            result += "📄 填充结果:\n"
            result += "=" * 50 + "\n"
            
            for field_name, filled_content in filled_fields.items():
                result += f"🔸 {field_name}:\n"
                result += f"   {filled_content[:100]}{'...' if len(filled_content) > 100 else ''}\n\n"
            
            result += "=" * 50 + "\n"
            result += f"💾 完整填充结果JSON:\n{json.dumps(filled_fields, ensure_ascii=False, indent=2)}"
            
            return result
            
        except Exception as e:
            logger.error(f"模板字段填充失败: {e}")
            return f"❌ 模板字段填充失败: {str(e)}"
    
    def _list_documents(self) -> str:
        """列出所有已上传的文档"""
        try:
            documents = self.vector_store.get_all_documents()
            
            if not documents:
                return "📚 当前没有已上传的文档"
            
            result = f"📚 已上传文档列表 (共 {len(documents)} 个):\n\n"
            
            for i, doc in enumerate(documents, 1):
                result += f"📄 文档 {i}:\n"
                result += f"   📁 文件名: {doc['metadata'].get('filename', '未知')}\n"
                result += f"   🆔 文档ID: {doc['id']}\n"
                result += f"   📊 分块数: {doc['chunks']} 个\n"
                result += f"   ⏰ 上传时间: {doc['metadata'].get('upload_time', '未知')}\n\n"
            
            return result
            
        except Exception as e:
            logger.error(f"获取文档列表失败: {e}")
            return f"❌ 获取文档列表失败: {str(e)}"
    
    def _clear_documents(self) -> str:
        """清空所有文档"""
        try:
            self.vector_store.clear_all()
            return "✅ 所有文档已清空"
        except Exception as e:
            logger.error(f"清空文档失败: {e}")
            return f"❌ 清空文档失败: {str(e)}" 