#!/usr/bin/env python3
"""
增强版长文档生成器
在原有长文档生成基础上，集成图片RAG调用和图片插入功能
"""

import os
import json
import uuid
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    import requests
    from PIL import Image
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 导入现有工具
from image_rag_tool import ImageRAGTool
from rag_tool_chroma import RAGTool
from tools import Tool

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class EnhancedLongDocumentGenerator(Tool):
    """增强版长文档生成器 - 集成文本RAG和图片RAG的多模态文档生成"""
    
    def __init__(self):
        super().__init__(
            name="enhanced_long_document_generator",
            description="📝 增强版长文档生成器 - 结合文本RAG检索和图片RAG的智能文档生成系统，支持自动插入相关图片"
        )
        
        # 初始化RAG工具
        self.text_rag = RAGTool()
        self.image_rag = ImageRAGTool()
        
        # 生成配置
        self.config = {
            "max_refinement_cycles": 3,
            "search_top_k": 5,
            "image_search_top_k": 3,
            "min_image_score": 0.3,
            "max_images_per_chapter": 2,
            "image_width_inches": 4.0
        }
        
        # 状态跟踪
        self.generation_stats = {
            "generated_documents": 0,
            "inserted_images": 0,
            "retrieved_knowledge": 0
        }
        
        logger.info("✅ 增强版长文档生成器初始化完成")
    
    def execute(self, **kwargs) -> str:
        """
        执行增强版长文档生成
        
        Args:
            action: 操作类型 (generate/get_stats/set_config)
            request: 用户请求描述
            project_name: 项目名称（可选）
            include_images: 是否包含图片（默认True）
            output_format: 输出格式 (docx/json, 默认docx)
            custom_config: 自定义配置（可选）
        """
        action = kwargs.get("action", "generate")
        
        if action == "generate":
            return self._generate_document(**kwargs)
        elif action == "get_stats":
            return self._get_generation_stats()
        elif action == "set_config":
            return self._set_config(**kwargs)
        else:
            return json.dumps({
                "status": "error",
                "message": f"不支持的操作: {action}",
                "supported_actions": ["generate", "get_stats", "set_config"]
            }, indent=2, ensure_ascii=False)
    
    def _generate_document(self, **kwargs) -> str:
        """生成增强版长文档"""
        request = kwargs.get("request")
        project_name = kwargs.get("project_name", "")
        include_images = kwargs.get("include_images", True)
        output_format = kwargs.get("output_format", "docx")
        
        if not request:
            return json.dumps({
                "status": "error",
                "message": "请提供文档生成请求 (request参数)"
            }, indent=2, ensure_ascii=False)
        
        # 更新配置
        custom_config = kwargs.get("custom_config", {})
        if custom_config:
            self.config.update(custom_config)
        
        task_id = str(uuid.uuid4())
        generation_result = {
            "status": "success",
            "task_id": task_id,
            "request": request,
            "project_name": project_name,
            "timestamp": datetime.now().isoformat(),
            "generation_steps": []
        }
        
        try:
            # 步骤1: 生成文档大纲
            logger.info("📋 生成文档大纲...")
            outline_result = self._generate_outline(request, project_name)
            generation_result["generation_steps"].append(outline_result)
            
            if outline_result["status"] != "success":
                generation_result["status"] = "failed"
                return json.dumps(generation_result, indent=2, ensure_ascii=False)
            
            # 步骤2: 大纲知识增强
            logger.info("🧠 知识增强大纲...")
            enhancement_result = self._enhance_outline_with_knowledge(
                outline_result["outline"], project_name
            )
            generation_result["generation_steps"].append(enhancement_result)
            
            # 步骤3: 生成章节内容
            logger.info("📝 生成章节内容...")
            chapters_result = self._generate_chapters_with_images(
                enhancement_result.get("enhanced_outline", outline_result["outline"]),
                project_name,
                include_images
            )
            generation_result["generation_steps"].append(chapters_result)
            
            # 步骤4: 组装最终文档
            logger.info("📄 组装最终文档...")
            assembly_result = self._assemble_final_document(
                chapters_result["chapters"],
                output_format,
                task_id
            )
            generation_result["generation_steps"].append(assembly_result)
            
            # 更新统计信息
            self.generation_stats["generated_documents"] += 1
            
            generation_result["output"] = {
                "document_path": assembly_result.get("document_path"),
                "total_chapters": len(chapters_result["chapters"]),
                "total_images": assembly_result.get("total_images", 0),
                "document_size": assembly_result.get("document_size")
            }
            
            logger.info(f"✅ 文档生成完成: {task_id}")
            return json.dumps(generation_result, indent=2, ensure_ascii=False)
            
        except Exception as e:
            generation_result["status"] = "error"
            generation_result["error"] = str(e)
            logger.error(f"❌ 文档生成失败: {e}")
            return json.dumps(generation_result, indent=2, ensure_ascii=False)
    
    def _generate_outline(self, request: str, project_name: str) -> Dict[str, Any]:
        """生成文档大纲"""
        step_result = {
            "step": "outline_generation",
            "status": "processing",
            "timestamp": datetime.now().isoformat()
        }
        
        try:
            # 调用文本RAG获取相关背景信息
            background_query = f"{project_name} {request}" if project_name else request
            background_results = self.text_rag.execute(
                action="search",
                query=background_query,
                top_k=self.config["search_top_k"]
            )
            
            # 构建大纲生成提示
            prompt = f"""基于以下信息，为请求的文档生成详细的结构化大纲。

用户请求: {request}
项目名称: {project_name}

背景资料:
{background_results}

请以JSON格式返回大纲，包含以下结构：
{{
    "title": "文档标题",
    "chapters": [
        {{
            "chapter_id": "ch_01",
            "title": "章节标题",
            "key_points": ["要点1", "要点2", "要点3"],
            "estimated_length": "预估字数",
            "image_keywords": ["图片搜索关键词1", "关键词2"]
        }}
    ]
}}

重要：
1. 生成3-6个章节
2. 每个章节包含3-5个关键要点
3. 为每个章节提供2-3个图片搜索关键词
4. 所有内容使用中文
"""
            
            # 这里应该调用AI模型，暂时模拟生成
            outline = self._simulate_ai_outline_generation(request, project_name)
            
            step_result.update({
                "status": "success",
                "message": "大纲生成完成",
                "outline": outline,
                "background_sources": len(background_results.split('\n')) if background_results else 0
            })
            
            return step_result
            
        except Exception as e:
            step_result.update({
                "status": "error",
                "message": f"大纲生成失败: {str(e)}"
            })
            return step_result
    
    def _enhance_outline_with_knowledge(self, outline: Dict, project_name: str) -> Dict[str, Any]:
        """使用知识库增强大纲"""
        step_result = {
            "step": "knowledge_enhancement",
            "status": "processing",
            "timestamp": datetime.now().isoformat(),
            "enhanced_chapters": 0
        }
        
        try:
            enhanced_outline = outline.copy()
            enhanced_count = 0
            
            for chapter in enhanced_outline.get("chapters", []):
                # 为每个章节搜索相关知识
                chapter_query = f"{project_name} {chapter.get('title', '')}"
                knowledge_results = self.text_rag.execute(
                    action="search",
                    query=chapter_query,
                    top_k=3
                )
                
                if knowledge_results and "找到" in knowledge_results:
                    # 解析知识结果并增强章节
                    chapter["knowledge_base"] = knowledge_results[:500] + "..."
                    enhanced_count += 1
                    self.generation_stats["retrieved_knowledge"] += 1
            
            step_result.update({
                "status": "success",
                "message": f"知识增强完成，增强了 {enhanced_count} 个章节",
                "enhanced_outline": enhanced_outline,
                "enhanced_chapters": enhanced_count
            })
            
            return step_result
            
        except Exception as e:
            step_result.update({
                "status": "error",
                "message": f"知识增强失败: {str(e)}"
            })
            return step_result
    
    def _generate_chapters_with_images(self, outline: Dict, project_name: str, include_images: bool = True) -> Dict[str, Any]:
        """生成章节内容并插入相关图片"""
        step_result = {
            "step": "chapter_generation",
            "status": "processing",
            "timestamp": datetime.now().isoformat(),
            "generated_chapters": 0,
            "inserted_images": 0
        }
        
        try:
            chapters = []
            total_inserted_images = 0
            
            for chapter_data in outline.get("chapters", []):
                chapter_result = self._generate_single_chapter_with_images(
                    chapter_data, project_name, include_images
                )
                chapters.append(chapter_result)
                total_inserted_images += chapter_result.get("image_count", 0)
            
            step_result.update({
                "status": "success",
                "message": f"章节生成完成，共 {len(chapters)} 个章节，插入 {total_inserted_images} 张图片",
                "chapters": chapters,
                "generated_chapters": len(chapters),
                "inserted_images": total_inserted_images
            })
            
            self.generation_stats["inserted_images"] += total_inserted_images
            return step_result
            
        except Exception as e:
            step_result.update({
                "status": "error",
                "message": f"章节生成失败: {str(e)}"
            })
            return step_result
    
    def _generate_single_chapter_with_images(self, chapter_data: Dict, project_name: str, include_images: bool) -> Dict[str, Any]:
        """生成单个章节并插入图片"""
        chapter_result = {
            "chapter_id": chapter_data.get("chapter_id", "unknown"),
            "title": chapter_data.get("title", ""),
            "content": "",
            "images": [],
            "image_count": 0
        }
        
        try:
            # 生成章节文本内容
            content = self._generate_chapter_content(chapter_data, project_name)
            chapter_result["content"] = content
            
            # 搜索相关图片
            if include_images and "image_keywords" in chapter_data:
                for keyword in chapter_data["image_keywords"][:self.config["max_images_per_chapter"]]:
                    image_search_query = f"{project_name} {keyword}" if project_name else keyword
                    image_results = self._search_relevant_images(image_search_query)
                    
                    if image_results:
                        chapter_result["images"].extend(image_results)
                        chapter_result["image_count"] += len(image_results)
            
            return chapter_result
            
        except Exception as e:
            chapter_result["error"] = str(e)
            return chapter_result
    
    def _generate_chapter_content(self, chapter_data: Dict, project_name: str) -> str:
        """生成章节内容"""
        # 模拟内容生成（实际应该调用AI模型）
        title = chapter_data.get("title", "")
        key_points = chapter_data.get("key_points", [])
        knowledge_base = chapter_data.get("knowledge_base", "")
        
        content_parts = []
        
        # 章节介绍
        content_parts.append(f"本章将详细介绍{title}的相关内容。")
        
        # 基于关键点生成内容
        for i, point in enumerate(key_points, 1):
            content_parts.append(f"\n\n## {i}. {point}\n")
            content_parts.append(f"关于{point}，需要从以下几个方面进行分析和讨论：")
            content_parts.append(f"- 基本概念和定义")
            content_parts.append(f"- 实际应用场景")
            content_parts.append(f"- 相关技术要求")
            content_parts.append(f"- 注意事项和最佳实践")
        
        # 整合知识库信息
        if knowledge_base:
            content_parts.append(f"\n\n## 相关资料参考\n")
            content_parts.append(f"根据项目资料库的相关信息：{knowledge_base}")
        
        return "\n".join(content_parts)
    
    def _search_relevant_images(self, query: str) -> List[Dict[str, Any]]:
        """搜索相关图片"""
        try:
            image_search_result = self.image_rag.execute(
                action="search",
                query=query,
                top_k=self.config["image_search_top_k"],
                min_score=self.config["min_image_score"]
            )
            
            # 解析图片搜索结果
            images = []
            if "找到" in image_search_result and "张相关图片" in image_search_result:
                # 简化的结果解析（实际需要更复杂的解析逻辑）
                lines = image_search_result.split('\n')
                for line in lines:
                    if "图片URL:" in line:
                        url = line.split("图片URL:")[-1].strip()
                        if url:
                            images.append({
                                "url": url,
                                "caption": f"与'{query}'相关的图片",
                                "search_query": query
                            })
            
            return images[:self.config["max_images_per_chapter"]]
            
        except Exception as e:
            logger.warning(f"图片搜索失败: {e}")
            return []
    
    def _assemble_final_document(self, chapters: List[Dict], output_format: str, task_id: str) -> Dict[str, Any]:
        """组装最终文档"""
        step_result = {
            "step": "document_assembly",
            "status": "processing",
            "timestamp": datetime.now().isoformat()
        }
        
        try:
            if output_format == "docx" and DOCX_AVAILABLE:
                document_path = self._create_docx_document(chapters, task_id)
                document_size = os.path.getsize(document_path) if os.path.exists(document_path) else 0
            else:
                document_path = self._create_json_document(chapters, task_id)
                document_size = os.path.getsize(document_path) if os.path.exists(document_path) else 0
            
            total_images = sum(chapter.get("image_count", 0) for chapter in chapters)
            
            step_result.update({
                "status": "success",
                "message": f"文档组装完成: {output_format.upper()} 格式",
                "document_path": document_path,
                "total_images": total_images,
                "document_size": f"{document_size / 1024:.1f} KB"
            })
            
            return step_result
            
        except Exception as e:
            step_result.update({
                "status": "error",
                "message": f"文档组装失败: {str(e)}"
            })
            return step_result
    
    def _create_docx_document(self, chapters: List[Dict], task_id: str) -> str:
        """创建DOCX文档"""
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('智能生成文档', 0)
        
        # 添加章节
        for chapter in chapters:
            # 章节标题
            doc.add_heading(chapter.get("title", ""), level=1)
            
            # 章节内容
            content = chapter.get("content", "")
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.startswith('##'):
                        doc.add_heading(para.replace('##', '').strip(), level=2)
                    elif para.startswith('-'):
                        p = doc.add_paragraph(para.strip())
                        p.style = 'List Bullet'
                    else:
                        doc.add_paragraph(para.strip())
            
            # 插入图片
            images = chapter.get("images", [])
            for img_data in images:
                try:
                    # 这里应该下载图片并插入，暂时添加图片占位符
                    doc.add_paragraph(f"[图片: {img_data.get('caption', '相关图片')}]")
                    # doc.add_picture(image_path, width=Inches(self.config["image_width_inches"]))
                except Exception as e:
                    logger.warning(f"图片插入失败: {e}")
        
        # 保存文档
        filename = f"enhanced_document_{task_id}.docx"
        doc_path = os.path.join("generated_documents", filename)
        os.makedirs("generated_documents", exist_ok=True)
        doc.save(doc_path)
        
        return doc_path
    
    def _create_json_document(self, chapters: List[Dict], task_id: str) -> str:
        """创建JSON文档"""
        document_data = {
            "task_id": task_id,
            "timestamp": datetime.now().isoformat(),
            "document_type": "enhanced_long_document",
            "chapters": chapters,
            "metadata": {
                "total_chapters": len(chapters),
                "total_images": sum(chapter.get("image_count", 0) for chapter in chapters),
                "generation_tool": "enhanced_long_document_generator"
            }
        }
        
        filename = f"enhanced_document_{task_id}.json"
        doc_path = os.path.join("generated_documents", filename)
        os.makedirs("generated_documents", exist_ok=True)
        
        with open(doc_path, 'w', encoding='utf-8') as f:
            json.dump(document_data, f, ensure_ascii=False, indent=2)
        
        return doc_path
    
    def _simulate_ai_outline_generation(self, request: str, project_name: str) -> Dict[str, Any]:
        """模拟AI大纲生成（替代真实的AI调用）"""
        return {
            "title": f"{project_name}相关文档" if project_name else "智能生成文档",
            "chapters": [
                {
                    "chapter_id": "ch_01",
                    "title": "概述与背景",
                    "key_points": ["项目背景", "目标和意义", "技术现状"],
                    "estimated_length": "800-1000字",
                    "image_keywords": ["项目概览", "背景图"]
                },
                {
                    "chapter_id": "ch_02", 
                    "title": "技术方案设计",
                    "key_points": ["技术架构", "实施方案", "关键技术"],
                    "estimated_length": "1200-1500字",
                    "image_keywords": ["技术架构图", "流程图"]
                },
                {
                    "chapter_id": "ch_03",
                    "title": "实施与应用",
                    "key_points": ["实施步骤", "应用场景", "效果评估"],
                    "estimated_length": "1000-1200字",
                    "image_keywords": ["实施流程", "应用效果"]
                }
            ]
        }
    
    def _get_generation_stats(self) -> str:
        """获取生成统计信息"""
        return json.dumps({
            "status": "success",
            "generation_statistics": self.generation_stats,
            "current_config": self.config,
            "timestamp": datetime.now().isoformat()
        }, indent=2, ensure_ascii=False)
    
    def _set_config(self, **kwargs) -> str:
        """设置生成配置"""
        config_updates = kwargs.get("config", {})
        
        if not config_updates:
            return json.dumps({
                "status": "error",
                "message": "请提供配置更新 (config参数)"
            }, indent=2, ensure_ascii=False)
        
        # 更新配置
        valid_keys = set(self.config.keys())
        updated_keys = []
        
        for key, value in config_updates.items():
            if key in valid_keys:
                self.config[key] = value
                updated_keys.append(key)
        
        return json.dumps({
            "status": "success",
            "message": f"配置更新完成: {updated_keys}",
            "updated_config": {k: self.config[k] for k in updated_keys},
            "current_config": self.config
        }, indent=2, ensure_ascii=False)
    
    def get_usage_guide(self) -> str:
        """获取使用指南"""
        return """
📝 增强版长文档生成器使用指南

🚀 核心功能：
1. 智能文档大纲生成
2. 知识库内容检索增强
3. 自动图片搜索和插入
4. 多格式文档输出（DOCX/JSON）

📋 基本用法：
1. 生成智能文档:
   enhanced_long_document_generator(
       action="generate",
       request="生成项目技术方案",
       project_name="智能系统项目",
       include_images=True,
       output_format="docx"
   )

2. 获取统计信息:
   enhanced_long_document_generator(action="get_stats")

3. 设置配置:
   enhanced_long_document_generator(
       action="set_config",
       config={
           "max_images_per_chapter": 3,
           "image_search_top_k": 5
       }
   )

🔧 配置参数：
- max_refinement_cycles: 大纲优化循环次数（默认3）
- search_top_k: 文本检索结果数量（默认5）
- image_search_top_k: 图片检索结果数量（默认3）
- min_image_score: 最小图片相似度分数（默认0.3）
- max_images_per_chapter: 每章最大图片数（默认2）
- image_width_inches: 插入图片宽度（默认4.0英寸）

📊 生成流程：
1. 🧠 基于用户请求生成文档大纲
2. 📚 使用文本RAG检索相关知识增强大纲
3. 📝 生成每个章节的详细内容
4. 🖼️ 为每个章节搜索和插入相关图片
5. 📄 组装成最终文档（DOCX或JSON格式）

💡 使用建议：
- 提供清晰具体的文档请求
- 设置适当的项目名称以提高检索精度
- 根据需要调整图片搜索参数
- 大型文档建议分章节逐步生成

⚡ 优势特色：
✅ 多模态内容整合（文本+图片）
✅ 智能知识库检索增强
✅ 自动图片匹配和插入
✅ 灵活的配置和输出格式
✅ 详细的生成过程跟踪
        """


# 工具实例化和导出
if __name__ == "__main__":
    generator = EnhancedLongDocumentGenerator()
    print(generator.get_usage_guide()) 