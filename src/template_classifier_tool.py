#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import shutil
import re
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import logging
from docx import Document
import PyPDF2
from .tools import Tool

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TemplateClassifierTool(Tool):
    """
    模板分类工具 - 智能判断文档类型并进行相应处理
    
    功能：
    1. 分析上传文档内容，判断是模板文档还是资料文档
    2. 模板文档：保存到templates文件夹，保留原始文件名
    3. 资料文档：调用RAG工具进行embedding处理
    """
    
    def __init__(self):
        super().__init__(
            name="template_classifier",
            description="模板分类工具 - 智能判断上传文档类型并进行相应处理"
        )
        
        # 创建模板存储目录
        self.templates_dir = Path("templates_storage")
        self.templates_dir.mkdir(exist_ok=True)
        
        # 模板文档特征关键词
        self.template_indicators = {
            # 空白字段标识符
            'blank_fields': [
                r'____+',  # 下划线空白
                r'＿＿+',   # 全角下划线
                r'\[.*?\]',  # 方括号占位符
                r'【.*?】',   # 中文方括号
                r'\(.*?\)',  # 圆括号占位符
                r'（.*?）',   # 中文圆括号
                r'\{.*?\}',  # 花括号占位符
            ],
            
            # 表格表单关键词
            'form_keywords': [
                '记录', '表格', '申请', '审批', '检查', '验收', 
                '登记', '填写', '签名', '日期', '编号',
                '项目名称', '负责人', '联系方式', '备注',
                '表1-1', '表格', '序号', '内容', '结果',
                '复核', '监理', '施工', '质检'
            ],
            
            # 模板文档标题特征
            'template_titles': [
                '记录表', '申请表', '审批表', '检查表', 
                '验收表', '登记表', '报告表', '统计表',
                '现场.*记录', '施工.*记录', '质量.*记录',
                '安全.*检查', '材料.*申请', '变更.*申请'
            ]
        }
        
        # 资料文档特征关键词
        self.document_indicators = {
            # 内容丰富的文档特征
            'content_keywords': [
                '根据', '按照', '依据', '为了', '通过',
                '分析', '研究', '调查', '统计', '评估',
                '方案', '计划', '报告', '总结', '建议',
                '技术', '工艺', '材料', '设备', '施工'
            ],
            
            # 长文档特征
            'document_structure': [
                '第.*章', '第.*节', '第.*条',
                '一、', '二、', '三、', '四、', '五、',
                '（一）', '（二）', '（三）',
                '1.', '2.', '3.', '4.', '5.'
            ]
        }
    
    def execute(self, **kwargs) -> str:
        """
        执行文档分类处理
        
        参数:
        - file_path: 上传文档的路径
        - action: 可选，指定操作类型 (classify/process)
        """
        try:
            file_path = kwargs.get('file_path')
            action = kwargs.get('action', 'classify')
            
            if not file_path:
                return "❌ 错误：未提供文件路径"
            
            if not os.path.exists(file_path):
                return f"❌ 错误：文件不存在 - {file_path}"
            
            if action == 'classify':
                return self._classify_and_process_document(file_path)
            else:
                return f"❌ 错误：不支持的操作类型 - {action}"
                
        except Exception as e:
            logger.error(f"模板分类工具执行错误: {e}")
            return f"❌ 模板分类工具执行错误: {str(e)}"
    
    def _classify_and_process_document(self, file_path: str) -> str:
        """
        分类并处理文档
        """
        try:
            # 提取文档内容
            content = self._extract_document_content(file_path)
            if not content:
                return f"❌ 无法提取文档内容: {file_path}"
            
            # 判断文档类型
            doc_type, confidence, features = self._classify_document_type(content, file_path)
            
            # 根据类型进行处理
            if doc_type == "template":
                result = self._handle_template_document(file_path, features)
                return f"📋 模板文档处理完成\n" \
                       f"文件: {os.path.basename(file_path)}\n" \
                       f"置信度: {confidence:.2f}\n" \
                       f"特征: {', '.join(features[:3])}\n" \
                       f"处理结果: {result}"
            else:
                result = self._handle_resource_document(file_path, features)
                return f"📚 资料文档处理完成\n" \
                       f"文件: {os.path.basename(file_path)}\n" \
                       f"置信度: {confidence:.2f}\n" \
                       f"特征: {', '.join(features[:3])}\n" \
                       f"处理结果: {result}"
                       
        except Exception as e:
            logger.error(f"文档分类处理错误: {e}")
            return f"❌ 文档分类处理错误: {str(e)}"
    
    def _extract_document_content(self, file_path: str) -> str:
        """
        提取文档内容
        """
        try:
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext in ['.docx']:
                return self._extract_docx_content(file_path)
            elif file_ext in ['.doc']:
                # 对于.doc文件，先尝试转换为.docx
                return self._extract_doc_content(file_path)
            elif file_ext in ['.pdf']:
                return self._extract_pdf_content(file_path)
            elif file_ext in ['.txt']:
                return self._extract_txt_content(file_path)
            else:
                logger.warning(f"不支持的文件格式: {file_ext}")
                return ""
                
        except Exception as e:
            logger.error(f"内容提取错误: {e}")
            return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """提取DOCX文档内容"""
        try:
            doc = Document(file_path)
            content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return "\n".join(content)
            
        except Exception as e:
            logger.error(f"DOCX内容提取错误: {e}")
            return ""
    
    def _extract_doc_content(self, file_path: str) -> str:
        """提取DOC文档内容（通过转换）"""
        try:
            # 这里可以集成之前的DOC转换功能
            # 暂时返回基础信息
            return f"DOC文件: {os.path.basename(file_path)}"
        except Exception as e:
            logger.error(f"DOC内容提取错误: {e}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """提取PDF文档内容"""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        content.append(text.strip())
            return "\n".join(content)
        except Exception as e:
            logger.error(f"PDF内容提取错误: {e}")
            return ""
    
    def _extract_txt_content(self, file_path: str) -> str:
        """提取TXT文档内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"TXT内容提取错误: {e}")
                return ""
    
    def _classify_document_type(self, content: str, file_path: str) -> Tuple[str, float, list]:
        """
        分类文档类型
        
        返回: (类型, 置信度, 特征列表)
        """
        template_score = 0
        document_score = 0
        features = []
        
        # 检查模板特征
        # 1. 空白字段检查
        blank_field_count = 0
        for pattern in self.template_indicators['blank_fields']:
            matches = re.findall(pattern, content)
            blank_field_count += len(matches)
        
        if blank_field_count > 0:
            template_score += min(blank_field_count * 2, 20)  # 最多20分
            features.append(f"空白字段({blank_field_count}个)")
        
        # 2. 表单关键词检查
        form_keyword_count = 0
        for keyword in self.template_indicators['form_keywords']:
            if keyword in content:
                form_keyword_count += 1
        
        if form_keyword_count > 0:
            template_score += min(form_keyword_count * 3, 30)  # 最多30分
            features.append(f"表单关键词({form_keyword_count}个)")
        
        # 3. 模板标题检查
        title_match_count = 0
        for pattern in self.template_indicators['template_titles']:
            if re.search(pattern, content):
                title_match_count += 1
        
        if title_match_count > 0:
            template_score += min(title_match_count * 10, 30)  # 最多30分
            features.append(f"模板标题特征({title_match_count}个)")
        
        # 4. 文件名检查
        filename = os.path.basename(file_path).lower()
        if any(keyword in filename for keyword in ['记录', '表', '申请', '审批', '检查']):
            template_score += 10
            features.append("文件名含模板特征")
        
        # 检查资料文档特征
        # 1. 内容关键词检查
        content_keyword_count = 0
        for keyword in self.document_indicators['content_keywords']:
            content_keyword_count += content.count(keyword)
        
        if content_keyword_count > 0:
            document_score += min(content_keyword_count * 1, 25)  # 最多25分
            features.append(f"内容关键词({content_keyword_count}个)")
        
        # 2. 文档结构检查
        structure_count = 0
        for pattern in self.document_indicators['document_structure']:
            structure_count += len(re.findall(pattern, content))
        
        if structure_count > 0:
            document_score += min(structure_count * 3, 25)  # 最多25分
            features.append(f"文档结构({structure_count}个)")
        
        # 3. 文档长度检查
        if len(content) > 1000:  # 长文档更可能是资料文档
            document_score += 10
            features.append("长文档内容")
        
        # 4. 段落数量检查
        paragraph_count = len([p for p in content.split('\n') if p.strip()])
        if paragraph_count > 10:
            document_score += 10
            features.append(f"多段落({paragraph_count}个)")
        
        # 决定文档类型
        total_score = template_score + document_score
        if total_score == 0:
            return "document", 0.5, ["无明显特征"]
        
        template_confidence = template_score / total_score
        document_confidence = document_score / total_score
        
        if template_confidence > 0.6:
            return "template", template_confidence, features
        else:
            return "document", document_confidence, features
    
    def _handle_template_document(self, file_path: str, features: list) -> str:
        """
        处理模板文档 - 保存到模板文件夹
        """
        try:
            filename = os.path.basename(file_path)
            destination = self.templates_dir / filename
            
            # 如果文件已存在，添加序号
            if destination.exists():
                name_part = destination.stem
                ext_part = destination.suffix
                counter = 1
                while destination.exists():
                    new_name = f"{name_part}_{counter}{ext_part}"
                    destination = self.templates_dir / new_name
                    counter += 1
            
            # 复制文件到模板目录
            shutil.copy2(file_path, destination)
            
            logger.info(f"模板文档已保存: {destination}")
            return f"已保存到模板库: {destination.name}"
            
        except Exception as e:
            logger.error(f"模板文档处理错误: {e}")
            return f"处理失败: {str(e)}"
    
    def _handle_resource_document(self, file_path: str, features: list) -> str:
        """
        处理资料文档 - 调用RAG工具
        """
        try:
            # 这里调用RAG工具进行处理
            from .rag_tool_chroma import RAGTool
            
            rag_tool = RAGTool()
            result = rag_tool.execute(
                operation="add_document",
                file_path=file_path
            )
            
            logger.info(f"资料文档已添加到RAG: {file_path}")
            return f"已添加到知识库: {result}"
            
        except Exception as e:
            logger.error(f"资料文档处理错误: {e}")
            return f"RAG处理失败: {str(e)}"
    
    def get_templates_list(self) -> list:
        """
        获取已保存的模板列表
        """
        try:
            templates = []
            for file_path in self.templates_dir.iterdir():
                if file_path.is_file():
                    templates.append({
                        'name': file_path.name,
                        'size': file_path.stat().st_size,
                        'modified': file_path.stat().st_mtime
                    })
            return templates
        except Exception as e:
            logger.error(f"获取模板列表错误: {e}")
            return []
    
    def get_classification_stats(self) -> dict:
        """
        获取分类统计信息
        """
        try:
            templates_count = len(list(self.templates_dir.iterdir()))
            return {
                'templates_stored': templates_count,
                'templates_directory': str(self.templates_dir),
                'supported_formats': ['.docx', '.doc', '.pdf', '.txt']
            }
        except Exception as e:
            logger.error(f"获取统计信息错误: {e}")
            return {}

def create_template_classifier_tool() -> TemplateClassifierTool:
    """创建模板分类工具实例"""
    return TemplateClassifierTool()

if __name__ == "__main__":
    # 测试代码
    tool = create_template_classifier_tool()
    print("模板分类工具初始化完成")
    print(f"模板存储目录: {tool.templates_dir}") 